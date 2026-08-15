# -*- coding: utf-8 -*-
"""Export AZ stories for translation and import target-locale DOCX packages.

Export:
  python tools/i18n_story_pipeline.py export

Import a translated JSON package (list of {stem, title, paragraphs}):
  python tools/i18n_story_pipeline.py import --lang en path/to/batch.json
  python tools/i18n_story_pipeline.py import --lang tr path/to/batch.json
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(Path(__file__).resolve().parent))
from i18n_config import TARGET_LANGS  # noqa: E402
AZ_STORIES = ROOT / "az" / "stories"
MAP_JSON = ROOT / "tools" / "story-mapping.json"
MANIFEST = ROOT / "docs" / "i18n" / "translation_manifest.json"


def read_docx(path: Path) -> tuple[str, list[str]]:
    paras = [
        (p.text or "").strip()
        for p in Document(str(path)).paragraphs
        if (p.text or "").strip()
    ]
    if not paras:
        return "", []
    return paras[0], paras[1:]


def write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    doc = Document()
    doc.add_paragraph(title)
    for para in paragraphs:
        doc.add_paragraph(para)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def load_manifest() -> dict:
    if MANIFEST.is_file():
        return json.loads(MANIFEST.read_text(encoding="utf-8"))
    return {"stems": {}}


def save_manifest(data: dict) -> None:
    MANIFEST.parent.mkdir(parents=True, exist_ok=True)
    MANIFEST.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def cmd_export(out_dir: Path) -> None:
    mapping = json.loads(MAP_JSON.read_text(encoding="utf-8"))
    out_dir.mkdir(parents=True, exist_ok=True)
    rows = []
    for row in mapping["rows"]:
        stem = row["en_stem"]
        path = AZ_STORIES / f"{stem}.docx"
        if not path.is_file():
            continue
        title, paragraphs = read_docx(path)
        item = {
            "stem": stem,
            "az_title": title,
            "paragraphs": paragraphs,
        }
        for lang in TARGET_LANGS:
            item[f"{lang}_title"] = ""
            item[f"{lang}_paragraphs"] = []
        rows.append(item)
        (out_dir / f"{stem}.json").write_text(
            json.dumps(item, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
        )
    index = {"count": len(rows), "stems": [r["stem"] for r in rows]}
    (out_dir / "_index.json").write_text(
        json.dumps(index, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    print(f"Exported {len(rows)} stories -> {out_dir}")


def cmd_import(lang: str, package: Path) -> None:
    if lang not in TARGET_LANGS:
        raise SystemExit(f"--lang must be one of: {', '.join(TARGET_LANGS)}")
    data = json.loads(package.read_text(encoding="utf-8"))
    if isinstance(data, dict) and "stories" in data:
        stories = data["stories"]
    elif isinstance(data, list):
        stories = data
    else:
        raise SystemExit("Package must be a list or {stories:[...]}")

    dest = ROOT / lang / "stories"
    dest.mkdir(parents=True, exist_ok=True)
    manifest = load_manifest()
    stems = manifest.setdefault("stems", {})
    n = 0
    for item in stories:
        stem = item["stem"]
        title = item.get("title") or item.get(f"{lang}_title")
        paragraphs = item.get("paragraphs") or item.get(f"{lang}_paragraphs") or []
        if not title or not paragraphs:
            print(f"skip incomplete {stem}")
            continue
        write_docx(dest / f"{stem}.docx", title, paragraphs)
        meta = stems.setdefault(stem, {})
        meta[f"text_{lang}"] = "done"
        meta[f"audio_{lang}"] = meta.get(f"audio_{lang}") or "pending"
        meta[f"illustration_{lang}"] = meta.get(f"illustration_{lang}") or "pending"
        n += 1
    save_manifest(manifest)
    print(f"Imported {n} {lang} stories -> {dest}")
    print(f"Manifest: {MANIFEST}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_ex = sub.add_parser("export", help="Export AZ stories for translation")
    p_ex.add_argument("--out", type=Path, default=ROOT / "docs" / "i18n" / "export")

    p_im = sub.add_parser("import", help="Import translated stories into a target locale tree")
    p_im.add_argument("--lang", required=True, choices=TARGET_LANGS)
    p_im.add_argument("package", type=Path)

    args = parser.parse_args()
    if args.cmd == "export":
        cmd_export(args.out)
    else:
        cmd_import(args.lang, args.package)


if __name__ == "__main__":
    main()
