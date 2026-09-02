# -*- coding: utf-8 -*-
"""Local-only story edit API: save title/body/moral to DOCX + published assets.

Run:  python tools/dev_story_edit_server.py
API:  http://127.0.0.1:8768/api/dev/ping
      http://127.0.0.1:8768/api/dev/save-story  (POST JSON)

Only binds to 127.0.0.1. Intended for development, never production.
"""
from __future__ import annotations

import html
import json
import re
import sys
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
HOST = "127.0.0.1"
PORT = 8768
LANGS = frozenset({"az", "en", "ru", "ky"})
STEM_RE = re.compile(r"^[a-z0-9]+(?:-[a-z0-9]+)*$")
MORAL_RE = re.compile(r"^(ibrət|ibret|moral|мораль|үлгү|сабак)\s*:", re.I)
SOURCE_RE = re.compile(
    r"(internet\s+sources|internet\s+mənb|открыт\w*\s+источник|интернет|"
    r"(?:source|mənbə|kaynak|источник|булак|булагы)\s*:|"
    r"hekayə\s+açıq|окуя\s+ачык|история\s+взята|story\s+was\s+obtained|"
    r"hekayələr\s+açıq|рассказы\s+взяты|stories\s+are\s+taken|"
    r"аңгемелер\s+ачык)",
    re.I,
)
SITE_SOURCE = {
    "az": "Hekayələr açıq internet mənbələrindən alınmış, illüstrasiyalar isə süni intellektlə yaradılmışdır.",
    "en": "Stories are taken from open Internet sources; illustrations were created with AI.",
    "ru": "Рассказы взяты из открытых интернет-источников, а иллюстрации созданы с помощью ИИ.",
    "ky": "Аңгемелер ачык интернет булактарынан алынган, иллюстрациялар болсо ЖИ менен түзүлгөн.",
}


def _set_paragraph_text(paragraph, text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _visible_paragraphs(doc: Document) -> list:
    return [p for p in doc.paragraphs if (p.text or "").strip()]


def update_docx(path: Path, title: str, body: list[str], moral: str) -> None:
    if not path.is_file():
        raise FileNotFoundError(f"missing docx: {path}")
    doc = Document(str(path))
    visible = _visible_paragraphs(doc)
    if len(visible) < 2:
        raise ValueError("docx has too few paragraphs")

    last = visible[-1].text.strip()
    keep_source = bool(SOURCE_RE.search(last)) and not MORAL_RE.match(last)
    source_text = last if keep_source else None

    target = [title, *[p for p in body if p.strip()], moral]
    if keep_source and source_text:
        target.append(source_text)

    # Rewrite existing visible paragraphs; blank extras; append if needed.
    for idx, para in enumerate(visible):
        if idx < len(target):
            _set_paragraph_text(para, target[idx])
        else:
            _set_paragraph_text(para, "")

    if len(visible) < len(target):
        # Append after the last document paragraph.
        for text in target[len(visible) :]:
            doc.add_paragraph(text)

    doc.save(str(path))


def _load_json_assignment(path: Path, prefix: str) -> tuple[str, Any, str]:
    raw = path.read_text(encoding="utf-8")
    m = re.match(rf"({re.escape(prefix)}\s*=\s*)(\{{.*\}}|\[.*\])(\s*;?\s*)$", raw, re.S)
    if not m:
        raise ValueError(f"cannot parse {path.name}")
    return m.group(1), json.loads(m.group(2)), m.group(3) or ";\n"


def update_stories_data(lang: str, stem: str, title: str, body: list[str], moral: str) -> str | None:
    path = ROOT / lang / "assets" / "stories-data.js"
    prefix, data, suffix = _load_json_assignment(path, "window.__BIRINCI_STORIES__")
    category_slug = None
    found = False
    source = SITE_SOURCE[lang]
    paragraphs = list(body) + [moral, source]
    for cat in data.get("categories") or []:
        for story in cat.get("stories") or []:
            if story.get("stem") != stem:
                continue
            story["title"] = title
            story["paragraphs"] = paragraphs
            category_slug = cat.get("slug")
            found = True
            break
        if found:
            break
    if not found:
        raise ValueError(f"stem {stem!r} not in {lang} stories-data")
    path.write_text(
        prefix + json.dumps(data, ensure_ascii=False, separators=(",", ":")) + suffix,
        encoding="utf-8",
        newline="\n",
    )
    return category_slug


def update_search_index(lang: str, stem: str, title: str, body: list[str], moral: str, category: str | None) -> None:
    path = ROOT / lang / "assets" / "search-index.js"
    prefix, entries, suffix = _load_json_assignment(path, "window.__BIRINCI_SEARCH__")
    source = SITE_SOURCE[lang]
    hay = " ".join(p for p in [title, category or "", *body, moral, source] if p).lower()
    found = False
    for entry in entries:
        if entry.get("stem") != stem:
            continue
        entry["title"] = title
        entry["hay"] = hay
        if category:
            entry["category"] = category
        found = True
        break
    if not found:
        raise ValueError(f"stem {stem!r} not in {lang} search-index")
    path.write_text(
        prefix + json.dumps(entries, ensure_ascii=False, separators=(",", ":")) + suffix,
        encoding="utf-8",
        newline="\n",
    )


def _story_body_html(body: list[str], moral: str, lang: str) -> str:
    parts = [f"<p>{html.escape(p)}</p>" for p in body]
    parts.append(f'<p class="story__moral">{html.escape(moral)}</p>')
    parts.append(f'<p class="story__source">{html.escape(SITE_SOURCE[lang])}</p>')
    return "".join(parts)


def _replace_story_text_in_html(text: str, stem: str, body_html: str) -> str:
    """Replace .story__text for a stem.

    Category pages often use id=\"text-{stem}\", but many older articles use a
    prefixed id such as text-the-{stem} / text-a-{stem}. Match by article first.
    """
    exact = re.compile(
        rf'(<div class="story__text card-text" id="text-{re.escape(stem)}">)([\s\S]*?)(</div>)'
    )
    text2, n = exact.subn(
        rf"\1\n          {body_html}\n        \3",
        text,
        count=1,
    )
    if n == 1:
        return text2

    art = re.search(
        rf'<article class="story news-card" id="{re.escape(stem)}"[^>]*>[\s\S]*?</article>',
        text,
    )
    if not art:
        raise ValueError(f"story article not found for {stem}")
    article = art.group(0)
    article2, n = re.subn(
        r'(<div class="story__text card-text" id="text-[^"]+">)([\s\S]*?)(</div>)',
        rf"\1\n          {body_html}\n        \3",
        article,
        count=1,
    )
    if n != 1:
        raise ValueError(f"story text block not found for {stem}")
    return text[: art.start()] + article2 + text[art.end() :]


def update_category_html(
    lang: str,
    slug: str,
    stem: str,
    title: str,
    body: list[str],
    moral: str,
    old_title: str | None,
) -> None:
    path = ROOT / lang / "categories" / f"{slug}.html"
    if not path.is_file():
        raise FileNotFoundError(f"missing category page: {path}")
    text = path.read_text(encoding="utf-8")
    blurb = body[0] if body else title
    body_html = _story_body_html(body, moral, lang)

    text = _replace_story_text_in_html(text, stem, body_html)

    text = re.sub(
        rf'(data-stem="{re.escape(stem)}" data-title=")([^"]*)(")',
        rf'\1{html.escape(title, quote=True)}\3',
        text,
    )
    text = re.sub(
        rf'(id="{re.escape(stem)}"[^>]*data-title=")([^"]*)(")',
        rf'\1{html.escape(title, quote=True)}\3',
        text,
    )
    text = re.sub(
        rf'(<li data-stem="{re.escape(stem)}"[^>]*>\s*<a href="#{re.escape(stem)}">)([^<]*)(</a>)',
        rf"\1{html.escape(title)}\3",
        text,
    )
    text = re.sub(
        rf'(href="#{re.escape(stem)}"[^>]*data-blurb=")([^"]*)(")',
        rf'\1{html.escape(blurb, quote=True)}\3',
        text,
    )
    text, n = re.subn(
        rf'(<a class="cat-card page-card" href="#{re.escape(stem)}"[\s\S]*?<h2 class="card-title">)([^<]*)(</h2>\s*<div class="card-desc">)([^<]*)(</div>)',
        rf"\1{html.escape(title)}\3{html.escape(blurb)}\5",
        text,
        count=1,
    )
    text, n2 = re.subn(
        rf'(<article class="story news-card" id="{re.escape(stem)}"[\s\S]*?<h2 class="card-title story__title">)([^<]*)(</h2>)',
        rf"\1{html.escape(title)}\3",
        text,
        count=1,
    )
    if n2 != 1:
        raise ValueError(f"article title not found for {stem}")

    if old_title and old_title != title:
        text = text.replace(old_title, title)

    path.write_text(text, encoding="utf-8", newline="\n")


def update_sitemap(lang: str, stem: str, slug: str, title: str) -> None:
    path = ROOT / lang / "sitemap.html"
    if not path.is_file():
        return
    text = path.read_text(encoding="utf-8")
    text2, n = re.subn(
        rf'(href="categories/{re.escape(slug)}\.html#{re.escape(stem)}"><span>)([^<]*)(</span>)',
        rf"\1{html.escape(title)}\3",
        text,
        count=1,
    )
    if n:
        path.write_text(text2, encoding="utf-8", newline="\n")


def category_title_for_slug(lang: str, slug: str) -> str | None:
    path = ROOT / lang / "assets" / "stories-data.js"
    _, data, _ = _load_json_assignment(path, "window.__BIRINCI_STORIES__")
    for cat in data.get("categories") or []:
        if cat.get("slug") == slug:
            return cat.get("title")
    return None


def find_story_meta(lang: str, stem: str) -> tuple[str, str]:
    """Return (category_slug, current_title)."""
    path = ROOT / lang / "assets" / "stories-data.js"
    _, data, _ = _load_json_assignment(path, "window.__BIRINCI_STORIES__")
    for cat in data.get("categories") or []:
        for story in cat.get("stories") or []:
            if story.get("stem") == stem:
                return str(cat.get("slug") or ""), str(story.get("title") or "")
    raise ValueError(f"stem {stem!r} not found in {lang}")


def save_story(payload: dict[str, Any]) -> dict[str, Any]:
    lang = str(payload.get("lang") or "").strip().lower()
    stem = str(payload.get("stem") or "").strip().lower()
    title = " ".join(str(payload.get("title") or "").split()).strip()
    moral = " ".join(str(payload.get("moral") or "").split()).strip()
    raw_body = payload.get("body") or []
    if not isinstance(raw_body, list):
        raise ValueError("body must be a list of strings")
    body = [" ".join(str(p).split()).strip() for p in raw_body]
    body = [p for p in body if p]

    if lang not in LANGS:
        raise ValueError("invalid lang")
    if not STEM_RE.match(stem):
        raise ValueError("invalid stem")
    if not title:
        raise ValueError("title required")
    if not body:
        raise ValueError("body required")
    if not moral:
        raise ValueError("moral required")

    slug, old_title = find_story_meta(lang, stem)
    if not slug:
        raise ValueError("category slug missing")

    docx_path = ROOT / lang / "wisdom-stories" / f"{stem}.docx"
    update_docx(docx_path, title, body, moral)
    update_stories_data(lang, stem, title, body, moral)
    cat_title = category_title_for_slug(lang, slug)
    update_search_index(lang, stem, title, body, moral, cat_title)
    update_category_html(lang, slug, stem, title, body, moral, old_title)
    update_sitemap(lang, stem, slug, title)

    return {
        "ok": True,
        "lang": lang,
        "stem": stem,
        "title": title,
        "slug": slug,
        "docx": str(docx_path.relative_to(ROOT)).replace("\\", "/"),
        "body_paragraphs": len(body),
    }


class Handler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):  # noqa: A003
        sys.stderr.write("%s - %s\n" % (self.address_string(), format % args))

    def _cors(self) -> None:
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def _json(self, code: int, payload: dict[str, Any]) -> None:
        raw = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(code)
        self._cors()
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def do_OPTIONS(self) -> None:  # noqa: N802
        self.send_response(204)
        self._cors()
        self.end_headers()

    def do_GET(self) -> None:  # noqa: N802
        path = self.path.split("?", 1)[0]
        if path in {"/api/dev/ping", "/ping"}:
            self._json(200, {"ok": True, "service": "dev-story-edit", "port": PORT})
            return
        self._json(404, {"ok": False, "error": "not found"})

    def do_POST(self) -> None:  # noqa: N802
        path = self.path.split("?", 1)[0]
        if path not in {"/api/dev/save-story", "/save-story"}:
            self._json(404, {"ok": False, "error": "not found"})
            return
        length = int(self.headers.get("Content-Length") or 0)
        if length <= 0 or length > 500_000:
            self._json(400, {"ok": False, "error": "bad content-length"})
            return
        try:
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
            if not isinstance(payload, dict):
                raise ValueError("JSON object required")
            result = save_story(payload)
            self._json(200, result)
        except Exception as exc:  # noqa: BLE001 — return to browser
            self._json(400, {"ok": False, "error": str(exc)})


def main() -> None:
    # Refuse non-loopback bind by construction.
    try:
        server = ThreadingHTTPServer((HOST, PORT), Handler)
    except OSError as exc:
        print(
            f"ERROR: could not bind http://{HOST}:{PORT} ({exc}).\n"
            "Another process is already using that port — stop old "
            "dev_story_edit_server.py instances and try again.",
            file=sys.stderr,
        )
        raise SystemExit(1) from exc
    print(f"dev story edit API on http://{HOST}:{PORT}")
    print("  GET  /api/dev/ping")
    print("  POST /api/dev/save-story")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nstopped")


if __name__ == "__main__":
    main()
