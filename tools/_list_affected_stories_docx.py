# -*- coding: utf-8 -*-
from __future__ import annotations

import json
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
stats = json.loads((ROOT / "docs" / "_az_lang_analysis.json").read_text(encoding="utf-8"))
OUT = ROOT / "docs" / "STORIES_AFFECTED_BY_LANGUAGE_FIXES.docx"


def font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def p(doc, text, *, size=11, bold=False, center=False):
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    font(run, size=size, bold=bold)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    return para


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                font(run, size=9, bold=True)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
            for paragraph in cells[c_i].paragraphs:
                for run in paragraph.runs:
                    font(run, size=9)
    doc.add_paragraph()


def main():
    long = OrderedDict()
    for item in stats["long_paras"]:
        prev = long.get(item["stem"])
        if not prev or item["len"] > prev["len"]:
            long[item["stem"]] = item

    mixed = OrderedDict()
    for item in stats["mixed"]:
        mixed[item["stem"]] = item

    both = [stem for stem in long if stem in mixed]
    tier2 = OrderedDict((stem, item) for stem, item in mixed.items() if stem not in long)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Stories affected by recommended language modifications")
    font(r, size=16, bold=True, color=RGBColor(0x00, 0x3C, 0x6E))

    p(
        doc,
        "Birİnci · Derived from the Azerbaijani language review · 13 August 2026",
        size=10,
        center=True,
    )

    doc.add_heading("Regeneration rule of thumb", level=1)
    add_table(
        doc,
        ["Change type", "Illustrations (.webp)", "Audio (.mp3)"],
        [
            [
                "Punctuation / quote style only (—, «», ASCII)",
                "No",
                "No (spoken words unchanged)",
            ],
            [
                "Paragraph splits without rephrasing",
                "No",
                "Usually no",
            ],
            [
                "Wording / fluency rewrite",
                "Only if scene/meaning changes",
                "Yes",
            ],
            [
                "New scenes or character actions",
                "Yes",
                "Yes",
            ],
        ],
    )

    p(
        doc,
        f"Summary: Tier 1 = {len(long)} stories (structure / possible rewrite). "
        f"Tier 2 = {len(tier2)} stories (dialogue punctuation only). "
        f"{len(both)} stories appear in both tiers.",
    )
    p(
        doc,
        "If you only apply punctuation + paragraph-split recommendations without rewriting meaning, expect:",
    )
    p(doc, "• Illustrations to regenerate: 0 (unless story content later changes)")
    p(
        doc,
        f"• Audio to regenerate: mainly the Tier 1 set ({len(long)}), and only if phrasing changes when splitting",
    )
    p(
        doc,
        f"• Tier 2 ({len(tier2)}): text files only — no illustration/audio regen required for punctuation unification",
    )

    doc.add_heading(
        f"Tier 1 — Long paragraphs (highest chance of text change) — {len(long)} stories",
        level=1,
    )
    rows = []
    for i, (stem, item) in enumerate(long.items(), 1):
        rows.append(
            [
                i,
                item["cat"],
                item["title"],
                stem,
                item["len"],
                "yes" if stem in mixed else "no",
            ]
        )
    add_table(
        doc,
        ["#", "Category", "Title", "Stem", "Longest para (chars)", "Also mixed dialogue?"],
        rows,
    )

    doc.add_heading("Tier 1 file paths", level=2)
    for stem in long:
        p(doc, f"• az/illustrations/{stem}.webp", size=10)
        p(doc, f"• az/audio/{stem}.mp3", size=10)

    doc.add_heading(
        f"Tier 2 — Mixed dialogue punctuation only (not in Tier 1) — {len(tier2)} stories",
        level=1,
    )
    rows2 = []
    for i, (stem, item) in enumerate(tier2.items(), 1):
        rows2.append(
            [
                i,
                item["cat"],
                item["title"],
                stem,
                ", ".join(item["styles"]),
            ]
        )
    add_table(
        doc,
        ["#", "Category", "Title", "Stem", "Styles mixed"],
        rows2,
    )

    doc.add_heading("Stem lists (copy/paste)", level=1)
    doc.add_heading("Tier 1 stems", level=2)
    p(doc, "\n".join(long.keys()), size=9)
    doc.add_heading("Tier 2-only stems", level=2)
    p(doc, "\n".join(tier2.keys()), size=9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
