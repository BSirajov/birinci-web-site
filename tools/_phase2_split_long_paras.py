# -*- coding: utf-8 -*-
"""Phase 2: split long paragraphs in Tier-1 stories at sentence boundaries (no wording change)."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
STORIES_DIR = ROOT / "az" / "stories"
DATA_JSON = ROOT / "az" / "data" / "stories.json"
ANALYSIS = ROOT / "docs" / "_az_lang_analysis.json"
REPORT = ROOT / "docs" / "PHASE2_TIER1_SPLITS_REPORT.md"
THRESHOLD = 600

TIER1_ORDER = [
    "teacher-hello-do-you-remember-me",
    "the-mother-of-pearl-flower",
    "the-former-minister-at-the-seminar",
    "the-mullah-and-the-scholar",
    "the-blind-well",
    "the-road-to-the-cotton-field",
    "discussion-and-conflict",
    "how-to-ward-off-insults",
    "why-people-shout-when-they-argue",
    "not-leaving-the-right-path",
]


def _ranges_inside_guillemets(text: str) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    i = 0
    while True:
        a = text.find("«", i)
        if a < 0:
            break
        b = text.find("»", a + 1)
        if b < 0:
            break
        ranges.append((a, b + 1))
        i = b + 1
    return ranges


def _inside(pos: int, ranges: list[tuple[int, int]]) -> bool:
    return any(a <= pos < b for a, b in ranges)


def sentence_ends(text: str) -> list[int]:
    """Return indices just after sentence-ending punctuation (split points)."""
    guillemets = _ranges_inside_guillemets(text)
    ends: list[int] = []
    for m in re.finditer(r"([.!?…]+)(\s+)(?=[«\"“—A-ZƏĞÜŞÖÇIİÁÉ])", text):
        end = m.end(1)
        if _inside(end - 1, guillemets):
            continue
        ends.append(end)
    return ends


def split_long_paragraph(text: str, max_len: int = THRESHOLD) -> list[str]:
    """Split text into chunks preferably <= max_len at sentence boundaries."""
    text = text.strip()
    if len(text) <= max_len:
        return [text]

    ends = sentence_ends(text)
    if not ends:
        for m in re.finditer(r"([;:])(\s+)", text):
            ends.append(m.end(1))
    if not ends:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        remaining = text[start:].strip()
        if len(remaining) <= max_len:
            if remaining:
                chunks.append(remaining)
            break

        window_end = start + max_len
        candidates = [e for e in ends if start < e <= window_end]
        if not candidates:
            after = [e for e in ends if e > window_end]
            if not after:
                chunks.append(remaining)
                break
            cut = after[0]
        else:
            cut = candidates[-1]

        piece = text[start:cut].strip()
        if piece:
            chunks.append(piece)
        start = cut
        while start < len(text) and text[start].isspace():
            start += 1

    if not chunks:
        return [text]

    # Merge a tiny trailing fragment into the previous chunk when safe.
    merged: list[str] = []
    for ch in chunks:
        if merged and len(ch) < 100 and len(merged[-1]) + 1 + len(ch) <= int(max_len * 1.25):
            merged[-1] = f"{merged[-1]} {ch}".strip()
        else:
            merged.append(ch)
    return merged


def load_title(stem: str) -> str:
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    for c in data["categories"]:
        for s in c["stories"]:
            if s["stem"] == stem:
                return s["title"]
    return stem


def process_stem(stem: str) -> dict:
    path = STORIES_DIR / f"{stem}.docx"
    title = load_title(stem)
    result = {
        "stem": stem,
        "title": title,
        "path": str(path),
        "exists": path.is_file(),
        "splits": [],
        "saved": False,
    }
    if not path.is_file():
        return result

    doc = Document(str(path))
    # Rebuild paragraph list: when we split, we need to insert new paragraphs
    # python-docx doesn't make insert easy; rebuild document body paragraphs carefully.
    # Approach: collect all paragraph texts, transform, then rewrite all paragraphs.

    old_paras = [(p.text or "") for p in doc.paragraphs]
    new_paras: list[str] = []
    for raw in old_paras:
        stripped = raw.strip()
        if not stripped:
            # preserve empty paras lightly — skip empties to avoid bloat
            continue
        if len(stripped) >= THRESHOLD:
            parts = split_long_paragraph(stripped)
            if len(parts) > 1:
                result["splits"].append(
                    {
                        "original_len": len(stripped),
                        "parts": len(parts),
                        "part_lens": [len(x) for x in parts],
                        "preview": [x[:140] for x in parts],
                    }
                )
                new_paras.extend(parts)
            else:
                new_paras.append(stripped)
        else:
            new_paras.append(stripped)

    if not result["splits"]:
        return result

    # Clear and rewrite: simplest reliable way — clear all paragraph texts then set first N, add rest
    # Keep title paragraph if first para equals title
    body_paras = list(doc.paragraphs)
    # Set existing paragraphs
    for i, p in enumerate(body_paras):
        if i < len(new_paras):
            p.text = new_paras[i]
        else:
            p.text = ""

    # Add remaining paragraphs
    for i in range(len(body_paras), len(new_paras)):
        doc.add_paragraph(new_paras[i])

    # If we had more old paras than new, trailing empties remain — remove empty trailing by clearing
    doc.save(str(path))
    result["saved"] = True
    result["new_para_count"] = len(new_paras)
    return result


def sync_json(stems: set[str]) -> int:
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    updated = 0
    for cat in data["categories"]:
        for story in cat["stories"]:
            stem = story["stem"]
            if stem not in stems:
                continue
            path = STORIES_DIR / f"{stem}.docx"
            if not path.is_file():
                continue
            doc = Document(str(path))
            title = story["title"]
            paras = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    paras.append(t)
            while paras and paras[0].casefold() == title.casefold():
                paras = paras[1:]
            story["paragraphs"] = paras
            updated += 1
    DATA_JSON.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return updated


def write_report(results: list[dict]) -> None:
    changed = [r for r in results if r.get("saved")]
    skipped = [r for r in results if r.get("exists") and not r.get("saved")]
    lines = []
    lines.append("# Phase 2 · Tier-1 long paragraph splits")
    lines.append("")
    lines.append("House style: `docs/HOUSE_STYLE_DIALOGUE.md`")
    lines.append("")
    lines.append("## Method")
    lines.append("")
    lines.append("- Split paragraphs ≥600 characters at sentence boundaries")
    lines.append("- **No wording changes** (punctuation of sentences preserved)")
    lines.append("- Illustrations: unchanged")
    lines.append("- Audio: **review/regenerate** for modified stems (paragraph timing may change; wording same)")
    lines.append("")
    lines.append(f"- Stories processed: **{len(results)}**")
    lines.append(f"- Stories modified: **{len(changed)}**")
    lines.append(f"- No split needed / could not auto-split: **{len(skipped)}**")
    lines.append("")

    lines.append("## Modified stories")
    lines.append("")
    for r in changed:
        lines.append(f"### {r['title']} (`{r['stem']}`)")
        lines.append("")
        for s in r["splits"]:
            lines.append(
                f"- Original length: {s['original_len']} → {s['parts']} parts {s['part_lens']}"
            )
            for i, prev in enumerate(s["preview"], 1):
                lines.append(f"  - part {i}: {prev}…")
        lines.append("")

    if skipped:
        lines.append("## Unchanged in this pass")
        lines.append("")
        for r in skipped:
            lines.append(f"- `{r['stem']}` — {r['title']}")
        lines.append("")

    lines.append("## Audio regeneration list (modified stems)")
    lines.append("")
    lines.append("```")
    for r in changed:
        lines.append(r["stem"])
    lines.append("```")
    lines.append("")
    lines.append("## Next")
    lines.append("")
    lines.append("- Spot-check the modified stories")
    lines.append("- Optionally regenerate audio for the stems listed above")
    lines.append("- Illustrations: keep unless you later rewrite scenes")
    lines.append("")

    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    # Prefer stems from analysis long_paras unique set, ordered
    stats = json.loads(ANALYSIS.read_text(encoding="utf-8"))
    long_stems = []
    seen = set()
    for item in stats["long_paras"]:
        if item["stem"] not in seen:
            seen.add(item["stem"])
            long_stems.append(item["stem"])
    # Use canonical order where possible
    stems = [s for s in TIER1_ORDER if s in seen] + [s for s in long_stems if s not in TIER1_ORDER]

    results = [process_stem(stem) for stem in stems]
    changed = {r["stem"] for r in results if r.get("saved")}
    synced = sync_json(changed) if changed else 0
    write_report(results)
    print(f"processed={len(results)} modified={len(changed)} json_synced={synced}")
    print("report", REPORT)
    for r in results:
        if r.get("saved"):
            print("CHANGED", r["stem"], [s["original_len"] for s in r["splits"]], "->", [s["parts"] for s in r["splits"]])
        else:
            print("SKIP", r["stem"])


if __name__ == "__main__":
    main()
