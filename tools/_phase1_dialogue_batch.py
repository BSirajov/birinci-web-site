# -*- coding: utf-8 -*-
"""Phase 1 batch: normalize ASCII/curly quotes to guillemets; write DOCX + report."""
from __future__ import annotations

import json
import re
from collections import OrderedDict
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
STORIES_DIR = ROOT / "az" / "stories"
DATA_JSON = ROOT / "az" / "data" / "stories.json"
ANALYSIS = ROOT / "docs" / "_az_lang_analysis.json"
BATCH_SIZE = 20
BATCH_NUM = 5
REPORT = ROOT / "docs" / f"PHASE1_BATCH{BATCH_NUM}_REPORT.md"


def tier2_stems() -> list[str]:
    stats = json.loads(ANALYSIS.read_text(encoding="utf-8"))
    long = {i["stem"] for i in stats["long_paras"]}
    stems: list[str] = []
    seen = set()
    for i in stats["mixed"]:
        stem = i["stem"]
        if stem in long or stem in seen:
            continue
        seen.add(stem)
        stems.append(stem)
    return stems


def normalize_paragraph(text: str) -> tuple[str, list[str]]:
    """Return (new_text, list of change tags)."""
    original = text
    changes: list[str] = []
    t = text

    # Collapse runs of 2+ spaces (not newlines)
    if re.search(r" {2,}", t):
        t = re.sub(r" {2,}", " ", t)
        changes.append("multi_space")

    # Odd ellipsis of 4+ dots → …
    if re.search(r"\.{4,}", t):
        t = re.sub(r"\.{4,}", "…", t)
        changes.append("ellipsis")

    # Curly double quotes → guillemets (pair-wise where possible)
    if "“" in t or "”" in t:
        # Replace opening/closing curly with guillemets
        t2 = t.replace("“", "«").replace("”", "»")
        if t2 != t:
            t = t2
            changes.append("curly_to_guillemets")

    # Straight ASCII doubles: convert paired "..." to «...»
    if '"' in t:
        parts = t.split('"')
        if len(parts) >= 3 and (len(parts) - 1) % 2 == 0:
            out = []
            for i, part in enumerate(parts):
                if i == 0:
                    out.append(part)
                elif i % 2 == 1:
                    out.append("«" + part + "»")
                else:
                    out.append(part)
            t2 = "".join(out)
            if t2 != t:
                t = t2
                changes.append("ascii_to_guillemets")
        else:
            # Repair common leftover: opening « with closing "
            t2 = t
            t2 = t2.replace('?"', "?»").replace('!"', "!»").replace('."', ".»").replace(',"', ",»")
            if t2.endswith('"') and "«" in t2:
                t2 = t2[:-1] + "»"
            if t2 != t:
                t = t2
                changes.append("ascii_closing_repaired")
            else:
                changes.append("ascii_unpaired_skipped")

    # Normalize en-dash used as leading dialogue marker to em-dash
    if re.match(r"^\s*–\s+", t):
        t = re.sub(r"^(\s*)–\s+", r"\1— ", t)
        changes.append("en_dash_to_em_dash")

    if t != original and not changes:
        changes.append("other")
    if t == original:
        return original, []
    return t, changes


def process_docx(stem: str) -> dict:
    path = STORIES_DIR / f"{stem}.docx"
    result = {
        "stem": stem,
        "path": str(path),
        "exists": path.is_file(),
        "changed_paras": 0,
        "changes": [],
        "samples": [],
    }
    if not path.is_file():
        return result

    doc = Document(str(path))
    changed = False
    for p in doc.paragraphs:
        raw = p.text or ""
        if not raw.strip():
            continue
        new, tags = normalize_paragraph(raw)
        if not tags or new == raw:
            # still record unpaired skip
            if "ascii_unpaired_skipped" in tags:
                result["changes"].append("ascii_unpaired_skipped")
            continue
        if len(result["samples"]) < 3:
            result["samples"].append({"before": raw[:180], "after": new[:180], "tags": tags})
        p.text = new
        changed = True
        result["changed_paras"] += 1
        result["changes"].extend(tags)

    if changed:
        doc.save(str(path))
    result["saved"] = changed
    result["change_counts"] = {}
    for tag in result["changes"]:
        result["change_counts"][tag] = result["change_counts"].get(tag, 0) + 1
    return result


def sync_stories_json(stems: set[str]) -> int:
    """Re-extract paragraphs for edited stems into stories.json without full rebuild wipe risk."""
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    updated = 0
    for cat in data["categories"]:
        for story in cat["stories"]:
            stem = story["stem"]
            if stem not in stems:
                continue
            path = STORIES_DIR / f"{stem}.docx"
            if not path.is_file():
                continue
            doc = Document(str(path))
            title = story["title"]
            paras = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    paras.append(t)
            while paras and paras[0].casefold() == title.casefold():
                paras = paras[1:]
            story["paragraphs"] = paras
            updated += 1
    DATA_JSON.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return updated


def write_report(batch: list[str], results: list[dict]) -> None:
    changed = [r for r in results if r.get("saved")]
    unchanged = [r for r in results if r.get("exists") and not r.get("saved")]
    missing = [r for r in results if not r.get("exists")]

    lines = []
    lines.append(f"# Phase 1 · Batch {BATCH_NUM} report")
    lines.append("")
    lines.append("House style: `docs/HOUSE_STYLE_DIALOGUE.md`")
    lines.append("")
    lines.append("## What this batch did")
    lines.append("")
    lines.append("- Converted ASCII/curly double quotes → «»")
    lines.append("- Normalized leading en-dash dialogue markers → em dash —")
    lines.append("- Collapsed accidental double spaces / overlong ellipsis where found")
    lines.append("- **Did not** remove correct inline «» or rewrite wording")
    lines.append("- **Did not** regenerate illustrations or audio")
    lines.append("")
    lines.append(
        "Note: many Tier-2 “mixed” stories already follow the house style "
        "(— for turns + «» for inline quotes). Those correctly show as unchanged."
    )
    lines.append("")
    lines.append(f"- Batch size: **{len(batch)}**")
    lines.append(f"- Files modified: **{len(changed)}**")
    lines.append(f"- Already compliant / no ASCII to fix: **{len(unchanged)}**")
    lines.append(f"- Missing DOCX: **{len(missing)}**")
    lines.append("")

    lines.append("## Modified stems")
    lines.append("")
    if not changed:
        lines.append("_None in this batch._")
    else:
        for r in changed:
            lines.append(
                f"- `{r['stem']}` — {r['changed_paras']} paragraph(s); tags: {r.get('change_counts', {})}"
            )
            for s in r.get("samples", []):
                lines.append(f"  - before: {s['before']}")
                lines.append(f"  - after:  {s['after']}")
    lines.append("")

    lines.append("## Unchanged stems (already house-style compliant for this pass)")
    lines.append("")
    for r in unchanged:
        lines.append(f"- `{r['stem']}`")
    lines.append("")

    if missing:
        lines.append("## Missing DOCX")
        lines.append("")
        for r in missing:
            lines.append(f"- `{r['stem']}`")
        lines.append("")

    lines.append("## Next")
    lines.append("")
    lines.append("1. Spot-check modified stories on the site after rebuild")
    lines.append(f"2. Reply **continue Phase 1 batch {BATCH_NUM + 1}** for the next 20 Tier-2 stems")
    lines.append("3. After Phase 1 completes → Phase 2 (Tier 1 long paragraphs)")
    lines.append("")

    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    stems = tier2_stems()
    start = (BATCH_NUM - 1) * BATCH_SIZE
    batch = stems[start : start + BATCH_SIZE]
    results = [process_docx(stem) for stem in batch]
    changed_stems = {r["stem"] for r in results if r.get("saved")}
    synced = sync_stories_json(changed_stems) if changed_stems else 0
    write_report(batch, results)
    print(f"batch={BATCH_NUM} size={len(batch)} modified={len(changed_stems)} json_synced={synced}")
    print("report", REPORT)
    for r in results:
        if r.get("saved"):
            print("CHANGED", r["stem"], r.get("change_counts"))


if __name__ == "__main__":
    main()
