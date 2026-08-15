# -*- coding: utf-8 -*-
"""Build static website for az moral stories (landing + one page per category)."""
from __future__ import annotations

import html
import json
import re
import shutil
import sys
from collections import defaultdict
from pathlib import Path
from urllib.parse import quote

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _az_story_categories import CATEGORY_ORDER, STORY_CATEGORY  # noqa: E402
from i18n_config import (  # noqa: E402
    SUPPORTED_LANGS,
    is_implemented,
    language_by_code,
    load_locale,
    locale_root,
    stories_ready,
    switcher_languages,
)

TOOLS = Path(__file__).resolve().parent
SITE_ROOT = TOOLS.parent
AZ_ROOT = SITE_ROOT / "az"
LANG = "az"
LANG_ROOT = AZ_ROOT
STORIES = AZ_ROOT / "stories"
ILLUSTRATIONS = AZ_ROOT / "illustrations"
AUDIO_DIR = AZ_ROOT / "audio"
MAP_JSON = Path(__file__).resolve().parent / "story-mapping.json"
DATA_JSON = AZ_ROOT / "data" / "stories.json"
PAGES_DIR = AZ_ROOT / "categories"
ASSETS = AZ_ROOT / "assets"
SHARED_ASSETS = SITE_ROOT / "assets"
LOCALE_ASSET_FILES = frozenset({"search-index.js", "stories-data.js", "site.js"})
LOCALE: dict = {}
UI: dict = {}

SITE_NAME = "Bir inci"
SITE_TITLE = "İbrətamiz deyimlər və hekayələr"
NAV_LABEL = "Ədəbiyyat"
NAV_STORIES_LABEL = "İbrətamiz hekayələr"
NAV_STORIES_DESC = "Kateqoriyalar üzrə ibrətamiz hekayələr"
NAV_ARTS_LABEL = "İncəsənət"
NAV_ARTS_DESC = "Rəssamlıq, memarlıq, musiqi və kino"
NAV_ARTS_ITEMS: list[dict[str, str]] = [
    {"label": "Rəssamlıq", "icon": "palette", "desc": "Rəng, forma və təsviri sənət"},
    {"label": "Memarlıq", "icon": "building", "desc": "Bina, məkan və dizayn"},
    {"label": "Musiqi", "icon": "music", "desc": "Səs, ritm və melodiya"},
    {"label": "Kino", "icon": "film", "desc": "Kino və ekran sənəti"},
]
NAV_SCIENCE_LABEL = "Elm"
NAV_NATURAL_SCIENCE_LABEL = "Təbiət elmləri"
NAV_NATURAL_SCIENCE_DESC = "Fizika, kimya və biologiya"
NAV_SOCIAL_SCIENCE_LABEL = "İctimai elmlər"
NAV_SOCIAL_SCIENCE_DESC = "Cəmiyyət, iqtisadiyyat və insan davranışı"
NAV_HUMANITIES_LABEL = "Humanitar elmlər"
NAV_HUMANITIES_DESC = "Tarix, fəlsəfə, dil və mədəniyyət"
NAV_MEDICAL_SCIENCE_LABEL = "Tibb elmləri"
NAV_MEDICAL_SCIENCE_DESC = "İnsan orqanizmi, diaqnostika və müalicə"
NAV_INFORMATICS_LABEL = "İnformatika"
NAV_INFORMATICS_DESC = "Alqoritmlər, data və proqram sistemləri"
NAV_MATH_LABEL = "Riyaziyyat"
NAV_MATH_DESC = "Rəqəm, düstur və məntiq"
# Curriculum order (not alphabetical)
NAV_NATURAL_SCIENCE_ITEMS: list[dict[str, str]] = [
    {
        "label": "Fizika",
        "icon": "atom",
        "desc": "Maddə, enerji, hərəkət və qüvvələrin qanunauyğunluqları",
    },
    {
        "label": "Kimya",
        "icon": "flask",
        "desc": "Maddələrin tərkibi, xassələri və çevrilmələri",
    },
    {
        "label": "Biologiya",
        "icon": "dna",
        "desc": "Canlı orqanizmlər və həyat prosesləri",
    },
    {
        "label": "Astronomiya",
        "icon": "telescope",
        "desc": "Kainat, ulduzlar, planetlər",
    },
    {
        "label": "Geologiya (Yer elmləri)",
        "icon": "mountain",
        "desc": "Yerin quruluşu, süxurlar, geoloji proseslər",
    },
    {
        "label": "Ekologiya",
        "icon": "leaf",
        "desc": "Canlıların ətraf mühitlə qarşılıqlı əlaqəsi",
    },
]
NAV_SOCIAL_SCIENCE_ITEMS: list[dict[str, str]] = [
    {
        "label": "İqtisadiyyat",
        "icon": "chart",
        "desc": "İstehsal, bölgü və maliyyə",
    },
    {
        "label": "Sosiologiya",
        "icon": "users",
        "desc": "Cəmiyyət, qruplar və sosial münasibətlər",
    },
    {
        "label": "Siyasi elmlər",
        "icon": "scale",
        "desc": "Hakimiyyət, dövlət və siyasi proseslər",
    },
    {
        "label": "Psixologiya",
        "icon": "brain",
        "desc": "Psixi proseslər və insan davranışı",
    },
    {
        "label": "Demoqrafiya",
        "icon": "globe",
        "desc": "Əhali, doğum, ölüm və miqrasiya",
    },
]
NAV_HUMANITIES_ITEMS: list[dict[str, str]] = [
    {
        "label": "Tarix",
        "icon": "landmark",
        "desc": "Keçmiş hadisələr və sivilizasiyalar",
    },
    {
        "label": "Fəlsəfə",
        "icon": "lightbulb",
        "desc": "Varlıq, bilik və dəyərlər haqqında düşüncə",
    },
    {
        "label": "Dilçilik (filologiya)",
        "icon": "message",
        "desc": "Dil quruluşu, inkişafı və işlənməsi",
    },
    {
        "label": "Ədəbiyyatşünaslıq",
        "icon": "book",
        "desc": "Ədəbi əsərlərin təhlili və tədqiqi",
    },
    {
        "label": "Sənətşünaslıq",
        "icon": "palette",
        "desc": "İncəsənət tarixi və nəzəriyyəsi",
    },
    {
        "label": "Dinşünaslıq (teologiya)",
        "icon": "sparkles",
        "desc": "Dinlər, inanclar və teoloji tədqiqat",
    },
    {
        "label": "Antropologiya və etnoqrafiya",
        "icon": "users",
        "desc": "İnsan mədəniyyətləri və etnik qruplar",
    },
    {
        "label": "Arxeologiya",
        "icon": "mountain",
        "desc": "Qədim mədəniyyətlərin maddi qalıqları",
    },
]
NAV_MEDICAL_SCIENCE_ITEMS: list[dict[str, str]] = [
    {
        "label": "Anatomiya",
        "icon": "activity",
        "desc": "Orqanizmin quruluşu və orqanlar",
    },
    {
        "label": "Fiziologiya",
        "icon": "heart",
        "desc": "Orqan və sistemlərin funksiyaları",
    },
    {
        "label": "Patologiya",
        "icon": "microscope",
        "desc": "Xəstəliklərin səbəb və mexanizmləri",
    },
    {
        "label": "Farmakologiya",
        "icon": "pill",
        "desc": "Dərmanlar və onların təsiri",
    },
    {
        "label": "Cərrahiyyə",
        "icon": "scissors",
        "desc": "Cərrahi müalicə və əməliyyatlar",
    },
    {
        "label": "Daxili xəstəliklər (terapiya)",
        "icon": "stethoscope",
        "desc": "Daxili orqan xəstəliklərinin müalicəsi",
    },
    {
        "label": "Pediatriya",
        "icon": "baby",
        "desc": "Uşaq sağlamlığı və xəstəlikləri",
    },
    {
        "label": "Kardiologiya",
        "icon": "heart",
        "desc": "Ürək və qan-damar sistemi",
    },
    {
        "label": "Nevrologiya",
        "icon": "brain",
        "desc": "Sinir sistemi və beyin",
    },
    {
        "label": "Onkologiya",
        "icon": "microscope",
        "desc": "Şiş xəstəlikləri və müalicəsi",
    },
    {
        "label": "Psixiatriya",
        "icon": "brain",
        "desc": "Psixi sağlamlıq və pozuntular",
    },
    {
        "label": "Genetika (tibbi)",
        "icon": "dna",
        "desc": "İrsi xəstəliklər və genlər",
    },
    {
        "label": "Epidemiologiya",
        "icon": "globe",
        "desc": "Xəstəliklərin yayılması və profilaktikası",
    },
]
# Azerbaijani alphabetical (A, D, K, P, S, Ş)
NAV_INFORMATICS_ITEMS: list[dict[str, str]] = [
    {
        "label": "Alqoritmlər və data strukturları",
        "icon": "layers",
        "desc": "Hesablama problemlərinin həlli üsulları və verilənlərin səmərəli təşkili",
    },
    {
        "label": "Data bazaları",
        "icon": "database",
        "desc": "Verilənlərin saxlanması, təşkili, axtarışı və idarə edilməsi",
    },
    {
        "label": "Data elmi",
        "icon": "chart",
        "desc": "Böyük məlumat massivlərinin toplanması, emalı, təhlili və onlardan bilik əldə edilməsi",
    },
    {
        "label": "Kibertəhlükəsizlik",
        "icon": "shield",
        "desc": "Məlumatların, şəbəkələrin və kompüter sistemlərinin qorunması",
    },
    {
        "label": "Kompüter qrafikası və vizuallaşdırma",
        "icon": "palette",
        "desc": "Vizual məlumatların yaradılması, emalı, təsviri və təqdim edilməsi",
    },
    {
        "label": "Proqram mühəndisliyi",
        "icon": "cpu",
        "desc": "Proqram təminatının tələblərinin müəyyənləşdirilməsi, layihələndirilməsi, hazırlanması, test edilməsi və idarə olunması",
    },
    {
        "label": "Süni intellekt",
        "icon": "brain",
        "desc": "Maşın öyrənməsi, neyron şəbəkələr, təbii dilin emalı və intellektual sistemlər",
    },
    {
        "label": "Şəbəkə texnologiyaları",
        "icon": "network",
        "desc": "İnternet, kompüter şəbəkələri, rabitə protokolları və paylanmış sistemlər",
    },
]
# Azerbaijani alphabetical (C, D, E, Ə, H, R, T)
NAV_MATH_ITEMS: list[dict[str, str]] = [
    {
        "label": "Cəbr",
        "icon": "calculator",
        "desc": "Tənliklər və cəbri strukturlar",
    },
    {
        "label": "Diferensial tənliklər və riyazi modelləşdirmə",
        "icon": "activity",
        "desc": "Real proseslərin riyazi təsviri",
    },
    {
        "label": "Diskret riyaziyyat və riyazi məntiq",
        "icon": "layers",
        "desc": "Kombinatorika, qraflar və formal düşüncə",
    },
    {
        "label": "Ehtimal nəzəriyyəsi və statistika",
        "icon": "chart",
        "desc": "Təsadüfilik və məlumatların təhlili",
    },
    {
        "label": "Ədədlər nəzəriyyəsi",
        "icon": "lightbulb",
        "desc": "Tam ədədlərin qanunauyğunluqları",
    },
    {
        "label": "Həndəsə və topologiya",
        "icon": "mountain",
        "desc": "Forma, fəza və onların xassələri",
    },
    {
        "label": "Riyazi analiz",
        "icon": "atom",
        "desc": "Funksiyalar və dəyişən proseslər",
    },
    {
        "label": "Tətbiqi və hesablama riyaziyyatı",
        "icon": "cpu",
        "desc": "Praktik məsələlərin riyazi və kompüter üsulları ilə həlli",
    },
]
NAV_SCIENCE_BRANCHES: dict[str, list[dict[str, str]]] = {
    "natural": NAV_NATURAL_SCIENCE_ITEMS,
    "social": NAV_SOCIAL_SCIENCE_ITEMS,
    "humanities": NAV_HUMANITIES_ITEMS,
    "medical": NAV_MEDICAL_SCIENCE_ITEMS,
    "informatics": NAV_INFORMATICS_ITEMS,
    "math": NAV_MATH_ITEMS,
}
_AZ_SCIENCE_BRANCHES: dict[str, list[dict[str, str]]] = {
    key: [dict(item) for item in items]
    for key, items in NAV_SCIENCE_BRANCHES.items()
}
_SCIENCE_BRANCH_LOCALE_KEYS = {
    "natural": "natural_items",
    "social": "social_items",
    "humanities": "humanities_items",
    "medical": "medical_items",
    "informatics": "informatics_items",
    "math": "math_items",
}
# Azerbaijani alphabetical: H, İ (c before n), R, T (ə before i)
NAV_SCIENCE_ITEMS: list[dict[str, str]] = [
    {
        "label": NAV_HUMANITIES_LABEL,
        "icon": "scroll",
        "desc": NAV_HUMANITIES_DESC,
        "branch": "humanities",
    },
    {
        "label": NAV_SOCIAL_SCIENCE_LABEL,
        "icon": "users",
        "desc": NAV_SOCIAL_SCIENCE_DESC,
        "branch": "social",
    },
    {
        "label": NAV_INFORMATICS_LABEL,
        "icon": "cpu",
        "desc": NAV_INFORMATICS_DESC,
        "branch": "informatics",
    },
    {
        "label": NAV_MATH_LABEL,
        "icon": "calculator",
        "desc": NAV_MATH_DESC,
        "branch": "math",
    },
    {
        "label": NAV_NATURAL_SCIENCE_LABEL,
        "icon": "flask",
        "desc": NAV_NATURAL_SCIENCE_DESC,
        "branch": "natural",
    },
    {
        "label": NAV_MEDICAL_SCIENCE_LABEL,
        "icon": "stethoscope",
        "desc": NAV_MEDICAL_SCIENCE_DESC,
        "branch": "medical",
    },
]
TOP_NAV_LINKS: list[dict[str, str]] = [
    {"label": "Tanınmış şəxsiyyətlər", "icon": "landmark"},
    {"label": "Kəşf və ixtiralar", "icon": "lightbulb"},
    {"label": "Ümumi biliklər", "icon": "globe"},
    {"label": "Haqqımızda", "icon": "info"},
    {"label": "Bizi dəstəkləyin", "icon": "hand-heart"},
]
HOME_CRUMB = "Ana səhifə"
ASSET_VERSION = "20260815j"


def shared_asset_href(prefix: str, filename: str) -> str:
    """Href from a locale page ({lang}/ or {lang}/categories/) to /assets/."""
    return f"../{prefix}assets/{quote(filename)}"


def sync_shared_assets() -> None:
    """Collect identical chrome files once at the project-root assets folder."""
    SHARED_ASSETS.mkdir(parents=True, exist_ok=True)
    az_assets = SITE_ROOT / "az" / "assets"
    if az_assets.is_dir():
        for path in az_assets.iterdir():
            if not path.is_file() or path.name in LOCALE_ASSET_FILES:
                continue
            dest = SHARED_ASSETS / path.name
            if not dest.exists():
                shutil.copy2(path, dest)
    (SHARED_ASSETS / "site.css").write_text(CSS, encoding="utf-8")


def prune_locale_assets() -> None:
    """Keep only per-locale generated JS under {lang}/assets/."""
    if not ASSETS.is_dir():
        return
    for path in list(ASSETS.iterdir()):
        if path.is_file() and path.name not in LOCALE_ASSET_FILES:
            path.unlink()


def apply_locale(lang: str) -> None:
    """Point path globals and chrome strings at a locale."""
    global LANG, LANG_ROOT, STORIES, ILLUSTRATIONS, AUDIO_DIR, DATA_JSON, PAGES_DIR, ASSETS
    global LOCALE, UI, SITE_NAME, SITE_TITLE, NAV_LABEL, NAV_STORIES_LABEL, NAV_STORIES_DESC
    global NAV_ARTS_LABEL, NAV_ARTS_DESC, NAV_SCIENCE_LABEL
    global NAV_NATURAL_SCIENCE_LABEL, NAV_NATURAL_SCIENCE_DESC
    global NAV_SOCIAL_SCIENCE_LABEL, NAV_SOCIAL_SCIENCE_DESC
    global NAV_HUMANITIES_LABEL, NAV_HUMANITIES_DESC
    global NAV_MEDICAL_SCIENCE_LABEL, NAV_MEDICAL_SCIENCE_DESC
    global NAV_INFORMATICS_LABEL, NAV_INFORMATICS_DESC
    global NAV_MATH_LABEL, NAV_MATH_DESC
    global NAV_ARTS_ITEMS, TOP_NAV_LINKS, HOME_CRUMB, CATEGORY_META
    global NAV_SCIENCE_ITEMS, NAV_SCIENCE_BRANCHES

    loc = load_locale(lang)
    LANG = lang
    LANG_ROOT = locale_root(lang)
    STORIES = LANG_ROOT / "stories"
    ILLUSTRATIONS = LANG_ROOT / "illustrations"
    AUDIO_DIR = LANG_ROOT / "audio"
    DATA_JSON = LANG_ROOT / "data" / "stories.json"
    PAGES_DIR = LANG_ROOT / "categories"
    ASSETS = LANG_ROOT / "assets"
    LOCALE = loc
    UI = dict(loc.get("ui") or {})

    SITE_NAME = loc.get("site_name") or SITE_NAME
    SITE_TITLE = loc.get("site_title") or SITE_TITLE
    NAV_LABEL = loc.get("nav_label") or NAV_LABEL
    NAV_STORIES_LABEL = loc.get("nav_stories_label") or NAV_STORIES_LABEL
    NAV_STORIES_DESC = loc.get("nav_stories_desc") or NAV_STORIES_DESC
    NAV_ARTS_LABEL = loc.get("nav_arts_label") or NAV_ARTS_LABEL
    NAV_ARTS_DESC = loc.get("nav_arts_desc") or NAV_ARTS_DESC
    NAV_SCIENCE_LABEL = loc.get("nav_science_label") or NAV_SCIENCE_LABEL
    NAV_NATURAL_SCIENCE_LABEL = loc.get("nav_natural_science_label") or NAV_NATURAL_SCIENCE_LABEL
    NAV_NATURAL_SCIENCE_DESC = loc.get("nav_natural_science_desc") or NAV_NATURAL_SCIENCE_DESC
    NAV_SOCIAL_SCIENCE_LABEL = loc.get("nav_social_science_label") or NAV_SOCIAL_SCIENCE_LABEL
    NAV_SOCIAL_SCIENCE_DESC = loc.get("nav_social_science_desc") or NAV_SOCIAL_SCIENCE_DESC
    NAV_HUMANITIES_LABEL = loc.get("nav_humanities_label") or NAV_HUMANITIES_LABEL
    NAV_HUMANITIES_DESC = loc.get("nav_humanities_desc") or NAV_HUMANITIES_DESC
    NAV_MEDICAL_SCIENCE_LABEL = loc.get("nav_medical_science_label") or NAV_MEDICAL_SCIENCE_LABEL
    NAV_MEDICAL_SCIENCE_DESC = loc.get("nav_medical_science_desc") or NAV_MEDICAL_SCIENCE_DESC
    NAV_INFORMATICS_LABEL = loc.get("nav_informatics_label") or NAV_INFORMATICS_LABEL
    NAV_INFORMATICS_DESC = loc.get("nav_informatics_desc") or NAV_INFORMATICS_DESC
    NAV_MATH_LABEL = loc.get("nav_math_label") or NAV_MATH_LABEL
    NAV_MATH_DESC = loc.get("nav_math_desc") or NAV_MATH_DESC
    HOME_CRUMB = loc.get("home_crumb") or HOME_CRUMB

    cats = loc.get("categories") or {}
    for meta in CATEGORY_META:
        slug = meta["slug"]
        if slug in cats:
            meta["title"] = cats[slug].get("title") or meta["title"]
            meta["blurb"] = cats[slug].get("blurb") or meta["blurb"]

    if loc.get("arts_items"):
        NAV_ARTS_ITEMS = list(loc["arts_items"])
    if loc.get("top_nav"):
        TOP_NAV_LINKS = list(loc["top_nav"])

    # Rebuild science top-level labels from locale (branch lists keep AZ structure/icons).
    NAV_SCIENCE_ITEMS = [
        {
            "label": NAV_HUMANITIES_LABEL,
            "icon": "scroll",
            "desc": NAV_HUMANITIES_DESC,
            "branch": "humanities",
        },
        {
            "label": NAV_SOCIAL_SCIENCE_LABEL,
            "icon": "users",
            "desc": NAV_SOCIAL_SCIENCE_DESC,
            "branch": "social",
        },
        {
            "label": NAV_INFORMATICS_LABEL,
            "icon": "cpu",
            "desc": NAV_INFORMATICS_DESC,
            "branch": "informatics",
        },
        {
            "label": NAV_MATH_LABEL,
            "icon": "calculator",
            "desc": NAV_MATH_DESC,
            "branch": "math",
        },
        {
            "label": NAV_NATURAL_SCIENCE_LABEL,
            "icon": "flask",
            "desc": NAV_NATURAL_SCIENCE_DESC,
            "branch": "natural",
        },
        {
            "label": NAV_MEDICAL_SCIENCE_LABEL,
            "icon": "stethoscope",
            "desc": NAV_MEDICAL_SCIENCE_DESC,
            "branch": "medical",
        },
    ]

    branches: dict[str, list[dict[str, str]]] = {}
    for branch, key in _SCIENCE_BRANCH_LOCALE_KEYS.items():
        items = loc.get(key)
        if isinstance(items, list) and items:
            branches[branch] = list(items)
        else:
            branches[branch] = [dict(item) for item in _AZ_SCIENCE_BRANCHES[branch]]
    NAV_SCIENCE_BRANCHES = branches


def t_ui(key: str, default: str = "") -> str:
    return str(UI.get(key) or default or key)


# Immediate home view toggle (not deferred). Survives site.js load races / failures.
HOME_VIEW_BOOTSTRAP = r"""
<script>
(function () {
  function hideEl(el, hide) {
    if (!el) return;
    el.hidden = !!hide;
    if (hide) el.setAttribute("hidden", "");
    else el.removeAttribute("hidden");
  }

  function esc(s) {
    return String(s || "")
      .replace(/&/g, "&amp;")
      .replace(/</g, "&lt;")
      .replace(/>/g, "&gt;")
      .replace(/"/g, "&quot;");
  }

  function scrollHomeToolsIntoView() {
    var run = function () {
      window.scrollTo(0, 0);
      document.documentElement.scrollTop = 0;
      if (document.body) document.body.scrollTop = 0;
    };
    // Cards hide shrinks the document; wait for layout, then again after stories paint.
    requestAnimationFrame(function () {
      requestAnimationFrame(function () {
        run();
        window.setTimeout(run, 80);
        window.setTimeout(run, 250);
      });
    });
  }

  function applyView(view) {
    view = view === "list" ? "list" : "cards";
    var prev = window.__birinciHomeView;
    var cards = document.querySelector('[data-view="cards"]');
    var list = document.querySelector('[data-view="list"]');
    hideEl(cards, view !== "cards");
    hideEl(list, view !== "list");
    document.querySelectorAll("[data-home-view]").forEach(function (btn) {
      btn.setAttribute(
        "aria-pressed",
        btn.getAttribute("data-home-view") === view ? "true" : "false"
      );
    });
    document.querySelectorAll("[data-home-list-only]").forEach(function (el) {
      hideEl(el, view !== "list");
    });
    window.__birinciHomeView = view;
    try {
      localStorage.setItem("birinci-home-view", view);
    } catch (_) {}
    if (view === "list") loadStoriesFallback();
    if (view === "list" && prev !== "list") scrollHomeToolsIntoView();
    return false;
  }

  window.__birinciScrollHomeTools = scrollHomeToolsIntoView;

  function renderFallbackStories(catalog) {
    var host = document.querySelector('[data-view="list"] [data-stories-list]');
    if (!host) return;
    if (host.querySelector(".story")) return;
    var rows = [];
    (catalog.categories || []).forEach(function (cat) {
      (cat.stories || []).forEach(function (story) {
        rows.push(story);
      });
    });
    rows.sort(function (a, b) {
      return String(a.title || "").localeCompare(String(b.title || ""), document.documentElement.lang || "az", {
        sensitivity: "base",
      });
    });
    var page = rows;
    var sizeEl = document.querySelector('[data-tools="home"] [data-home-batch-size]');
    var raw = "";
    var allMode = false;
    var pageSize = 12;
    try {
      allMode = localStorage.getItem("birinci-home-batch-all") === "1";
      raw = localStorage.getItem("birinci-home-batch-size") || "";
      if (!raw) {
        var legacy = localStorage.getItem("birinci-home-page-size") || "";
        if (legacy && legacy !== "all") raw = legacy;
        else if (legacy === "all") allMode = true;
      }
    } catch (_) {}
    var cap = Math.max(1, rows.length || 1);
    var n = Number(raw);
    if (Number.isFinite(n) && n > 0) pageSize = Math.min(Math.floor(n), cap);
    if (sizeEl) {
      sizeEl.max = String(cap);
      sizeEl.value = String(pageSize);
    }
    page = allMode ? rows : rows.slice(0, pageSize);
    host.innerHTML = page
      .map(function (story) {
        var list = story.paragraphs || [];
        var last = list.length - 1;
        var foldAzI = function (s) { return String(s || "").replace(/[İIı]/g, "i"); };
        var srcRe = /(internet\s+sources|internet\s+mənb|internet\s+kaynak|открыт\w*\s+источник|интернет|(?:source|mənbə|kaynak|источник|булак|булагы)\s*:)/i;
        var moralRe = /^(ibrət|ibret|moral|мораль|үлгү)\s*:/i;
        var authorSrcStems = { "everyone-has-work-to-do": 1, "weeds-must-be-pulled-from-the-root": 1, "the-silent-corridor": 1 };
        var authorSrc = !!(story.stem && authorSrcStems[story.stem]);
        var lastIsSrc = last >= 0 && (authorSrc || srcRe.test(foldAzI(list[last] || "")));
        var srcLabel = (window.__BIRINCI_I18N__ && window.__BIRINCI_I18N__.ui && window.__BIRINCI_I18N__.ui.story_source) || "";
        var moralI = -1;
        for (var j = lastIsSrc ? last - 1 : last; j >= 0; j--) {
          if (moralRe.test(foldAzI(String(list[j] || "").trim()))) { moralI = j; break; }
        }
        if (moralI < 0) moralI = lastIsSrc && last >= 1 ? last - 1 : last;
        var paras = list
          .map(function (p, i) {
            var isSrc = lastIsSrc && i === last;
            var cls = isSrc ? "story__source" : i === moralI ? "story__moral" : "";
            var text = isSrc && srcLabel && !authorSrc ? srcLabel : p;
            return "<p" + (cls ? ' class="' + cls + '"' : "") + ">" + esc(text) + "</p>";
          })
          .join("");
        return (
          '<article class="story news-card" id="' +
          esc(story.stem) +
          '" data-stem="' +
          esc(story.stem) +
          '" data-title="' +
          esc(story.title) +
          '"><div class="card-header"><h2 class="card-title story__title">' +
          esc(story.title) +
          '</h2></div><div class="card-body"><div class="story__content"><div class="story__panel"><div class="story__text card-text">' +
          paras +
          "</div></div></div></div></article>"
        );
      })
      .join("");
  }

  function loadStoriesFallback() {
    var panel = document.querySelector('[data-view="list"]');
    var host = panel && panel.querySelector("[data-stories-list]");
    if (!panel || !host) return;
    if (host.querySelector(".story")) return;
    if (window.__birinciHomeStoriesLoading) return;
    if (window.__BIRINCI_STORIES__) {
      renderFallbackStories(window.__BIRINCI_STORIES__);
      return;
    }
    window.__birinciHomeStoriesLoading = true;
    var scriptUrl = panel.getAttribute("data-stories-script");
    var jsonUrl = panel.getAttribute("data-stories-url") || "data/stories.json";

    function done() {
      window.__birinciHomeStoriesLoading = false;
    }

    function fail() {
      done();
    }

    function loadViaScript() {
      if (!scriptUrl) return Promise.reject(new Error("no-script"));
      return new Promise(function (resolve, reject) {
        var s = document.createElement("script");
        s.src = scriptUrl;
        s.async = true;
        s.onload = function () {
          if (window.__BIRINCI_STORIES__) resolve(window.__BIRINCI_STORIES__);
          else reject(new Error("empty-stories"));
        };
        s.onerror = function () {
          reject(new Error("script-error"));
        };
        document.head.appendChild(s);
      });
    }

    loadViaScript()
      .catch(function () {
        return fetch(jsonUrl).then(function (res) {
          if (!res.ok) throw new Error("fetch-failed");
          return res.json();
        });
      })
      .then(function (catalog) {
        window.__BIRINCI_STORIES__ = catalog;
        renderFallbackStories(catalog);
        done();
      })
      .catch(fail);
  }

  window.__birinciSetHomeView = applyView;

  document.addEventListener(
    "click",
    function (ev) {
      if (!document.body || !document.body.classList.contains("page-home")) return;
      var t = ev.target;
      if (!t || !t.closest) return;
      var btn = t.closest("[data-home-view]");
      if (!btn) return;
      var next = btn.getAttribute("data-home-view");
      if (next !== "list" && next !== "cards") return;
      applyView(next);
    },
    true
  );
})();
</script>
""".strip()

# Inline Lucide-style stroke icons (24x24 viewBox) for menu items.
CATEGORY_ICONS: dict[str, str] = {
    # Top-nav stories menu
    "layers": '<path d="M12.83 2.18a2 2 0 0 0-1.66 0L2.6 6.08a1 1 0 0 0 0 1.83l8.58 3.91a2 2 0 0 0 1.66 0l8.58-3.9a1 1 0 0 0 0-1.83z"/><path d="M2 12a1 1 0 0 0 .58.91l8.6 3.91a2 2 0 0 0 1.65 0l8.58-3.9A1 1 0 0 0 22 12"/><path d="M2 17a1 1 0 0 0 .58.91l8.6 3.91a2 2 0 0 0 1.65 0l8.58-3.9A1 1 0 0 0 22 17"/>',
    # Spirituality / faith
    "sparkles": '<path d="M9.937 15.5A2 2 0 0 0 8.5 14.063l-6.135-1.582a.5.5 0 0 1 0-.962L8.5 9.936A2 2 0 0 0 9.937 8.5l1.582-6.135a.5.5 0 0 1 .963 0L14.063 8.5A2 2 0 0 0 15.5 9.937l6.135 1.581a.5.5 0 0 1 0 .964L15.5 14.063a2 2 0 0 0-1.437 1.437l-1.582 6.135a.5.5 0 0 1-.963 0z"/><path d="M20 3v4"/><path d="M22 5h-4"/><path d="M4 17v2"/><path d="M5 18H3"/>',
    # Family / home
    "home": '<path d="M15 21v-8a1 1 0 0 0-1-1h-4a1 1 0 0 0-1 1v8"/><path d="M3 10a2 2 0 0 1 .709-1.528l7-5.999a2 2 0 0 1 2.582 0l7 5.999A2 2 0 0 1 21 10v9a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2z"/>',
    # Love
    "heart": '<path d="M19 14c1.49-1.46 3-3.21 3-5.5A5.5 5.5 0 0 0 16.5 3c-1.76 0-3 .5-4.5 2-1.5-1.5-2.74-2-4.5-2A5.5 5.5 0 0 0 2 8.5c0 2.3 1.5 4.05 3 5.5l7 7Z"/>',
    # Friendship
    "users": '<path d="M16 21v-2a4 4 0 0 0-4-4H6a4 4 0 0 0-4 4v2"/><circle cx="9" cy="7" r="4"/><path d="M22 21v-2a4 4 0 0 0-3-3.87"/><path d="M16 3.13a4 4 0 0 1 0 7.75"/>',
    # Character / ethics
    "shield": '<path d="M20 13c0 5-3.5 7.5-7.66 8.95a1 1 0 0 1-.67-.01C7.5 20.5 4 18 4 13V6a1 1 0 0 1 1-1c2 0 4.5-1.2 6.24-2.72a1.17 1.17 0 0 1 1.52 0C14.51 3.81 17 5 19 5a1 1 0 0 1 1 1z"/>',
    # Communication
    "message": '<path d="M7.9 20A9 9 0 1 0 4 16.1L2 22Z"/>',
    # Wisdom
    "lightbulb": '<path d="M15 14c.2-1 .7-1.7 1.5-2.5 1-.9 1.5-2.2 1.5-3.5A6 6 0 0 0 6 8c0 1 .2 2.2 1.5 3.5.7.7 1.3 1.5 1.5 2.5"/><path d="M9 18h6"/><path d="M10 22h4"/>',
    # Work / wealth
    "briefcase": '<path d="M16 20V4a2 2 0 0 0-2-2h-4a2 2 0 0 0-2 2v16"/><rect width="20" height="14" x="2" y="6" rx="2"/>',
    # Justice
    "scale": '<path d="m16 16 3-8 3 8c-.87.65-1.92 1-3 1s-2.13-.35-3-1Z"/><path d="m2 16 3-8 3 8c-.87.65-1.92 1-3 1s-2.13-.35-3-1Z"/><path d="M7 21h10"/><path d="M12 3v18"/><path d="M3 7h2c2 0 5-1 7-2 2 1 5 2 7 2h2"/>',
    # Time / aging
    "clock": '<circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/>',
    # Parables / literature
    "book": '<path d="M12 7v14"/><path d="M3 18a1 1 0 0 1-1-1V4a1 1 0 0 1 1-1h5a4 4 0 0 1 4 4 4 4 0 0 1 4-4h5a1 1 0 0 1 1 1v13a1 1 0 0 1-1 1h-6a3 3 0 0 0-3 3 3 3 0 0 0-3-3z"/>',
    # History / figures
    "landmark": '<path d="M10 18v-7"/><path d="M11.12 2.198a2 2 0 0 1 1.76.006l7.866 3.847c.10.0.0.5-.22.949H3.474c-.53 0-.695-.716-.22-.949z"/><path d="M14 18v-7"/><path d="M18 18v-7"/><path d="M3 22h18"/><path d="M6 18v-7"/>',
    # Math
    "calculator": '<rect width="16" height="20" x="4" y="2" rx="2"/><line x1="8" x2="16" y1="6" y2="6"/><line x1="16" x2="16" y1="14" y2="18"/><path d="M16 10h.01"/><path d="M12 10h.01"/><path d="M8 10h.01"/><path d="M12 14h.01"/><path d="M8 14h.01"/><path d="M12 18h.01"/><path d="M8 18h.01"/>',
    # Nature / science
    "flask": '<path d="M10 2v7.527a2 2 0 0 1-.211.896L4.72 20.55a1 1 0 0 0 .9 1.45h12.76a1 1 0 0 0 .9-1.45l-5.069-10.127A2 2 0 0 1 14 9.527V2"/><path d="M8.5 2h7"/><path d="M7 16h10"/>',
    # Science umbrella
    "atom": '<circle cx="12" cy="12" r="1"/><path d="M20.2 20.2c2.04-2.03.02-7.36-4.5-11.9-4.54-4.52-9.87-6.54-11.9-4.5-2.04 2.03-.02 7.36 4.5 11.9 4.54 4.52 9.87 6.54 11.9 4.5Z"/><path d="M15.7 15.7c4.52-4.54 6.54-9.87 4.5-11.9-2.03-2.04-7.36-.02-11.9 4.5-4.52 4.54-6.54 9.87-4.5 11.9 2.03 2.04 7.36.02 11.9-4.5Z"/>',
    # Medicine
    "stethoscope": '<path d="M11 2v2"/><path d="M5 2v2"/><path d="M5 3H4a2 2 0 0 0-2 2v4a6 6 0 0 0 12 0V5a2 2 0 0 0-2-2h-1"/><path d="M8 15a6 6 0 0 0 12 0v-3"/><circle cx="20" cy="10" r="2"/>',
    # IT
    "cpu": '<rect width="16" height="16" x="4" y="4" rx="2"/><rect width="6" height="6" x="9" y="9" rx="1"/><path d="M15 2v2"/><path d="M15 20v2"/><path d="M2 15h2"/><path d="M2 9h2"/><path d="M20 15h2"/><path d="M20 9h2"/><path d="M9 2v2"/><path d="M9 20v2"/>',
    "database": '<ellipse cx="12" cy="5" rx="9" ry="3"/><path d="M3 5V19A9 3 0 0 0 21 19V5"/><path d="M3 12A9 3 0 0 0 21 12"/>',
    "network": '<rect x="16" y="16" width="6" height="6" rx="1"/><rect x="2" y="16" width="6" height="6" rx="1"/><rect x="9" y="2" width="6" height="6" rx="1"/><path d="M5 16v-3a1 1 0 0 1 1-1h12a1 1 0 0 1 1 1v3"/><path d="M12 12V8"/>',
    # Economics
    "chart": '<path d="M3 3v18h18"/><path d="m19 9-5 5-4-4-3 3"/>',
    # General knowledge
    "globe": '<circle cx="12" cy="12" r="10"/><path d="M12 2a14.5 14.5 0 0 0 0 20 14.5 14.5 0 0 0 0-20"/><path d="M2 12h20"/>',
    "website": '<circle cx="12" cy="12" r="10"/><path d="M12 2a14.5 14.5 0 0 0 0 20 14.5 14.5 0 0 0 0-20"/><path d="M2 12h20"/>',
    "phone": '<path d="M22 16.92v3a2 2 0 0 1-2.18 2 19.79 19.79 0 0 1-8.63-3.07 19.5 19.5 0 0 1-6-6 19.79 19.79 0 0 1-3.07-8.67A2 2 0 0 1 4.11 2h3a2 2 0 0 1 2 1.72 12.84 12.84 0 0 0 .7 2.81 2 2 0 0 1-.45 2.11L8.09 9.91a16 16 0 0 0 6 6l1.27-1.27a2 2 0 0 1 2.11-.45 12.84 12.84 0 0 0 2.81.7A2 2 0 0 1 22 16.92z"/>',
    "map-pin": '<path d="M20 10c0 6-8 12-8 12s-8-6-8-12a8 8 0 0 1 16 0Z"/><circle cx="12" cy="10" r="3"/>',
    # About
    "info": '<circle cx="12" cy="12" r="10"/><path d="M12 16v-4"/><path d="M12 8h.01"/>',
    # Support
    "hand-heart": '<path d="M11 14h2a2 2 0 0 0 0-4h-3c-.6 0-1.1.2-1.4.6L3 16"/><path d="m7 20 1.6-1.4c.3-.4.8-.6 1.4-.6h4c1.1 0 2.1-.4 2.8-1.2l4.6-4.4a2 2 0 0 0-2.75-2.91l-4.2 3.9"/><path d="m2 15 6 6"/><path d="M19.5 8.5c.7-.7 1.5-1.6 1.5-2.7A2.73 2.73 0 0 0 16 4a2.78 2.78 0 0 0-5 1.8c0 1.2.8 2 1.5 2.7L16 12Z"/>',
    # Arts
    "palette": '<circle cx="13.5" cy="6.5" r=".5" fill="currentColor"/><circle cx="17.5" cy="10.5" r=".5" fill="currentColor"/><circle cx="8.5" cy="7.5" r=".5" fill="currentColor"/><circle cx="6.5" cy="12.5" r=".5" fill="currentColor"/><path d="M12 2C6.5 2 2 6.5 2 12s4.5 10 10 10c.926 0 1.648-.746 1.648-1.688 0-.437-.18-.835-.437-1.125-.29-.289-.438-.652-.438-1.125a1.64 1.64 0 0 1 1.668-1.668h1.996c3.051 0 5.555-2.503 5.555-5.554C21.965 6.012 17.461 2 12 2z"/>',
    "building": '<rect width="16" height="20" x="4" y="2" rx="2" ry="2"/><path d="M9 22v-4h6v4"/><path d="M8 6h.01"/><path d="M16 6h.01"/><path d="M12 6h.01"/><path d="M12 10h.01"/><path d="M12 14h.01"/><path d="M16 10h.01"/><path d="M16 14h.01"/><path d="M8 10h.01"/><path d="M8 14h.01"/>',
    "music": '<path d="M9 18V5l12-2v13"/><circle cx="6" cy="18" r="3"/><circle cx="18" cy="16" r="3"/>',
    "film": '<rect width="18" height="18" x="3" y="3" rx="2"/><path d="M7 3v18"/><path d="M3 7.5h4"/><path d="M3 12h18"/><path d="M3 16.5h4"/><path d="M17 3v18"/><path d="M17 7.5h4"/><path d="M17 16.5h4"/>',
    # Natural sciences
    "dna": '<path d="m10 16 1.5 1.5"/><path d="m14 8-1.5-1.5"/><path d="M15 2c-1.798 1.998-2.518 3.995-2.807 5.993"/><path d="m16.5 10.5 1 1"/><path d="m17 6-2.899-2.899"/><path d="M2 15c6.667-6 13.333 0 20-6"/><path d="m2 9 2.899-2.899"/><path d="M20 9c-1.798 1.998-2.518 3.995-2.807 5.993"/><path d="m8.5 13.5-1-1"/><path d="m9 18 2.899 2.899"/>',
    "telescope": '<path d="m10.065 12.493-6.18 1.318a.934.934 0 0 1-1.108-.702l-.537-2.15a1.07 1.07 0 0 1 .691-1.265l13.504-4.44"/><path d="m13.56 11.747 4.332-.924"/><path d="m16 21-3.105-6.21"/><path d="M16.485 5.94a2 2 0 0 1 1.455-2.425l1.09-.272a1 1 0 0 1 1.212.727l1.515 6.06a1 1 0 0 1-.727 1.213l-1.09.272a2 2 0 0 1-2.425-1.455z"/><path d="m6.158 8.633 1.114 4.456"/><path d="m8 21 3.105-6.21"/><circle cx="12" cy="13" r="2"/>',
    "mountain": '<path d="m8 3 4 8 5-5 5 15H2L8 3z"/>',
    "leaf": '<path d="M11 20A7 7 0 0 1 9.8 6.1C15.5 5 17 4.48 19 2c1 2 2 4.18 2 8 0 5.5-4.78 10-10 10Z"/><path d="M2 21c0-3 1.85-5.36 5.08-6C9.5 14.52 12 13 13 12"/>',
    # Humanities / social / medical extras
    "scroll": '<path d="M19 17V5a2 2 0 0 0-2-2H4"/><path d="M8 21h12a2 2 0 0 0 2-2v-1a1 1 0 0 0-1-1H11a1 1 0 0 0-1 1v1a2 2 0 1 1-4 0V5a2 2 0 1 0-4 0v2a1 1 0 0 0 1 1h3"/>',
    "brain": '<path d="M12 18V5"/><path d="M15 13a4.17 4.17 0 0 1-3-4 4.17 4.17 0 0 1-3 4"/><path d="M17.598 6.5A3 3 0 1 0 12 5a3 3 0 1 0-5.598 1.5"/><path d="M17.997 5.125a4 4 0 0 1 2.526 5.77"/><path d="M18 18a4 4 0 0 0 2-7.464"/><path d="M19.967 17.483A4 4 0 1 1 12 18a4 4 0 1 1-7.967-.517"/><path d="M6 18a4 4 0 0 1-2-7.464"/><path d="M6.003 5.125a4 4 0 0 0-2.526 5.77"/>',
    "activity": '<path d="M22 12h-2.48a2 2 0 0 0-1.93 1.46l-2.35 8.36a.25.25 0 0 1-.48 0L9.24 2.18a.25.25 0 0 0-.48 0l-2.35 8.36A2 2 0 0 1 4.49 12H2"/>',
    "microscope": '<path d="M6 18h8"/><path d="M3 22h18"/><path d="M14 22a7 7 0 1 0 0-14h-1"/><path d="M9 14h2"/><path d="M9 12a2 2 0 0 1-2-2V6h6v4a2 2 0 0 1-2 2Z"/><path d="M12 6V3a1 1 0 0 0-1-1H9a1 1 0 0 0-1 1v3"/>',
    "pill": '<path d="m10.5 20.5 10-10a4.95 4.95 0 1 0-7-7l-10 10a4.95 4.95 0 1 0 7 7Z"/><path d="m8.5 8.5 7 7"/>',
    "scissors": '<circle cx="6" cy="6" r="3"/><path d="M8.12 8.12 12 12"/><path d="M20 4 8.12 15.88"/><circle cx="6" cy="18" r="3"/><path d="M14.8 14.8 20 20"/>',
    "baby": '<path d="M10 16c3.31 0 6-2.69 6-6"/><path d="M2 16h20"/><path d="M12 2v4"/><circle cx="12" cy="10" r="4"/><path d="m8 22 2-4"/><path d="m16 22-2-4"/>',
}

# Per-icon accent colors for 3D colorful badges.
ICON_COLORS: dict[str, dict[str, str]] = {
    "layers": {"from": "#38bdf8", "to": "#0369a1", "glow": "#7dd3fc"},
    "sparkles": {"from": "#7c5cff", "to": "#b44dff", "glow": "#a78bfa"},
    "home": {"from": "#0ea5e9", "to": "#2563eb", "glow": "#38bdf8"},
    "heart": {"from": "#f43f5e", "to": "#e11d48", "glow": "#fb7185"},
    "users": {"from": "#14b8a6", "to": "#0d9488", "glow": "#2dd4bf"},
    "shield": {"from": "#f59e0b", "to": "#d97706", "glow": "#fbbf24"},
    "message": {"from": "#3b82f6", "to": "#1d4ed8", "glow": "#60a5fa"},
    "lightbulb": {"from": "#fbbf24", "to": "#f97316", "glow": "#fcd34d"},
    "briefcase": {"from": "#64748b", "to": "#334155", "glow": "#94a3b8"},
    "scale": {"from": "#6366f1", "to": "#4f46e5", "glow": "#818cf8"},
    "clock": {"from": "#a855f7", "to": "#7e22ce", "glow": "#c084fc"},
    "book": {"from": "#10b981", "to": "#059669", "glow": "#34d399"},
    "landmark": {"from": "#eab308", "to": "#ca8a04", "glow": "#facc15"},
    "calculator": {"from": "#0ea5e9", "to": "#0369a1", "glow": "#7dd3fc"},
    "flask": {"from": "#22c55e", "to": "#15803d", "glow": "#86efac"},
    "atom": {"from": "#06b6d4", "to": "#0e7490", "glow": "#67e8f9"},
    "stethoscope": {"from": "#ef4444", "to": "#b91c1c", "glow": "#fca5a5"},
    "cpu": {"from": "#6366f1", "to": "#4338ca", "glow": "#a5b4fc"},
    "database": {"from": "#0ea5e9", "to": "#0369a1", "glow": "#7dd3fc"},
    "network": {"from": "#14b8a6", "to": "#0f766e", "glow": "#5eead4"},
    "chart": {"from": "#10b981", "to": "#047857", "glow": "#6ee7b7"},
    "globe": {"from": "#14b8a6", "to": "#0f766e", "glow": "#5eead4"},
    "website": {"from": "#3b82f6", "to": "#1d4ed8", "glow": "#93c5fd"},
    "phone": {"from": "#22c55e", "to": "#15803d", "glow": "#86efac"},
    "map-pin": {"from": "#f43f5e", "to": "#be123c", "glow": "#fda4af"},
    "info": {"from": "#3b82f6", "to": "#1d4ed8", "glow": "#93c5fd"},
    "hand-heart": {"from": "#f43f5e", "to": "#be123c", "glow": "#fda4af"},
    "palette": {"from": "#f97316", "to": "#c2410c", "glow": "#fdba74"},
    "building": {"from": "#64748b", "to": "#334155", "glow": "#94a3b8"},
    "music": {"from": "#8b5cf6", "to": "#6d28d9", "glow": "#c4b5fd"},
    "film": {"from": "#0ea5e9", "to": "#0369a1", "glow": "#7dd3fc"},
    "dna": {"from": "#ec4899", "to": "#be185d", "glow": "#f9a8d4"},
    "telescope": {"from": "#6366f1", "to": "#4338ca", "glow": "#a5b4fc"},
    "mountain": {"from": "#a16207", "to": "#78350f", "glow": "#d97706"},
    "leaf": {"from": "#22c55e", "to": "#15803d", "glow": "#86efac"},
    "scroll": {"from": "#a78bfa", "to": "#6d28d9", "glow": "#c4b5fd"},
    "brain": {"from": "#f472b6", "to": "#db2777", "glow": "#f9a8d4"},
    "activity": {"from": "#ef4444", "to": "#b91c1c", "glow": "#fca5a5"},
    "microscope": {"from": "#64748b", "to": "#334155", "glow": "#94a3b8"},
    "pill": {"from": "#06b6d4", "to": "#0e7490", "glow": "#67e8f9"},
    "scissors": {"from": "#f59e0b", "to": "#d97706", "glow": "#fbbf24"},
    "baby": {"from": "#38bdf8", "to": "#0284c7", "glow": "#7dd3fc"},
}

CATEGORY_META: list[dict[str, str]] = [
    {
        "title": "İman və mənəviyyat",
        "slug": "iman-ve-meneviyyat",
        "blurb": "İnanç, ibadət və ruhi yol axtarışına dair hekayələr.",
        "icon": "sparkles",
    },
    {
        "title": "Ailə və tərbiyə",
        "slug": "aile-ve-terbiye",
        "blurb": "Ana-ata sevgisi, övlad tərbiyəsi və yuva qurmaq.",
        "icon": "home",
    },
    {
        "title": "Sevgi və evlilik",
        "slug": "sevgi-ve-evlilik",
        "blurb": "Eşq, vəfa və birgə ömrün sınaqları.",
        "icon": "heart",
    },
    {
        "title": "Dostluq və insan münasibətləri",
        "slug": "dostluq-ve-insan-munasibetleri",
        "blurb": "Dost seçimi, mərhəmət və insanlara münasibət.",
        "icon": "users",
    },
    {
        "title": "Əxlaq və xarakter",
        "slug": "exlaq-ve-xarakter",
        "blurb": "Dürüstlük, səxavət və şəxsi kamillik.",
        "icon": "shield",
    },
    {
        "title": "Söz, sükut və ünsiyyət",
        "slug": "soz-sukut-ve-unsiyyet",
        "blurb": "Sözün gücü, sükutun hikməti və doğru ünsiyyət.",
        "icon": "message",
    },
    {
        "title": "Hikmət və həyat dərsləri",
        "slug": "hikmet-ve-heyat-dersleri",
        "blurb": "Həyatın mənası, seçimlər və düşüncə tərzi.",
        "icon": "lightbulb",
    },
    {
        "title": "Əmək, ruzi və sərvət",
        "slug": "emek-ruzi-ve-servet",
        "blurb": "Zəhmət, ruzi və var-dövlətə baxış.",
        "icon": "briefcase",
    },
    {
        "title": "Ədalət və cəmiyyət",
        "slug": "edalet-ve-cemiyyet",
        "blurb": "Haqq, cəmiyyət və ortaq məsuliyyət.",
        "icon": "scale",
    },
    {
        "title": "Yaşlanma və zaman",
        "slug": "yaslanma-ve-zaman",
        "blurb": "Ömrün fəsilləri və zamanın dəyişdirdikləri.",
        "icon": "clock",
    },
    {
        "title": "Təmsillər və məsəllər",
        "slug": "temsiller-ve-meseller",
        "blurb": "Heyvanlar və obrazlarla öyrədən məsəllər.",
        "icon": "book",
    },
    {
        "title": "Tarix və tanınmış şəxsiyyətlər",
        "slug": "tarix-ve-taninmish-shexsiyyetler",
        "blurb": "Tarixdən və tanınmış adlardan ibrət.",
        "icon": "landmark",
    },
]
# Preserve AZ category titles for STORY_CATEGORY grouping after locale swap.
for _meta in CATEGORY_META:
    _meta.setdefault("_az_title", _meta["title"])


def menu_icon(name: str) -> str:
    paths = CATEGORY_ICONS.get(name, CATEGORY_ICONS["book"])
    colors = ICON_COLORS.get(name, ICON_COLORS["book"])
    return (
        f'<span class="menu-icon menu-icon--{esc(name)}" aria-hidden="true" '
        f'style="--icon-from:{colors["from"]};--icon-to:{colors["to"]};--icon-glow:{colors["glow"]}">'
        f'<svg class="menu-icon__svg" viewBox="0 0 24 24" width="18" height="18" '
        f'fill="none" stroke="#fff" stroke-width="2.15" '
        f'stroke-linecap="round" stroke-linejoin="round">'
        f"{paths}</svg></span>"
    )


def extract_paragraphs(docx_path: Path, title: str) -> list[str]:
    doc = Document(str(docx_path))
    paras: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            paras.append(t)
    while paras and paras[0].casefold() == title.casefold():
        paras = paras[1:]
    return paras


def load_catalog() -> dict:
    if not stories_ready(LANG):
        return {
            "site_title": SITE_TITLE,
            "nav_label": NAV_STORIES_LABEL,
            "categories": [{**meta, "count": 0, "stories": []} for meta in CATEGORY_META],
        }

    data = json.loads(MAP_JSON.read_text(encoding="utf-8"))
    by_stem = {r["en_stem"]: r for r in data["rows"]}
    assert [c["_az_title"] for c in CATEGORY_META] == CATEGORY_ORDER

    grouped: dict[str, list[dict]] = defaultdict(list)
    for stem, row in by_stem.items():
        cat = STORY_CATEGORY[stem]
        story_path = STORIES / f"{stem}.docx"
        if not story_path.is_file():
            continue
        raw_paras = [
            (p.text or "").strip()
            for p in Document(str(story_path)).paragraphs
            if (p.text or "").strip()
        ]
        if not raw_paras:
            continue
        title = raw_paras[0]
        paras = raw_paras[1:]
        override = story_source_override(stem)
        if override:
            if paras:
                paras[-1] = override
            else:
                paras.append(override)
        elif paras and is_source_paragraph(paras[-1]):
            paras[-1] = normalize_source_paragraph(paras[-1])
        ill_rel = f"../illustrations/{stem}.webp"
        grouped[cat].append(
            {
                "stem": stem,
                "title": title,
                "paragraphs": paras,
                "image": ill_rel,
                "image_from_root": f"illustrations/{stem}.webp",
                "has_image": (ILLUSTRATIONS / f"{stem}.webp").is_file(),
            }
        )

    categories = []
    for meta in CATEGORY_META:
        stories = sorted(
            grouped.get(meta["_az_title"], []), key=lambda s: s["title"].casefold()
        )
        categories.append({**meta, "count": len(stories), "stories": stories})

    return {
        "site_title": SITE_TITLE,
        "nav_label": NAV_STORIES_LABEL,
        "categories": categories,
    }


def esc(s: str) -> str:
    return html.escape(s, quote=True)


_SOURCE_RE = re.compile(
    r"(internet\s+sources|internet\s+mənb|internet\s+kaynak|"
    r"открыт\w*\s+источник|интернет|"
    r"(?:source|mənbə|kaynak|источник|булак|булагы)\s*:)",
    re.I,
)
_MORAL_RE = re.compile(r"^(ibrət|ibret|moral|мораль|үлгү)\s*:", re.I)


def _fold_az_i(text: str) -> str:
    """Map dotted/dotless I so AZ/TR İ matches ASCII i in role regexes."""
    return (text or "").replace("İ", "i").replace("I", "i").replace("ı", "i")


_SOURCE_QUOTE_PAIRS = (("«", "»"), ("„", "“"), ("“", "”"), ('"', '"'), ("'", "'"))

# Last-line author credit instead of the generic "Source: Internet" label.
_AUTHOR_SOURCE_BY_LANG = {
    "en": "Bakhtiyar Sirajov",
    "ru": "Бахтияр Сираджов",
    "az": "Bəxtiyar Siracov",
    "ky": "Бахтияр Сиражов",
}
_STORY_SOURCE_BY_STEM = {
    "everyone-has-work-to-do": _AUTHOR_SOURCE_BY_LANG,
    "weeds-must-be-pulled-from-the-root": _AUTHOR_SOURCE_BY_LANG,
    "the-silent-corridor": _AUTHOR_SOURCE_BY_LANG,
}


def story_source_override(stem: str, lang: str | None = None) -> str:
    """Return a per-stem author credit for the current (or given) locale."""
    by_lang = _STORY_SOURCE_BY_STEM.get(stem or "") or {}
    code = (lang or LANG or "").strip().lower()
    return str(by_lang.get(code) or "").strip()


def is_source_paragraph(text: str) -> bool:
    return bool(_SOURCE_RE.search(_fold_az_i((text or "").strip().strip("«»\"“”"))))


def normalize_source_paragraph(text: str, stem: str = "") -> str:
    """Replace a detected source line with the locale story_source label."""
    override = story_source_override(stem)
    if override:
        return override
    label = str(UI.get("story_source") or "").strip()
    if label:
        return label
    s = (text or "").strip()
    changed = True
    while changed and s:
        changed = False
        for left, right in _SOURCE_QUOTE_PAIRS:
            if len(s) >= len(left) + len(right) and s.startswith(left) and s.endswith(right):
                s = s[len(left) : -len(right)].strip()
                changed = True
                break
    s = s.strip("«»\"“”„''")
    s = s.rstrip(".。").rstrip()
    s = s.replace("История взята", "Рассказ взят")
    return s


def story_paragraph_roles(paragraphs: list[str], stem: str = "") -> list[str]:
    """Return a role per paragraph: '', 'moral', or 'source'."""
    if not paragraphs:
        return []
    last = len(paragraphs) - 1
    roles = [""] * len(paragraphs)
    last_is_source = bool(story_source_override(stem)) or is_source_paragraph(
        paragraphs[last]
    )
    moral_i = None
    search_from = last - 1 if last_is_source else last
    for i in range(search_from, -1, -1):
        if _MORAL_RE.match(_fold_az_i((paragraphs[i] or "").strip())):
            moral_i = i
            break
    if moral_i is None:
        moral_i = last - 1 if last_is_source and last >= 1 else last
    if last_is_source:
        roles[last] = "source"
    if moral_i is not None and 0 <= moral_i < len(paragraphs) and roles[moral_i] != "source":
        roles[moral_i] = "moral"
    return roles


def story_paragraphs_html(paragraphs: list[str], stem: str = "") -> str:
    """Render story body; highlight the moral and de-emphasize the source line."""
    if not paragraphs:
        return ""
    roles = story_paragraph_roles(paragraphs, stem)
    parts: list[str] = []
    for p, role in zip(paragraphs, roles):
        if role == "source":
            p = normalize_source_paragraph(p, stem)
        if role:
            parts.append(f'<p class="story__{role}">{esc(p)}</p>')
        else:
            parts.append(f"<p>{esc(p)}</p>")
    return "".join(parts)


def source_attribution_html() -> str:
    """Home-page disclaimer with intro__source styling."""
    text = t_ui(
        "intro_source",
        "Hekayələr açıq İnternet mənbələrindən əldə olunub.",
    )
    if not str(text or "").strip():
        return ""
    return (
        f'<p class="intro__source">'
        f'<span class="intro__source-ornament" aria-hidden="true"></span>'
        f'<span class="intro__source-text">{esc(text)}</span>'
        f"</p>"
    )


def breadcrumbs_html(crumbs: list[tuple[str, str | None]], prefix: str) -> str:
    """crumbs: list of (label, href_or_None). Last item is current page."""
    items = []
    for i, (label, href) in enumerate(crumbs):
        is_last = i == len(crumbs) - 1
        if is_last:
            items.append(
                f'<li class="breadcrumbs__item" aria-current="page">'
                f"<span>{esc(label)}</span></li>"
            )
        elif href:
            items.append(
                f'<li class="breadcrumbs__item">'
                f'<a href="{esc(href)}">{esc(label)}</a></li>'
            )
        else:
            items.append(
                f'<li class="breadcrumbs__item">'
                f"<span>{esc(label)}</span></li>"
            )
    return f"""
<nav class="breadcrumbs" aria-label="Breadcrumb">
  <div class="breadcrumbs__inner">
    <ol class="breadcrumbs__list">
      {"".join(items)}
    </ol>
  </div>
</nav>
""".strip()


def nav_html(active_slug: str | None, prefix: str) -> str:
    coming_soon = esc(t_ui("coming_soon", "Tezliklə"))
    # Mega panel: 3 columns (Forumlar-style), contiguous chunks.
    n_cols = 3
    cats = list(CATEGORY_META)
    col_size = (len(cats) + n_cols - 1) // n_cols
    columns: list[list[dict[str, str]]] = [
        cats[i : i + col_size] for i in range(0, len(cats), col_size)
    ]
    mega_cols = []
    for col in columns:
        links = []
        for c in col:
            href = f"{prefix}categories/{c['slug']}.html"
            active = " is-active" if c["slug"] == active_slug else ""
            links.append(
                f'<a class="nav-dropdown-link{active}" href="{href}">'
                f'{menu_icon(c["icon"])}'
                f'<span class="nav-dropdown-link-copy">'
                f'<span class="nav-dropdown-link-title">{esc(c["title"])}</span>'
                f'<span class="nav-dropdown-link-desc">{esc(c["blurb"])}</span>'
                f"</span>"
                f"</a>"
            )
        mega_cols.append(
            '<div class="nav-mega-col"><div class="nav-mega-links">'
            + "".join(links)
            + "</div></div>"
        )
    mega_grid = '<div class="nav-mega-grid">' + "".join(mega_cols) + "</div>"
    arts_links = "".join(
        f'<a class="nav-dropdown-link" href="#" aria-disabled="true" tabindex="-1" title="{coming_soon}">'
        f'{menu_icon(item["icon"])}'
        f'<span class="nav-dropdown-link-copy">'
        f'<span class="nav-dropdown-link-title">{esc(item["label"])}</span>'
        f'<span class="nav-dropdown-link-desc">{esc(item["desc"])}</span>'
        f"</span>"
        f"</a>"
        for item in NAV_ARTS_ITEMS
    )
    def science_child_links(items: list[dict[str, str]]) -> str:
        return "".join(
            f'<a class="nav-dropdown-link" href="#" aria-disabled="true" tabindex="-1" title="{coming_soon}">'
            f'{menu_icon(child["icon"])}'
            f'<span class="nav-dropdown-link-copy">'
            f'<span class="nav-dropdown-link-title">{esc(child["label"])}</span>'
            f'<span class="nav-dropdown-link-desc">{esc(child["desc"])}</span>'
            f"</span>"
            f"</a>"
            for child in items
        )

    def science_mega_grid(items: list[dict[str, str]], *, n_cols: int = 3, grid_class: str = "") -> str:
        # Contiguous chunks (same approach as stories mega), preserving item order.
        col_size = (len(items) + n_cols - 1) // n_cols
        columns: list[list[dict[str, str]]] = [
            items[i : i + col_size] for i in range(0, len(items), col_size)
        ]
        cols_html = []
        for col in columns:
            cols_html.append(
                '<div class="nav-mega-col"><div class="nav-mega-links">'
                + science_child_links(col)
                + "</div></div>"
            )
        cls = "nav-mega-grid" + (f" {grid_class}" if grid_class else "")
        return f'<div class="{cls}">' + "".join(cols_html) + "</div>"

    science_parts: list[str] = []
    for item in NAV_SCIENCE_ITEMS:
        branch = item.get("branch")
        if branch and branch in NAV_SCIENCE_BRANCHES:
            panel_id = f"{branch}-science-mega-panel"
            branch_items = NAV_SCIENCE_BRANCHES[branch]
            # Multi-column mega (İbrətamiz / Tibb pattern). Flat links stay outside branches.
            # Contiguous chunks preserve item list order.
            mega_layout = {
                "humanities": (3, "nav-mega-grid--humanities"),  # 8 → 3+3+2
                "social": (2, "nav-mega-grid--social"),  # 5 → 3+2
                "natural": (2, "nav-mega-grid--natural"),  # 6 → 3+3
                "medical": (3, "nav-mega-grid--medical"),  # 13 → 5+5+3
                "informatics": (3, "nav-mega-grid--informatics"),  # 8 → 3+3+2
                "math": (3, "nav-mega-grid--math"),  # 8 → 3+3+2
            }
            n_cols, grid_class = mega_layout[branch]
            panel_body = science_mega_grid(
                branch_items, n_cols=n_cols, grid_class=grid_class
            )
            science_parts.append(
                f'<div class="nav-dropdown--nested nav-dropdown--has-mega" data-nav-branch="{esc(branch)}">'
                f'<button type="button" class="nav-dropdown-toggle" aria-expanded="false" '
                f'aria-controls="{panel_id}" data-nav-mega-toggle>'
                f'{menu_icon(item["icon"])}'
                f'<span class="nav-dropdown-toggle__copy">'
                f'<span class="nav-dropdown-link-title">{esc(item["label"])}</span>'
                f'<span class="nav-dropdown-link-desc">{esc(item["desc"])}</span>'
                f"</span>"
                f'<span class="nav-dropdown-caret" aria-hidden="true"></span>'
                f"</button>"
                f'<div class="nav-dropdown-panel nav-dropdown-panel--mega nav-dropdown-panel--{esc(branch)}" '
                f'id="{panel_id}">'
                f"{panel_body}"
                f"</div>"
                f"</div>"
            )
        else:
            science_parts.append(
                f'<a class="nav-dropdown-link" href="#" aria-disabled="true" tabindex="-1" title="{coming_soon}">'
                f'{menu_icon(item["icon"])}'
                f'<span class="nav-dropdown-link-copy">'
                f'<span class="nav-dropdown-link-title">{esc(item["label"])}</span>'
                f'<span class="nav-dropdown-link-desc">{esc(item["desc"])}</span>'
                f"</span>"
                f"</a>"
            )
    science_links = "".join(science_parts)
    nested_open = " is-mega-open" if active_slug else ""
    nested_expanded = "true" if active_slug else "false"
    home = f"{prefix}index.html"
    data_url = f"{prefix}assets/search-index.js?v={ASSET_VERSION}"
    top_links = "\n".join(
        f'<a class="primary-nav__link" href="#" aria-disabled="true" tabindex="-1" title="{esc(t_ui("coming_soon", "Tezliklə"))}">'
        f'{menu_icon(item["icon"])}'
        f'<span>{esc(item["label"])}</span>'
        f"</a>"
        for item in TOP_NAV_LINKS
    )
    # Language dropdown is built from languages.json. The menu uses the Popover API
    # (manual) so a sticky header / overflow-x:clip cannot hide it on iOS Safari.
    current_lang = language_by_code(LANG) or {"code": LANG, "name": LANG, "flag": ""}
    flag_root = "../../" if prefix == "../" else "../"

    def lang_flag_html(meta: dict) -> str:
        rel = str(meta.get("flag") or "")
        if not rel:
            return '<span class="lang-switcher__flag-slot" aria-hidden="true"></span>'
        return (
            f'<img class="lang-switcher__flag" src="{esc(flag_root + rel)}" alt="" '
            f'width="20" height="14" decoding="async" />'
        )

    option_html = []
    for meta in switcher_languages():
        code = str(meta["code"])
        label = esc(meta.get("name") or meta.get("label") or code)
        if not is_implemented(meta):
            option_html.append(
                f'<span class="lang-switcher__option" role="option" aria-disabled="true" '
                f'title="{coming_soon}">'
                f"{lang_flag_html(meta)}<span>{label}</span></span>"
            )
            continue
        href = (
            f"../../{code}/categories/{active_slug}.html"
            if active_slug
            else f"../{code}/index.html"
        )
        selected = "true" if code == LANG else "false"
        option_html.append(
            f'<a class="lang-switcher__option" role="option" href="{esc(href)}" '
            f'hreflang="{esc(code)}" data-lang="{esc(code)}" aria-selected="{selected}">'
            f"{lang_flag_html(meta)}<span>{label}</span></a>"
        )
    lang_switcher = (
        f'<nav class="lang-switcher" aria-label="{esc(t_ui("lang_switcher_label", "Dil"))}">'
        f'<button type="button" class="lang-switcher__toggle" aria-expanded="false" '
        f'aria-haspopup="listbox" aria-controls="lang-switcher-menu">'
        f"{lang_flag_html(current_lang)}"
        f'<span class="lang-switcher__name">{esc(str(current_lang.get("name") or LANG))}</span>'
        f'<span class="lang-switcher__caret" aria-hidden="true"></span>'
        f"</button>"
        f'<div class="lang-switcher__menu" id="lang-switcher-menu" role="listbox" popover="manual" hidden>'
        f"{''.join(option_html)}"
        f"</div>"
        f"</nav>"
    )
    return f"""
<header class="site-header">
  <div class="site-header__inner">
    <button type="button" class="nav-toggle" id="nav-toggle" aria-expanded="false" aria-controls="primaryNav" aria-label="{esc(t_ui("open_menu", "Menyunu aç"))}">
      <span></span><span></span><span></span>
    </button>
    <a class="brand" href="{home}">
      <img class="brand__logo" src="{shared_asset_href(prefix, "pearl.webp")}" alt="" width="40" height="40" />
      <span class="brand__name">{esc(SITE_NAME)}</span>
    </a>
    <nav class="primary-nav" id="primaryNav" aria-label="{esc(t_ui("main_menu", "Əsas menyu"))}">
      <details class="nav-dropdown nav-dropdown--literature">
        <summary class="nav-dropdown__summary">
          {menu_icon("book")}
          <span>{esc(NAV_LABEL)}</span>
        </summary>
        <div class="nav-dropdown-panel">
          <div class="nav-dropdown--nested nav-dropdown--has-mega{nested_open}" data-nav-branch="stories">
            <button type="button" class="nav-dropdown-toggle" aria-expanded="{nested_expanded}" aria-controls="literature-mega-panel" data-nav-mega-toggle>
              {menu_icon("layers")}
              <span class="nav-dropdown-toggle__copy">
                <span class="nav-dropdown-link-title">{esc(NAV_STORIES_LABEL)}</span>
                <span class="nav-dropdown-link-desc">{esc(NAV_STORIES_DESC)}</span>
              </span>
              <span class="nav-dropdown-caret" aria-hidden="true"></span>
            </button>
            <div class="nav-dropdown-panel nav-dropdown-panel--mega" id="literature-mega-panel">
              {mega_grid}
            </div>
          </div>
        </div>
      </details>
      <details class="nav-dropdown nav-dropdown--science">
        <summary class="nav-dropdown__summary">
          {menu_icon("atom")}
          <span>{esc(NAV_SCIENCE_LABEL)}</span>
        </summary>
        <div class="nav-dropdown-panel nav-dropdown-panel--science">
          <div class="nav-mega-links nav-mega-links--science">
            {science_links}
          </div>
        </div>
      </details>
      <details class="nav-dropdown nav-dropdown--arts">
        <summary class="nav-dropdown__summary">
          {menu_icon("palette")}
          <span>{esc(NAV_ARTS_LABEL)}</span>
        </summary>
        <div class="nav-dropdown-panel nav-dropdown-panel--arts">
          <div class="nav-mega-links nav-mega-links--arts">
            {arts_links}
          </div>
        </div>
      </details>
      {top_links}
    </nav>
    <div class="site-header__actions">
      {lang_switcher}
      <button type="button" class="global-search-toggle" id="global-search-toggle" aria-expanded="false" aria-controls="global-search" title="{esc(t_ui("global_search_title_attr", "Axtar (Ctrl+K)"))}" aria-label="{esc(t_ui("global_search_toggle", "Qlobal axtarış, Ctrl+K"))}">
        <svg viewBox="0 0 24 24" width="22" height="22" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">
          <circle cx="11" cy="11" r="7"></circle>
          <path d="m20 20-3.5-3.5"></path>
        </svg>
        <span class="global-search-toggle__label">{esc(t_ui("search", "Axtar…"))}</span>
        <kbd class="global-search-toggle__kbd">Ctrl+K</kbd>
      </button>
    </div>
  </div>
</header>
<div class="global-search" id="global-search" hidden data-search-index="{esc(data_url)}">
  <button type="button" class="global-search__backdrop" data-global-search-close tabindex="-1" aria-label="{esc(t_ui("close_search", "Axtarışı bağla"))}"></button>
  <div class="global-search__panel" role="dialog" aria-modal="true" aria-labelledby="global-search-title">
    <div class="global-search__head">
      <p id="global-search-title" class="global-search__title">{esc(t_ui("global_search", "Qlobal axtarış"))}</p>
      <button type="button" class="global-search__close" data-global-search-close aria-label="{esc(t_ui("close", "Bağla"))}">×</button>
    </div>
    <label class="global-search__field">
      <span class="visually-hidden">{esc(t_ui("search_stories_label", "Hekayə axtar"))}</span>
      <input type="search" id="global-search-input" placeholder="{esc(t_ui("search_stories_placeholder", "Bütün hekayələrdə axtar…"))}" autocomplete="off" />
    </label>
    <p class="global-search__status" id="global-search-status" aria-live="polite"></p>
    <div class="global-search__results" id="global-search-results"></div>
  </div>
</div>
""".strip()


def _footer_contact_icon(name: str) -> str:
    return menu_icon(name)


def page_shell(
    *,
    title: str,
    description: str,
    active_slug: str | None,
    prefix: str,
    body: str,
    crumbs: list[tuple[str, str | None]],
    extra_body_class: str = "",
) -> str:
    page_title = title if SITE_NAME in title else f"{title} · {SITE_NAME}"
    html_lang = (LOCALE.get("html_lang") if LOCALE else None) or LANG or "az"
    return f"""<!DOCTYPE html>
<html lang="{esc(html_lang)}">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1, viewport-fit=cover" />
  <meta name="theme-color" content="#0069b4" />
  <meta name="color-scheme" content="light" />
  <title>{esc(page_title)}</title>
  <meta name="description" content="{esc(description)}" />
  <link rel="icon" href="{shared_asset_href(prefix, "favicon-32.png")}" type="image/png" sizes="32x32" />
  <link rel="icon" href="{shared_asset_href(prefix, "favicon-48.png")}" type="image/png" sizes="48x48" />
  <link rel="icon" href="{shared_asset_href(prefix, "favicon.png")}" type="image/png" sizes="192x192" />
  <link rel="icon" href="{shared_asset_href(prefix, "favicon.ico")}" sizes="any" />
  <link rel="apple-touch-icon" href="{shared_asset_href(prefix, "apple-touch-icon.png")}" sizes="180x180" />
  <link rel="preconnect" href="https://fonts.googleapis.com" />
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin />
  <link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,500;9..144,700&family=Source+Serif+4:opsz,wght@8..60,400;8..60,600&family=Source+Sans+3:wght@400;500;600;700&display=swap" rel="stylesheet" />
  <link rel="stylesheet" href="{shared_asset_href(prefix, "site.css")}?v={ASSET_VERSION}" />
</head>
<body class="{extra_body_class}" id="top" data-lang="{esc(LANG)}">
  <a class="skip-link" href="#main">{esc(t_ui("skip_to_content", "Məzmuna keç"))}</a>
  {nav_html(active_slug, prefix)}
  {breadcrumbs_html(crumbs, prefix)}
  <main id="main">
  {body}
  </main>
  <a class="back-to-top" href="#top" id="back-to-top" title="{esc(t_ui("back_to_top", "Səhifənin yuxarısına qayıt"))}" aria-label="{esc(t_ui("back_to_top", "Səhifənin yuxarısına qayıt"))}"></a>
  <footer class="footer-pro">
    <div class="footer-inner">
      <div class="footer-grid">
        <div class="footer-col footer-col--about">
          <p class="footer-about">{esc(t_ui("intro_lead", "Saytımızda bəşəriyyətin tarix boyu elm və texnologiya, təbiət elmləri, ictimai və humanitar elmlər, eləcə də ədəbiyyat və incəsənətin müxtəlif sahələrində qazandığı möhtəşəm nailiyyətlər, ümumbəşəri mənəvi dəyərlər, görkəmli şəxsiyyətlər, mühüm tarixi kəşf və ixtiralar haqqında zəngin məlumatlar təqdim olunur. Niyyətimiz əsrlər boyu toplanmış bu dəyərli irsi qorumaq, sistemləşdirmək və gələcək nəsillərə bilik, ibrət və ilham mənbəyi kimi çatdırmaqdır."))}</p>
        </div>
        <div class="footer-col footer-col--brand">
          <a class="footer-logo" href="{prefix}index.html">
            <img class="footer-logo__img" src="{shared_asset_href(prefix, "pearl.webp")}" alt="" width="72" height="72" decoding="async" />
            <span class="footer-logo__name">{esc(SITE_NAME)}</span>
          </a>
        </div>
        <div class="footer-col footer-col--contact">
          <h2 class="footer-contact__title">{esc(t_ui("footer_contact", "Əlaqə vasitələri"))}</h2>
          <ul class="footer-contact">
            <li>
              {_footer_contact_icon("phone")}
              <span class="footer-contact__label">{esc(t_ui("footer_phone", "Telefon"))}</span>
            </li>
            <li>
              {_footer_contact_icon("map-pin")}
              <span class="footer-contact__label">{esc(t_ui("footer_address", "Ünvan"))}</span>
            </li>
            <li>
              <a class="footer-contact__link" href="{esc(t_ui("footer_website_url", "https://birinci.cloud"))}" rel="noopener noreferrer">
                {_footer_contact_icon("website")}
                <span class="footer-contact__label">{esc(t_ui("footer_website", "Veb sayt"))}</span>
              </a>
            </li>
          </ul>
          <a class="footer-contact__url" href="{esc(t_ui("footer_website_url", "https://birinci.cloud"))}" rel="noopener noreferrer">birinci.cloud</a>
        </div>
      </div>
    </div>
    <div class="footer-bottom">
      <div class="footer-copy">© 2026 {esc(SITE_NAME)}</div>
    </div>
  </footer>
  <script src="{prefix}assets/site.js?v={ASSET_VERSION}" defer></script>
</body>
</html>
"""


def _tools_bar_glyph(name: str, extra_class: str = "") -> str:
    paths = {
        "eye": '<path d="M2 12s3.5-7 10-7 10 7 10 7-3.5 7-10 7S2 12 2 12Z"/><circle cx="12" cy="12" r="3"/>',
        "eye-off": '<path d="M3 3l18 18"/><path d="M10.6 10.6a3 3 0 0 0 4.2 4.2"/><path d="M9.9 5.1A11 11 0 0 1 12 5c6.5 0 10 7 10 7a19 19 0 0 1-3.2 4.1"/><path d="M6.1 6.1C3.6 7.8 2 12 2 12s3.5 7 10 7c1.6 0 3.1-.3 4.4-.9"/>',
        "text": '<path d="M7 3h7l5 5v13H7z"/><path d="M14 3v5h5"/><path d="M9 13h6"/><path d="M9 17h6"/>',
        "text-off": '<path d="M7 3h7l5 5v13H7z"/><path d="M14 3v5h5"/><path d="M9 13h6"/><path d="M9 17h6"/><path d="M5 5l14 14"/>',
        "chevron-left": '<path d="M15 6l-6 6 6 6"/>',
        "chevron-right": '<path d="M9 6l6 6-6 6"/>',
        "shuffle": '<path d="M16 3h5v5"/><path d="m21 3-7 7"/><path d="M4 20l7-7"/><path d="M16 16h5v5"/><path d="m21 21-6-6"/><path d="M4 4l5 5"/>',
        "grid": '<rect x="3" y="3" width="7" height="7" rx="1"/><rect x="14" y="3" width="7" height="7" rx="1"/><rect x="3" y="14" width="7" height="7" rx="1"/><rect x="14" y="14" width="7" height="7" rx="1"/>',
        "list": '<path d="M8 6h13"/><path d="M8 12h13"/><path d="M8 18h13"/><circle cx="4" cy="6" r="1"/><circle cx="4" cy="12" r="1"/><circle cx="4" cy="18" r="1"/>',
        "all": '<rect x="3" y="3" width="5" height="5" rx="1"/><rect x="9.5" y="3" width="5" height="5" rx="1"/><rect x="16" y="3" width="5" height="5" rx="1"/><rect x="3" y="9.5" width="5" height="5" rx="1"/><rect x="9.5" y="9.5" width="5" height="5" rx="1"/><rect x="16" y="9.5" width="5" height="5" rx="1"/><rect x="3" y="16" width="5" height="5" rx="1"/><rect x="9.5" y="16" width="5" height="5" rx="1"/><rect x="16" y="16" width="5" height="5" rx="1"/>',
        "play-visible": '<polygon points="5 4 15 12 5 20 5 4"/><path d="M18 6v12"/><path d="M21 8v8"/>',
        "listen": '<polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"/><path d="M15.54 8.46a5 5 0 0 1 0 7.07"/><path d="M19.07 4.93a10 10 0 0 1 0 14.14"/>',
        "stop": '<polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"/><path d="M15.54 8.46a5 5 0 0 1 0 7.07"/><path d="M19.07 4.93a10 10 0 0 1 0 14.14"/><path d="M3 3l18 18"/>',
    }
    cls = "tools-bar__glyph"
    if extra_class:
        cls = f"{cls} {extra_class}"
    return (
        f'<svg class="{cls}" viewBox="0 0 24 24" width="16" height="16" '
        f'aria-hidden="true" fill="none" stroke="currentColor" stroke-width="2" '
        f'stroke-linecap="round" stroke-linejoin="round">{paths[name]}</svg>'
    )


_STORY_ACTION_BTN = "tools-bar__view-btn tools-bar__view-btn--icon"


def _story_mode_pair(*, kind: str, controls_id: str = "") -> str:
    """Two-button story group: Show|Hide, or Listen|Stop. First button is the reveal/play mode."""
    if kind == "figure":
        mode_attr = "data-images-mode"
        first_label = esc(t_ui("show_image", "Şəkli göstər"))
        second_label = esc(t_ui("hide_image", "Şəkli gizlət"))
        group_label = esc(t_ui("story_image_label", "Şəkil"))
        first_glyph, second_glyph = "eye", "eye-off"
        first_mode, second_mode = "show", "hide"
        first_pressed, second_pressed = "true", "false"
        extra_class = ""
        extra_attrs = f' aria-controls="{controls_id}"' if controls_id else ""
    elif kind == "text":
        mode_attr = "data-texts-mode"
        first_label = esc(t_ui("show_text", "Mətni göstər"))
        second_label = esc(t_ui("hide_text", "Mətni gizlət"))
        group_label = esc(t_ui("story_text_label", "Mətn"))
        first_glyph, second_glyph = "text", "text-off"
        first_mode, second_mode = "show", "hide"
        first_pressed, second_pressed = "true", "false"
        extra_class = ""
        extra_attrs = f' aria-controls="{controls_id}"' if controls_id else ""
    elif kind == "tts":
        mode_attr = "data-tts-mode"
        first_label = esc(t_ui("listen", "Mətni dinlə"))
        second_label = esc(t_ui("stop", "Dayandır"))
        group_label = esc(t_ui("story_audio_label", "Səs"))
        first_glyph, second_glyph = "listen", "stop"
        first_mode, second_mode = "listen", "stop"
        first_pressed, second_pressed = "false", "true"
        extra_class = " story-tts"
        extra_attrs = " data-story-tts"
    else:
        raise ValueError(f"unsupported story mode pair: {kind}")
    return f"""
          <div class="story__action-group">
            <span class="tools-bar__label">{group_label}</span>
            <div class="tools-bar__views" role="group" aria-label="{group_label}">
              <button type="button" class="{_STORY_ACTION_BTN}{extra_class}" {mode_attr}="{first_mode}" aria-pressed="{first_pressed}"{extra_attrs} title="{first_label}" aria-label="{first_label}">{_tools_bar_glyph(first_glyph)}</button>
              <button type="button" class="{_STORY_ACTION_BTN}{extra_class}" {mode_attr}="{second_mode}" aria-pressed="{second_pressed}"{extra_attrs} title="{second_label}" aria-label="{second_label}">{_tools_bar_glyph(second_glyph)}</button>
            </div>
          </div>"""


def tools_bar_html(*, mode: str = "home") -> str:
    """Shared tools bar for home and category pages.

    mode:
      - "home": Axtar + Görüntü (Təsnifatlı/Ardıcıl) + list-only Şəkillər/Mətnlər/Sayı
      - "category": same chrome without Görüntü (page is already a single-category list);
        Şəkillər/Mətnlər/Sayı are always visible
    """
    if mode not in ("home", "category"):
        raise ValueError(f"unsupported tools bar mode: {mode}")
    is_home = mode == "home"
    list_only_attr = " data-home-list-only hidden" if is_home else ""
    view_block = ""
    cards_label = esc(t_ui("view_cards", "Təsnifatlı"))
    list_label = esc(t_ui("view_list", "Ardıcıl"))
    if is_home:
        view_block = f"""
  <div class="tools-bar__field">
    <span class="tools-bar__label" id="home-view-label">{esc(t_ui("view", "Görüntü"))}</span>
    <div class="tools-bar__views" role="group" aria-labelledby="home-view-label">
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-home-view="cards" aria-pressed="true" title="{cards_label}" aria-label="{cards_label}" onclick="return window.__birinciSetHomeView ? window.__birinciSetHomeView('cards') : false">{_tools_bar_glyph("grid")}</button>
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-home-view="list" aria-pressed="false" title="{list_label}" aria-label="{list_label}" onclick="return window.__birinciSetHomeView ? window.__birinciSetHomeView('list') : false">{_tools_bar_glyph("list")}</button>
    </div>
  </div>
""".rstrip()
    view_block_html = f"\n  {view_block}" if view_block else ""
    show_label = esc(t_ui("show", "Göstər"))
    hide_label = esc(t_ui("hide", "Gizlət"))
    count_label = esc(t_ui("batch_count", "Hekayə sayı"))
    prev_label = esc(t_ui("batch_prev", "Əvvəlki"))
    next_label = esc(t_ui("batch_next", "Növbəti"))
    random_label = esc(t_ui("batch_random", "Təsadüfi"))
    all_label = esc(t_ui("batch_all", "Hamısı"))
    listen_page_label = esc(t_ui("listen_page", "Səhifəni dinlə"))
    stop_label = esc(t_ui("stop", "Dayandır"))
    clear_filter_label = esc(t_ui("clear_search_filter", "Filtri təmizlə"))
    return f"""
<div class="tools-bar{" tools-bar--dense" if not is_home else ""}" data-tools="{esc(mode)}">
  <div class="tools-bar__search">
    <label class="tools-bar__search-field">
      <span class="visually-hidden">{esc(t_ui("search_aria", "Axtar"))}</span>
      <input type="search" data-tools-search placeholder="{esc(t_ui("search", "Axtar…"))}" autocomplete="off" />
    </label>
    <div class="tools-bar__search-chip" data-search-filter hidden>
      <span class="tools-bar__search-chip-dot" aria-hidden="true"></span>
      <span class="tools-bar__search-chip-text" data-search-filter-text aria-live="polite"></span>
      <button type="button" class="tools-bar__view-btn tools-bar__search-clear" data-search-filter-clear title="{clear_filter_label}" aria-label="{clear_filter_label}">×</button>
    </div>
  </div>{view_block_html}
  <div class="tools-bar__field"{list_only_attr}>
    <span class="tools-bar__label" id="tools-images-label">{esc(t_ui("images", "Şəkillər"))}</span>
    <div class="tools-bar__views tools-bar__images-toggle" role="group" aria-labelledby="tools-images-label" data-tools-images>
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-images-mode="show" aria-pressed="true" title="{show_label}" aria-label="{show_label}">{_tools_bar_glyph("eye")}</button>
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-images-mode="hide" aria-pressed="false" title="{hide_label}" aria-label="{hide_label}">{_tools_bar_glyph("eye-off")}</button>
    </div>
  </div>
  <div class="tools-bar__field"{list_only_attr}>
    <span class="tools-bar__label" id="tools-texts-label">{esc(t_ui("texts", "Mətnlər"))}</span>
    <div class="tools-bar__views tools-bar__texts-toggle" role="group" aria-labelledby="tools-texts-label" data-tools-texts>
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-texts-mode="show" aria-pressed="true" title="{show_label}" aria-label="{show_label}">{_tools_bar_glyph("text")}</button>
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-texts-mode="hide" aria-pressed="false" title="{hide_label}" aria-label="{hide_label}">{_tools_bar_glyph("text-off")}</button>
    </div>
  </div>
  <div class="tools-bar__field tools-bar__batch"{list_only_attr}>
    <span class="tools-bar__label" id="tools-batch-label">{esc(t_ui("batch", "Hekayələrin sayı"))}</span>
    <div class="tools-bar__batch-controls" role="group" aria-labelledby="tools-batch-label">
      <div class="tools-bar__stepper">
        <button type="button" class="tools-bar__stepper-btn" data-home-batch="dec" title="{esc(t_ui("batch_dec", "Azalt"))}" aria-label="{esc(t_ui("batch_dec", "Azalt"))}">−</button>
        <label class="tools-bar__batch-count">
          <span class="visually-hidden">{count_label}</span>
          <input
            type="number"
            class="tools-bar__batch-input"
            data-home-batch-size
            min="1"
            max="9999"
            value="12"
            step="1"
            inputmode="numeric"
            aria-label="{count_label}"
            title="{count_label}"
          />
        </label>
        <button type="button" class="tools-bar__stepper-btn" data-home-batch="inc" title="{esc(t_ui("batch_inc", "Artır"))}" aria-label="{esc(t_ui("batch_inc", "Artır"))}">+</button>
      </div>
      <div class="tools-bar__pager" role="navigation" aria-label="{esc(t_ui("batch_pager", "Səhifə"))}">
        <button type="button" class="tools-bar__pager-btn" data-home-batch="prev" title="{prev_label}" aria-label="{prev_label}">{_tools_bar_glyph("chevron-left")}</button>
        <span class="tools-bar__batch-range" data-home-batch-range aria-live="polite"></span>
        <button type="button" class="tools-bar__pager-btn" data-home-batch="next" title="{next_label}" aria-label="{next_label}">{_tools_bar_glyph("chevron-right")}</button>
      </div>
      <button type="button" class="tools-bar__icon-btn" data-home-batch="random" title="{random_label}" aria-label="{random_label}">{_tools_bar_glyph("shuffle")}</button>
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon tools-bar__batch-all" data-home-batch="all" title="{all_label}" aria-label="{all_label}" aria-pressed="false">{_tools_bar_glyph("all")}</button>
    </div>
  </div>
  <div class="tools-bar__field"{list_only_attr}>
    <span class="tools-bar__label" id="tools-listen-page-label">{listen_page_label}</span>
    <div class="tools-bar__views" role="group" aria-labelledby="tools-listen-page-label">
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-tools-play-visible data-tts-mode="listen" aria-pressed="false" title="{listen_page_label}" aria-label="{listen_page_label}">{_tools_bar_glyph("listen")}</button>
      <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-tools-play-visible data-tts-mode="stop" aria-pressed="true" title="{stop_label}" aria-label="{stop_label}">{_tools_bar_glyph("stop")}</button>
    </div>
  </div>
</div>
""".strip()


def home_tools_bar_html() -> str:
    return tools_bar_html(mode="home")


def content_placeholder_html() -> str:
    return f"""
<div class="content-placeholder" role="status">
  <p class="content-placeholder__kicker">{esc(t_ui("coming_soon", "Tezliklə"))}</p>
  <p class="content-placeholder__title">{esc(t_ui("stories_coming_soon", "Hekayələr tezliklə"))}</p>
  <p class="content-placeholder__lead">{esc(t_ui("stories_coming_soon_lead", "Bu dil üçün hekayələr hazırlanır. Səhifə quruluşu artıq hazırdır; mətnlər əlavə olunanda burada görünəcək."))}</p>
</div>
""".strip()


def build_landing(catalog: dict) -> str:
    ready = stories_ready(LANG)
    count_suffix = t_ui("stories_count_suffix", "hekayə")
    coming_soon = t_ui("coming_soon", "Tezliklə")
    cards = []
    for c in catalog["categories"]:
        icon = menu_icon(c["icon"])
        meta = coming_soon if not ready else f"{c['count']} {count_suffix}"
        meta_class = " cat-card__meta--soon" if not ready else ""
        cards.append(
            f"""
<a class="cat-card page-card" href="categories/{esc(c['slug'])}.html" data-title="{esc(c['title'])}" data-blurb="{esc(c['blurb'])}" data-count="{c['count']}">
  <div class="card-icon-wrap" aria-hidden="true">{icon}</div>
  <div class="card-body">
    <h2 class="card-title">{esc(c['title'])}</h2>
    <div class="card-desc">{esc(c['blurb'])}</div>
    <span class="cat-card__meta{meta_class}">{esc(meta)}</span>
  </div>
</a>
""".strip()
        )

    stories_nav = t_ui("stories_nav", "Hekayələr")
    tools = home_tools_bar_html() if ready else ""
    view_bootstrap = HOME_VIEW_BOOTSTRAP if ready else ""
    placeholder = content_placeholder_html() if not ready else ""
    list_view = ""
    if ready:
        list_view = f"""
    <div data-view="list" hidden
         data-stories-url="data/stories.json"
         data-stories-script="assets/stories-data.js?v={esc(ASSET_VERSION)}"
         data-asset-version="{esc(ASSET_VERSION)}">
      <div class="category-layout home-stories-layout">
        <aside class="story-nav sidebar" aria-label="{esc(stories_nav)}">
          <div class="sidebar-widget">
            <div class="widget-head">
              <span><span aria-hidden="true">📖</span> {esc(stories_nav)}</span>
              <button type="button" class="events-menu-toggle" aria-controls="homeStoryNavMenu" aria-expanded="false" aria-label="{esc(t_ui("stories_nav_open", "Hekayələr menyusunu aç"))}">
                <span></span><span></span><span></span>
              </button>
            </div>
            <div class="widget-body">
              <nav>
                <ul class="timeline-list" id="homeStoryNavMenu" data-home-nav></ul>
              </nav>
            </div>
          </div>
        </aside>
        <div class="home-stories-main">
          <div class="story-list" data-stories-list></div>
        </div>
      </div>
      <p class="tools-empty" data-home-list-empty hidden>{esc(t_ui("no_matching_story", "Uyğun hekayə tapılmadı."))}</p>
    </div>
"""

    body = f"""
<div class="page-home__content">
  <section class="intro">
    <div class="intro__atmosphere" aria-hidden="true"></div>
    <div class="intro__content">
      <div class="intro__copy">
        <h1 class="intro__brand">Bir <span>inci</span></h1>
        <p class="intro__lead">{esc(t_ui("intro_lead", "Saytımızda bəşəriyyətin tarix boyu elm və texnologiya, təbiət elmləri, ictimai və humanitar elmlər, eləcə də ədəbiyyat və incəsənətin müxtəlif sahələrində qazandığı möhtəşəm nailiyyətlər, ümumbəşəri mənəvi dəyərlər, görkəmli şəxsiyyətlər, mühüm tarixi kəşf və ixtiralar haqqında zəngin məlumatlar təqdim olunur. Niyyətimiz əsrlər boyu toplanmış bu dəyərli irsi qorumaq, sistemləşdirmək və gələcək nəsillərə bilik, ibrət və ilham mənbəyi kimi çatdırmaqdır."))}</p>
        {source_attribution_html()}
      </div>
      <div class="intro__visual">
        <img src="{shared_asset_href("", "Pearl with Background 3.png")}?v={ASSET_VERSION}" alt="" width="1536" height="1024" decoding="async" />
      </div>
    </div>
  </section>

  <section id="kateqoriyalar" class="section categories home-browser" aria-labelledby="home-categories-title">
    <h2 id="home-categories-title" class="section__title visually-hidden">{esc(t_ui("categories_heading", "Kateqoriyalar"))}</h2>
    {placeholder}
    {tools}
    {view_bootstrap}
    <div data-view="cards">
      <div class="cat-grid" data-tools-list>
        {"".join(cards)}
      </div>
      <p class="tools-empty" data-tools-empty hidden>{esc(t_ui("no_matching_category", "Uyğun kateqoriya tapılmadı."))}</p>
    </div>
    {list_view}
  </section>
</div>
"""
    return page_shell(
        title=SITE_NAME,
        description=t_ui("site_description", "Bir inci — ibrətamiz deyimlər və hekayələr toplusu."),
        active_slug=None,
        prefix="",
        body=body,
        crumbs=[(HOME_CRUMB, None)],
        extra_body_class="page-home",
    )


def build_category_page(cat: dict) -> str:
    stories_html = []
    for s in cat["stories"]:
        paras = story_paragraphs_html(s["paragraphs"], s.get("stem") or "")
        img = f"../illustrations/{esc(s['stem'])}.webp"
        audio_file = AUDIO_DIR / f"{s['stem']}.mp3"
        audio_attr = (
            f' data-audio="../audio/{esc(s["stem"])}.mp3?v={ASSET_VERSION}"'
            if audio_file.is_file()
            else ""
        )
        figure_toggle = ""
        figure_html = ""
        if s.get("has_image"):
            figure_toggle = _story_mode_pair(kind="figure", controls_id=f"figure-{esc(s['stem'])}")
            figure_html = f"""
    <figure class="story__figure" id="figure-{esc(s['stem'])}">
      <button type="button" class="story__figure-open" aria-label="{esc(s['title'])} şəklini böyüt">
        <img src="{img}" alt="{esc(s['title'])} illüstrasiyası" loading="lazy" width="1536" height="1024" />
      </button>
    </figure>"""
        stories_html.append(
            f"""
<article class="story news-card" id="{esc(s['stem'])}" data-stem="{esc(s['stem'])}" data-title="{esc(s['title'])}"{audio_attr}>
  <div class="card-header">
    <h2 class="card-title story__title">{esc(s['title'])}</h2>
  </div>
  <div class="card-body">
    <div class="story__content">
      <div class="story__panel">
        <div class="story__actions">
          {_story_mode_pair(kind="tts")}
          {figure_toggle}
          {_story_mode_pair(kind="text", controls_id="text-" + esc(s["stem"]))}
          <p class="story-tts__note" data-story-tts-note hidden></p>
        </div>
        <div class="story__text card-text" id="text-{esc(s['stem'])}">
          {paras}
        </div>
      </div>
    </div>
    {figure_html}
  </div>
</article>
""".strip()
        )

    nav_items = "\n".join(
        f'<li data-stem="{esc(s["stem"])}" data-title="{esc(s["title"])}">'
        f'<a href="#{esc(s["stem"])}">{esc(s["title"])}</a></li>'
        for s in cat["stories"]
    )
    stories_nav = t_ui("stories_nav", "Hekayələr")
    count_suffix = t_ui("stories_count_suffix", "hekayə")
    coming_soon = t_ui("coming_soon", "Tezliklə")
    ready = bool(cat["stories"]) and stories_ready(LANG)
    hero_meta = coming_soon if not ready else f'<span data-tools-count>{cat["count"]}</span> {esc(count_suffix)}'
    if ready:
        tools = tools_bar_html(mode="category")
        body = f"""
<div class="category-page">
  <div class="category-layout">
    <header class="category-hero">
      <h1>{esc(cat['title'])}</h1>
      <p class="category-hero__lead">{esc(cat['blurb'])} · {hero_meta}</p>
    </header>
    {tools}
    <aside class="story-nav sidebar" aria-label="{esc(stories_nav)}">
      <div class="sidebar-widget">
        <div class="widget-head">
          <span><span aria-hidden="true">📖</span> {esc(stories_nav)}</span>
          <button type="button" class="events-menu-toggle" aria-controls="storyNavMenu" aria-expanded="false" aria-label="{esc(t_ui("stories_nav_open", "Hekayələr menyusunu aç"))}">
            <span></span><span></span><span></span>
          </button>
        </div>
        <div class="widget-body">
          <nav>
            <ul class="timeline-list" id="storyNavMenu" data-tools-nav>
              {nav_items}
            </ul>
          </nav>
        </div>
      </div>
    </aside>
    <div class="story-list" data-tools-list>
      {"".join(stories_html)}
    </div>
  </div>
  <p class="tools-empty" data-tools-empty hidden>{esc(t_ui("no_matching_story", "Uyğun hekayə tapılmadı."))}</p>
</div>
"""
    else:
        body = f"""
<div class="category-page">
  <div class="category-layout category-layout--placeholder">
    <header class="category-hero">
      <h1>{esc(cat['title'])}</h1>
      <p class="category-hero__lead">{esc(cat['blurb'])} · {esc(coming_soon)}</p>
    </header>
    {content_placeholder_html()}
  </div>
</div>
"""
    return page_shell(
        title=f"{cat['title']} · {SITE_NAME}",
        description=cat["blurb"],
        active_slug=cat["slug"],
        prefix="../",
        body=body,
        crumbs=[
            (HOME_CRUMB, "../index.html"),
            (NAV_LABEL, "../index.html#kateqoriyalar"),
            (NAV_STORIES_LABEL, "../index.html#kateqoriyalar"),
            (cat["title"], None),
        ],
        extra_body_class="page-category",
    )


CSS = r"""
:root {
  /* DAAB activities palette (daab-tokens.css) */
  --ink: #08263b;
  --ink-soft: #345f86;
  --color-surface-news: #f5fbff;
  --color-site-bg-scrim: rgba(255, 255, 255, 0.07);
  --site-bg-image: url("diaspor-body-top-bg.png");
  --site-bg-position: top center;
  --site-bg-size: cover;
  --site-bg-repeat: no-repeat;
  --surface: #ffffff;
  --surface-muted: #eef7fc;
  --nav-blue: #0069b4;
  --nav-blue-deep: #005a9a;
  --nav-blue-soft: #dff2ff;
  --panel-blue: #e5f4fb;
  --panel-blue-deep: #d9eef9;
  --blue-400: #4eb4ee;
  --blue-600: #117fc8;
  --blue-800: #133f63;
  --blue-900: #06314e;
  --blue-soft: #9ed6f5;
  --gold: #c99b3b;
  --gold-soft: #fff0bf;
  --gold-bright: #f0c75e;
  --gold-grad-from: #d4922a;
  --gold-grad-mid: #e8b04a;
  --accent: var(--nav-blue);
  --accent-hover: var(--nav-blue-deep);
  --line: rgba(0, 105, 180, 0.18);
  --line-strong: rgba(0, 105, 180, 0.26);
  --line-10: rgba(0, 105, 180, 0.10);
  --line-14: rgba(0, 105, 180, 0.14);
  --line-16: rgba(0, 105, 180, 0.16);
  --ring: rgba(0, 105, 180, 0.28);
  --shadow: 0 12px 32px rgba(0, 78, 140, 0.14);
  --font-display: "Fraunces", Georgia, serif;
  --font-body: "Source Serif 4", Georgia, serif;
  --font-ui: "Source Sans 3", "Segoe UI", "Noto Sans", sans-serif;
  --max: 1120px;
  --max-wide: 1280px;
  --radius: 16px;
  --radius-sm: 12px;
  --radius-card: 24px;
  --radius-pill: 999px;
  --header-h: 4.25rem;
  --breadcrumb-h: 2.7rem;
  /* Reserved for a future third sticky layer; syncStickyChrome sets header/breadcrumb only. */
  --sticky-stack-h: 0rem;
  --sticky-stack-bottom: calc(var(--header-h) + var(--breadcrumb-h) + var(--sticky-stack-h));
  --space-1: 0.25rem;
  --space-2: 0.5rem;
  --space-3: 0.75rem;
  --space-4: 1rem;
  --space-5: 1.5rem;
  --space-6: 2rem;
  --space-7: 3rem;
  --duration-fast: 140ms;
  --duration: 160ms;
  --ease-standard: ease;
  --ease-outro: cubic-bezier(0.2, 0.8, 0.2, 1);
}

*, *::before, *::after { box-sizing: border-box; }
html {
  min-height: 100%;
  scroll-behavior: auto;
  overflow-x: clip;
}
html.smooth-scroll { scroll-behavior: smooth; }
html.no-smooth-scroll,
html.no-smooth-scroll * {
  scroll-behavior: auto !important;
}
body {
  margin: 0;
  display: flex;
  flex-direction: column;
  color: var(--ink);
  background-color: var(--color-surface-news, #f5fbff);
  font-family: var(--font-body);
  line-height: 1.55;
  min-height: 100vh;
  text-rendering: optimizeLegibility;
  -webkit-font-smoothing: antialiased;
}

/* Site-wide embroidery pattern (DAAB treatment) */
body::before {
  content: "";
  position: fixed;
  inset: 0;
  z-index: -2;
  pointer-events: none;
  background-color: var(--color-surface-news, #f5fbff);
  background-image:
    linear-gradient(
      0deg,
      var(--color-site-bg-scrim, rgba(255, 255, 255, 0.07)),
      var(--color-site-bg-scrim, rgba(255, 255, 255, 0.07))
    ),
    var(--site-bg-image, url("diaspor-body-top-bg.png"));
  background-position: center, var(--site-bg-position, top center);
  background-size: 100% 100%, var(--site-bg-size, cover);
  background-repeat: no-repeat, var(--site-bg-repeat, no-repeat);
  background-attachment: fixed, fixed;
}
body::after {
  content: "";
  position: fixed;
  inset: 0;
  z-index: -1;
  pointer-events: none;
  background: linear-gradient(to bottom, rgba(255, 255, 255, 0.18), rgba(255, 255, 255, 0));
  opacity: 0.34;
  mask-image: linear-gradient(to bottom, #000 0%, transparent 42%);
  -webkit-mask-image: linear-gradient(to bottom, #000 0%, transparent 42%);
}
body > main,
body > .footer-pro {
  position: relative;
  z-index: 1;
}
@media (max-width: 1060px) {
  body::before {
    background-attachment: scroll, scroll;
  }
}
@media (max-width: 900px) {
  body::before,
  body::after {
    position: absolute;
    min-height: 100%;
    background-attachment: scroll, scroll;
  }
}
@media (prefers-reduced-transparency: reduce) {
  body::before {
    background-image: linear-gradient(
      180deg,
      #f8fcff 0%,
      var(--color-surface-news, #f5fbff) 100%
    );
  }
  body::after {
    display: none;
  }
}

main {
  flex: 1 0 auto;
}
::selection {
  background: rgba(0, 105, 180, 0.18);
  color: var(--ink);
}
img { max-width: 100%; height: auto; display: block; }
a { color: var(--accent); text-decoration-thickness: 1px; text-underline-offset: 0.18em; }
a:hover { color: var(--accent-hover); }
:focus-visible {
  outline: 3px solid var(--gold-soft);
  outline-offset: 2px;
}
button,
a,
summary,
.tools-bar__view-btn,
.story-tts,
.nav-toggle,
.events-menu-toggle {
  touch-action: manipulation;
}

.skip-link {
  position: absolute;
  left: 1rem;
  top: -3rem;
  z-index: 100;
  padding: 0.65rem 0.9rem;
  border-radius: var(--radius-pill);
  background: var(--nav-blue);
  color: #fff;
  font-family: var(--font-ui);
  font-weight: 600;
  text-decoration: none;
}
.skip-link:focus { top: 0.75rem; }

.site-header {
  position: sticky;
  top: 0;
  z-index: 40;
  min-height: var(--header-h);
  overflow: visible;
  background: linear-gradient(180deg, var(--blue-600) 0%, var(--nav-blue) 55%, var(--nav-blue-deep) 100%);
  border-bottom: 3px solid var(--blue-900);
  color: #fff;
}
.site-header__inner {
  position: relative;
  width: 100%;
  max-width: none;
  margin: 0;
  padding: 0.5rem 0.45rem 0.5rem 0.35rem;
  box-sizing: border-box;
  display: flex;
  align-items: center;
  justify-content: flex-start;
  flex-wrap: nowrap;
  gap: 0.25rem;
  overflow: visible;
}
.site-header__actions {
  position: relative;
  z-index: 50;
  margin-left: auto;
  display: inline-flex;
  align-items: center;
  gap: 0.35rem;
  flex: 0 0 auto;
  overflow: visible;
}
.lang-switcher {
  position: relative;
  z-index: 50;
  font-family: var(--font-ui, "Source Sans 3", sans-serif);
}
.site-header:has(.lang-switcher.is-open) {
  z-index: 200;
}
.lang-switcher__toggle {
  display: inline-flex;
  align-items: center;
  gap: 0.4rem;
  min-height: 2.15rem;
  padding: 0.22rem 0.55rem 0.22rem 0.38rem;
  border: 1px solid rgba(255, 255, 255, 0.28);
  border-radius: 999px;
  background: rgba(255, 255, 255, 0.16);
  color: #fff;
  font: inherit;
  font-size: 0.78rem;
  font-weight: 700;
  letter-spacing: 0.01em;
  cursor: pointer;
  touch-action: manipulation;
  -webkit-tap-highlight-color: rgba(255, 255, 255, 0.18);
}
.lang-switcher__toggle:hover,
.lang-switcher__toggle[aria-expanded="true"] {
  background: rgba(255, 255, 255, 0.26);
}
.lang-switcher__flag,
.lang-switcher__flag-slot {
  display: block;
  width: 1.25rem;
  height: 0.86rem;
  border-radius: 2px;
  object-fit: cover;
  box-shadow: 0 0 0 1px rgba(8, 38, 59, 0.18);
  background: rgba(255, 255, 255, 0.35);
  flex: 0 0 auto;
}
.lang-switcher__name {
  max-width: 7.5rem;
  overflow: hidden;
  text-overflow: ellipsis;
  white-space: nowrap;
}
.lang-switcher__caret {
  width: 0.42rem;
  height: 0.42rem;
  margin-left: 0.05rem;
  border-right: 1.7px solid currentColor;
  border-bottom: 1.7px solid currentColor;
  transform: rotate(45deg) translateY(-1px);
  opacity: 0.85;
}
.lang-switcher.is-open .lang-switcher__caret {
  transform: rotate(225deg) translateY(-1px);
}
.lang-switcher__menu {
  display: none;
  position: fixed;
  inset: unset;
  top: calc(var(--header-h, 3.5rem) + 0.2rem);
  left: auto;
  right: 0.5rem;
  bottom: auto;
  transform: none;
  z-index: 400;
  width: max-content;
  min-width: 12.5rem;
  max-width: min(18rem, calc(100vw - 1rem));
  margin: 0;
  padding: 0.3rem;
  overflow: visible;
  border: 1px solid var(--line-16);
  border-radius: 14px;
  background: #fff;
  color: var(--ink);
  box-shadow: 0 16px 36px rgba(0, 78, 140, 0.18);
}
.lang-switcher.is-open .lang-switcher__menu,
.lang-switcher__menu:popover-open {
  display: block;
}
.lang-switcher__menu::before {
  content: "";
  position: absolute;
  left: 0;
  right: 0;
  bottom: 100%;
  height: 0.45rem;
}
.lang-switcher__option {
  display: flex;
  align-items: center;
  gap: 0.55rem;
  width: 100%;
  padding: 0.42rem 0.55rem;
  border: 0;
  border-radius: 10px;
  background: transparent;
  color: var(--ink);
  font: inherit;
  font-size: 0.86rem;
  font-weight: 650;
  text-align: left;
  text-decoration: none;
  cursor: pointer;
  touch-action: manipulation;
}
.lang-switcher__option:hover,
.lang-switcher__option:focus-visible {
  background: var(--nav-blue-soft);
  color: var(--nav-blue-deep);
}
.lang-switcher__option[aria-selected="true"] {
  background: linear-gradient(180deg, rgba(255, 240, 191, 0.55), rgba(223, 242, 255, 0.7));
  color: var(--blue-900);
}
.lang-switcher__option[aria-disabled="true"] {
  opacity: 0.42;
  cursor: not-allowed;
  color: var(--ink-soft);
}
.lang-switcher__option[aria-disabled="true"]:hover,
.lang-switcher__option[aria-disabled="true"]:focus-visible {
  background: transparent;
  color: var(--ink-soft);
}
.lang-switcher__option .lang-switcher__flag {
  box-shadow: 0 0 0 1px rgba(0, 105, 180, 0.16);
}
@media (pointer: coarse) {
  .lang-switcher__toggle {
    min-height: 44px;
    min-width: 44px;
  }
  .lang-switcher__option {
    min-height: 44px;
  }
}
.primary-nav {
  position: absolute;
  left: 50%;
  top: 50%;
  transform: translate(-50%, -50%);
  z-index: 1;
  display: flex;
  align-items: center;
  flex-wrap: nowrap;
  gap: 0.04rem;
  min-width: 0;
  flex: 0 0 auto;
  justify-content: center;
  overflow: visible;
}
.primary-nav__link {
  display: inline-flex;
  align-items: center;
  gap: 0.22rem;
  font-family: var(--font-ui);
  font-weight: 600;
  font-size: 0.78rem;
  line-height: 1.15;
  padding: 0.24rem 0.36rem 0.24rem 0.22rem;
  border: 1px solid transparent;
  border-radius: var(--radius-pill);
  color: #fff;
  text-decoration: none;
  white-space: nowrap;
  flex: 0 0 auto;
  transition: background 160ms ease, border-color 160ms ease, opacity 160ms ease;
}
.primary-nav__link .menu-icon,
.nav-dropdown > summary .menu-icon,
.nav-dropdown > .nav-dropdown__summary .menu-icon {
  width: 1.15rem;
  height: 1.15rem;
  border-radius: 0.35rem;
  background:
    linear-gradient(160deg, rgba(255, 255, 255, 0.92) 0%, rgba(255, 255, 255, 0.35) 42%, transparent 43%),
    linear-gradient(145deg, var(--icon-from), var(--icon-to));
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.55) inset,
    0 -1px 2px rgba(0, 0, 0, 0.14) inset,
    0 2px 6px color-mix(in srgb, var(--icon-glow) 50%, transparent);
  transform: none;
}
.primary-nav__link .menu-icon__svg,
.nav-dropdown > summary .menu-icon__svg,
.nav-dropdown > .nav-dropdown__summary .menu-icon__svg {
  width: 9px;
  height: 9px;
  stroke: var(--surface);
  filter: drop-shadow(0 1px 0 rgba(255, 255, 255, 0.3)) drop-shadow(0 1px 1px rgba(0, 0, 0, 0.2));
}
.primary-nav__link:hover {
  background: rgba(255, 255, 255, 0.16);
  border-color: rgba(255, 255, 255, 0.5);
}
.primary-nav__link[aria-disabled="true"] {
  cursor: default;
  opacity: 0.92;
}
.nav-toggle {
  display: none;
  width: 2.5rem;
  height: 2.5rem;
  padding: 0;
  border: 1px solid rgba(255, 255, 255, 0.5);
  border-radius: 0.7rem;
  background: rgba(255, 255, 255, 0.1);
  color: #fff;
  cursor: pointer;
  align-items: center;
  justify-content: center;
  flex-direction: column;
  gap: 0.28rem;
  transition: background 160ms ease, border-color 160ms ease;
}
.nav-toggle span {
  display: block;
  width: 1.15rem;
  height: 2px;
  border-radius: var(--radius-pill);
  background: currentColor;
  transition: transform 180ms ease, opacity 180ms ease;
}
.nav-toggle:hover,
.nav-toggle[aria-expanded="true"] {
  background: rgba(255, 255, 255, 0.2);
  border-color: rgba(255, 255, 255, 0.8);
}
.site-header.is-nav-open .nav-toggle span:nth-child(1) {
  transform: translateY(0.34rem) rotate(45deg);
}
.site-header.is-nav-open .nav-toggle span:nth-child(2) {
  opacity: 0;
}
.site-header.is-nav-open .nav-toggle span:nth-child(3) {
  transform: translateY(-0.34rem) rotate(-45deg);
}
.global-search-toggle {
  display: inline-flex;
  align-items: center;
  justify-content: flex-start;
  gap: 0.3rem;
  min-width: 0;
  width: auto;
  max-width: 9.5rem;
  min-height: 2.15rem;
  padding: 0.28rem 0.55rem;
  border: 1px solid rgba(255, 255, 255, 0.55);
  border-radius: var(--radius-pill);
  background: rgba(255, 255, 255, 0.14);
  color: #fff;
  font-family: var(--font-ui);
  font-size: 0.8rem;
  font-weight: 650;
  cursor: pointer;
  box-shadow: 0 4px 14px rgba(0, 45, 82, 0.16);
  transition: background 160ms ease, border-color 160ms ease, box-shadow 160ms ease;
}
.global-search-toggle:hover,
.global-search-toggle[aria-expanded="true"] {
  background: rgba(255, 255, 255, 0.22);
  border-color: rgba(255, 255, 255, 0.85);
  box-shadow: 0 6px 18px rgba(0, 45, 82, 0.2);
}
.global-search-toggle svg {
  display: block;
  flex: 0 0 auto;
  width: 1.05rem;
  height: 1.05rem;
}
.global-search-toggle__label {
  opacity: 0.92;
  white-space: nowrap;
  margin-right: 0.15rem;
}
.global-search-toggle__kbd {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-height: 1.2rem;
  padding: 0.08rem 0.35rem;
  border: 1px solid rgba(255, 255, 255, 0.45);
  border-radius: 0.35rem;
  background: rgba(8, 38, 59, 0.22);
  color: rgba(255, 255, 255, 0.92);
  font-family: var(--font-ui);
  font-size: 0.72rem;
  font-weight: 700;
  letter-spacing: 0.02em;
  line-height: 1;
  white-space: nowrap;
}
.global-search[hidden] { display: none !important; }
.global-search {
  position: fixed;
  inset: 0;
  z-index: 80;
  display: grid;
  place-items: start center;
  padding: calc(var(--header-h) + 1rem) 1rem 1.5rem;
}
.global-search__backdrop {
  position: absolute;
  inset: 0;
  border: 0;
  padding: 0;
  margin: 0;
  background: rgba(8, 38, 59, 0.45);
  cursor: pointer;
}
.global-search__panel {
  position: relative;
  z-index: 1;
  width: min(36rem, calc(100vw - 2rem));
  max-height: min(72vh, 34rem);
  display: flex;
  flex-direction: column;
  background: rgba(255, 255, 255, 0.98);
  border: 1px solid rgba(0, 105, 180, 0.18);
  border-radius: 18px;
  box-shadow: 0 24px 60px rgba(0, 45, 82, 0.28);
  overflow: hidden;
  color: var(--ink);
}
.global-search__head {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 0.75rem;
  padding: 0.9rem 1.05rem;
  background: linear-gradient(135deg, var(--blue-900) 0%, var(--nav-blue) 58%, var(--blue-400) 100%);
  border-bottom: 2px solid var(--gold-bright);
  color: #fff;
}
.global-search__head .global-search__title {
  margin: 0;
  font-family: var(--font-display);
  font-size: 1.05rem;
  font-weight: 800;
}
.global-search__close {
  width: 2.2rem;
  height: 2.2rem;
  border: 1px solid rgba(255, 255, 255, 0.35);
  border-radius: var(--radius-pill);
  background: rgba(255, 255, 255, 0.12);
  color: #fff;
  font-size: 1.4rem;
  line-height: 1;
  cursor: pointer;
}
.global-search__field {
  display: block;
  padding: 0.9rem 1.05rem 0.35rem;
}
.global-search__field input {
  width: 100%;
  min-height: 2.75rem;
  padding: 0.55rem 0.95rem;
  border: 1px solid rgba(0, 105, 180, 0.28);
  border-radius: var(--radius-pill);
  background: var(--panel-blue);
  color: var(--ink);
  font: inherit;
  font-size: 1rem;
}
.global-search__field input:focus {
  outline: 3px solid var(--ring);
  outline-offset: 1px;
}
.global-search__status {
  margin: 0;
  padding: 0.25rem 1.15rem 0.55rem;
  font-family: var(--font-ui);
  font-size: 0.85rem;
  font-weight: 600;
  color: var(--ink-soft);
}
.global-search__results {
  padding: 0 0.65rem 0.85rem;
  overflow: auto;
  display: grid;
  gap: 0.35rem;
}
.global-search__item {
  display: block;
  padding: 0.7rem 0.8rem;
  border-radius: 12px;
  text-decoration: none;
  color: var(--ink);
  border: 1px solid transparent;
}
.global-search__item:hover,
.global-search__item:focus-visible {
  background: var(--surface-muted);
  border-color: rgba(0, 105, 180, 0.18);
}
.global-search__item-title {
  display: block;
  font-family: var(--font-ui);
  font-weight: 700;
  font-size: 0.98rem;
  margin-bottom: 0.15rem;
}
.global-search__item-meta {
  display: block;
  font-family: var(--font-ui);
  font-size: 0.82rem;
  color: var(--ink-soft);
}
body.global-search-open { overflow: hidden; }
.brand {
  position: relative;
  z-index: 2;
  display: inline-flex;
  align-items: center;
  gap: 0.32rem;
  flex: 0 0 auto;
  flex-grow: 0;
  flex-shrink: 0;
  margin: 0;
  min-width: 0;
  max-width: none;
  width: auto;
  font-family: var(--font-display);
  font-weight: 700;
  font-size: clamp(1.05rem, 1.35vw, 1.28rem);
  color: #fff;
  text-decoration: none;
  letter-spacing: -0.01em;
  white-space: nowrap;
}
.brand__logo {
  width: 40px;
  height: 40px;
  object-fit: contain;
  flex: 0 0 auto;
}
.brand__name {
  color: #fff;
  white-space: nowrap;
}
.brand:hover { color: #fff; opacity: 0.95; }

.breadcrumbs {
  position: sticky;
  top: var(--header-h);
  z-index: 35;
  background: rgba(238, 248, 255, 0.96);
  backdrop-filter: blur(8px);
  border-bottom: 1px solid var(--line);
}
.breadcrumbs__inner {
  max-width: var(--max-wide);
  margin: 0 auto;
  padding: 0.7rem 1.25rem;
}
.breadcrumbs__list {
  list-style: none;
  margin: 0;
  padding: 0;
  display: flex;
  flex-wrap: wrap;
  align-items: center;
  gap: 0.35rem 0.15rem;
  font-family: var(--font-ui);
  font-size: 0.9rem;
}
.breadcrumbs__item {
  display: inline-flex;
  align-items: center;
  color: var(--ink-soft);
  min-width: 0;
}
.breadcrumbs__item:not(:last-child)::after {
  content: ">";
  margin: 0 0.45rem;
  color: rgba(0, 78, 140, 0.35);
}
.breadcrumbs__item a {
  color: var(--nav-blue);
  text-decoration: none;
  font-weight: 600;
}
.breadcrumbs__item a:hover {
  text-decoration: underline;
  color: var(--nav-blue-deep);
}
.breadcrumbs__item[aria-current="page"] span {
  color: var(--ink);
  font-weight: 600;
  overflow-wrap: anywhere;
}

.nav-dropdown { position: relative; font-family: var(--font-ui); }
.nav-dropdown > summary,
.nav-dropdown > .nav-dropdown__summary {
  list-style: none;
  cursor: pointer;
  display: inline-flex;
  align-items: center;
  gap: 0.22rem;
  font-weight: 600;
  font-size: 0.78rem;
  line-height: 1.15;
  padding: 0.24rem 0.42rem 0.24rem 0.22rem;
  border: 1px solid transparent;
  border-radius: var(--radius-pill);
  background: transparent;
  color: #fff;
  white-space: nowrap;
  flex: 0 0 auto;
  transition: background 160ms ease, border-color 160ms ease;
}
.nav-dropdown > summary:hover,
.nav-dropdown > .nav-dropdown__summary:hover,
.nav-dropdown[open] > summary,
.nav-dropdown[open] > .nav-dropdown__summary,
.nav-dropdown.is-hover-open > summary,
.nav-dropdown.is-hover-open > .nav-dropdown__summary {
  background: rgba(255, 255, 255, 0.16);
  border-color: rgba(255, 255, 255, 0.5);
}
.nav-dropdown > summary::-webkit-details-marker,
.nav-dropdown > .nav-dropdown__summary::-webkit-details-marker { display: none; }
.nav-dropdown > summary::after,
.nav-dropdown > .nav-dropdown__summary::after {
  content: none;
  display: none;
}

/* Forumlar-style literature / arts / science panels */
.nav-dropdown--literature > .nav-dropdown-panel,
.nav-dropdown--arts > .nav-dropdown-panel,
.nav-dropdown--science > .nav-dropdown-panel {
  position: absolute;
  left: 0;
  top: calc(100% + 0.35rem);
  z-index: 60;
  display: none;
  flex-direction: column;
  gap: 4px;
  min-width: 260px;
  max-width: min(320px, 90vw);
  margin: 0;
  padding: 6px;
  overflow: visible;
  background: #fff;
  border: 1px solid rgba(0, 105, 180, 0.16);
  border-radius: 12px;
  box-shadow: 0 16px 40px rgba(0, 45, 82, 0.22);
  color: var(--ink);
  animation: navDropdownIn 180ms ease-out;
}
.nav-dropdown--literature[open] > .nav-dropdown-panel,
.nav-dropdown--literature.is-hover-open > .nav-dropdown-panel,
.nav-dropdown--arts[open] > .nav-dropdown-panel,
.nav-dropdown--arts.is-hover-open > .nav-dropdown-panel,
.nav-dropdown--science[open] > .nav-dropdown-panel,
.nav-dropdown--science.is-hover-open > .nav-dropdown-panel {
  display: flex;
}
.nav-dropdown--literature > .nav-dropdown-panel::before,
.nav-dropdown--arts > .nav-dropdown-panel::before,
.nav-dropdown--science > .nav-dropdown-panel::before {
  content: "";
  position: absolute;
  top: -10px;
  left: 0;
  right: 0;
  height: 10px;
}
.nav-mega-links--science {
  display: flex;
  flex-direction: column;
  gap: 2px;
}
.nav-dropdown-panel--science {
  min-width: 17rem;
}
.nav-dropdown--nested {
  position: relative;
  width: 100%;
}
.nav-dropdown-toggle {
  width: 100%;
  margin: 0;
  padding: 8px 28px 8px 8px;
  border: 1px solid transparent;
  border-radius: 8px;
  background: transparent;
  color: #1e3d54;
  font-family: var(--font-ui);
  text-align: left;
  cursor: pointer;
  display: flex;
  flex-direction: row;
  align-items: center;
  gap: 0.55rem;
  position: relative;
  transition: background var(--duration-fast) var(--ease-standard), color var(--duration-fast) var(--ease-standard), border-color var(--duration-fast) var(--ease-standard);
}
.nav-dropdown-toggle__copy {
  display: flex;
  flex-direction: column;
  align-items: flex-start;
  gap: 2px;
  min-width: 0;
  padding-right: 0.2rem;
}
.nav-dropdown-toggle:hover,
.nav-dropdown-toggle:focus-visible,
.nav-dropdown--nested.is-mega-open > .nav-dropdown-toggle {
  color: var(--nav-blue-deep);
  background: rgba(0, 105, 180, 0.08);
  border-color: var(--line-14);
}
.nav-dropdown-link-title {
  font-size: 0.9rem;
  font-weight: 700;
  line-height: 1.25;
  color: inherit;
}
.nav-dropdown-link-desc {
  font-size: 0.75rem;
  font-weight: 400;
  line-height: 1.35;
  color: #5a7c93;
}
.nav-dropdown-toggle:hover .nav-dropdown-link-desc,
.nav-dropdown-toggle:focus-visible .nav-dropdown-link-desc {
  color: #2c5470;
}
.nav-dropdown-caret {
  position: absolute;
  top: 50%;
  right: 10px;
  width: 0.4rem;
  height: 0.4rem;
  border-right: 2px solid currentColor;
  border-bottom: 2px solid currentColor;
  transform: translateY(-60%) rotate(-45deg);
  transition: transform 160ms ease;
  opacity: 0.75;
}
.nav-dropdown--nested.is-mega-open > .nav-dropdown-toggle .nav-dropdown-caret {
  transform: translateY(-40%) rotate(45deg);
}
.nav-mega-links--arts {
  display: flex;
  flex-direction: column;
  gap: 2px;
  padding-left: 0;
  border-left: 0;
  min-width: 14rem;
}
.nav-dropdown-panel--arts {
  min-width: 15rem;
}
.nav-dropdown-panel--natural,
.nav-dropdown-panel--social,
.nav-dropdown-panel--humanities,
.nav-dropdown-panel--medical,
.nav-dropdown-panel--informatics,
.nav-dropdown-panel--math {
  min-width: 0;
}
.nav-dropdown-link[aria-disabled="true"] {
  cursor: default;
  opacity: 0.95;
}

.nav-dropdown-panel--mega {
  display: none;
  flex-direction: column;
  align-items: flex-start;
  margin: 0;
  padding: 8px;
  background: #fff;
  border: 1px solid rgba(0, 105, 180, 0.16);
  border-radius: 12px;
  box-shadow: 0 16px 40px rgba(0, 45, 82, 0.22);
}
.nav-dropdown--nested.is-mega-open > .nav-dropdown-panel--mega {
  display: flex;
}
.nav-mega-grid {
  display: grid;
  grid-template-columns: 1fr;
  gap: 6px 10px;
  width: 100%;
}
.nav-mega-col {
  display: flex;
  flex-direction: column;
  gap: 3px;
  min-width: 0;
}
.nav-mega-links {
  display: flex;
  flex-direction: column;
  gap: 2px;
  padding-left: 10px;
  border-left: 2px solid rgba(0, 105, 180, 0.22);
}
.nav-dropdown-link {
  display: flex;
  flex-direction: row;
  align-items: flex-start;
  gap: 0.45rem;
  padding: 6px 7px;
  border: 1px solid transparent;
  border-radius: 8px;
  text-decoration: none;
  color: var(--ink);
  font-family: var(--font-ui);
  transition: background var(--duration-fast) var(--ease-standard), color var(--duration-fast) var(--ease-standard), border-color var(--duration-fast) var(--ease-standard);
}
.nav-dropdown-link-copy {
  display: flex;
  flex-direction: column;
  align-items: flex-start;
  gap: 2px;
  min-width: 0;
}
.nav-dropdown-link .nav-dropdown-link-title {
  font-size: 0.78rem;
  font-weight: 500;
}
.nav-dropdown-link .nav-dropdown-link-desc {
  font-size: 0.7rem;
  color: #5a7c93;
}
.nav-dropdown-link .menu-icon {
  width: 1.2rem;
  height: 1.2rem;
  border-radius: 0.35rem;
  margin-top: 0.05rem;
  background:
    linear-gradient(160deg, rgba(255, 255, 255, 0.92) 0%, rgba(255, 255, 255, 0.35) 42%, transparent 43%),
    linear-gradient(145deg, var(--icon-from), var(--icon-to));
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.55) inset,
    0 -1px 2px rgba(0, 0, 0, 0.14) inset,
    0 2px 6px color-mix(in srgb, var(--icon-glow) 50%, transparent);
  transform: none;
}
.nav-dropdown-link .menu-icon__svg {
  width: 9px;
  height: 9px;
  stroke: var(--surface);
  filter: drop-shadow(0 1px 0 rgba(255, 255, 255, 0.3)) drop-shadow(0 1px 1px rgba(0, 0, 0, 0.2));
}
.nav-dropdown-link:hover,
.nav-dropdown-link:focus-visible {
  color: var(--nav-blue-deep);
  background: rgba(0, 105, 180, 0.08);
  border-color: var(--line-14);
}
.nav-dropdown-link:hover .nav-dropdown-link-desc,
.nav-dropdown-link:focus-visible .nav-dropdown-link-desc {
  color: #2c5470;
}
.nav-dropdown-link.is-active {
  color: #fff;
  background: var(--nav-blue);
  border-color: var(--nav-blue);
}
.nav-dropdown-link.is-active .nav-dropdown-link-desc {
  color: rgba(255, 255, 255, 0.85);
}

@media (hover: hover) and (pointer: fine) and (min-width: 1401px) {
  .nav-dropdown--literature:hover > .nav-dropdown-panel,
  .nav-dropdown--literature:focus-within > .nav-dropdown-panel,
  .nav-dropdown--arts:hover > .nav-dropdown-panel,
  .nav-dropdown--arts:focus-within > .nav-dropdown-panel,
  .nav-dropdown--science:hover > .nav-dropdown-panel,
  .nav-dropdown--science:focus-within > .nav-dropdown-panel {
    display: flex;
  }
  .nav-dropdown--nested.nav-dropdown--has-mega > .nav-dropdown-panel--mega {
    position: absolute;
    top: 0;
    left: calc(100% - 2px);
    width: max-content;
    max-width: min(780px, calc(100vw - 24px));
    opacity: 0;
    visibility: hidden;
    pointer-events: none;
    transition: opacity 140ms ease, visibility 140ms ease;
    display: flex;
  }
  .nav-dropdown--nested.nav-dropdown--has-mega.is-mega-open > .nav-dropdown-panel--mega,
  .nav-dropdown--nested.nav-dropdown--has-mega:hover > .nav-dropdown-panel--mega,
  .nav-dropdown--nested.nav-dropdown--has-mega:focus-within > .nav-dropdown-panel--mega {
    opacity: 1;
    visibility: visible;
    pointer-events: auto;
  }
  .nav-dropdown--nested .nav-mega-grid {
    grid-template-columns: repeat(3, minmax(14rem, 1fr));
    width: max-content;
    min-width: min(720px, calc(100vw - 48px));
    max-width: min(760px, calc(100vw - 40px));
    gap: 8px 14px;
  }
  .nav-dropdown--nested .nav-mega-col {
    width: auto;
    min-width: 14rem;
    max-width: 16.5rem;
  }
  .nav-dropdown--nested .nav-dropdown-link {
    width: 100%;
    max-width: none;
  }
  .nav-dropdown--nested .nav-dropdown-link-title {
    white-space: normal;
  }
  .nav-dropdown--nested .nav-dropdown-link-desc {
    display: -webkit-box;
    -webkit-line-clamp: 2;
    -webkit-box-orient: vertical;
    overflow: hidden;
  }
  .nav-dropdown-panel--arts {
    min-width: 16rem;
  }
  .nav-mega-links--arts {
    min-width: 15rem;
  }
  .nav-dropdown--science .nav-dropdown-panel--mega {
    max-height: min(70vh, 32rem);
    overflow-y: auto;
    overscroll-behavior: contain;
  }
  /* Tibb elmləri: 3-col mega (5+5+3), slightly tighter than stories. */
  .nav-dropdown--nested .nav-mega-grid--medical {
    grid-template-columns: repeat(3, minmax(13.5rem, 1fr));
    min-width: min(640px, calc(100vw - 48px));
    max-width: min(700px, calc(100vw - 40px));
  }
  /* Humanitar elmlər: 3-col mega (3+3+2). */
  .nav-dropdown--nested .nav-mega-grid--humanities {
    grid-template-columns: repeat(3, minmax(13.75rem, 1fr));
    min-width: min(620px, calc(100vw - 48px));
    max-width: min(700px, calc(100vw - 40px));
  }
  /* İnformatika: 3-col mega (3+3+2). */
  .nav-dropdown--nested .nav-mega-grid--informatics {
    grid-template-columns: repeat(3, minmax(13.75rem, 1fr));
    min-width: min(620px, calc(100vw - 48px));
    max-width: min(700px, calc(100vw - 40px));
  }
  /* Riyaziyyat: 3-col mega (3+3+2). */
  .nav-dropdown--nested .nav-mega-grid--math {
    grid-template-columns: repeat(3, minmax(13.75rem, 1fr));
    min-width: min(620px, calc(100vw - 48px));
    max-width: min(700px, calc(100vw - 40px));
  }
  /* İctimai elmlər: 2-col mega (3+2). */
  .nav-dropdown--nested .nav-mega-grid--social {
    grid-template-columns: repeat(2, minmax(14rem, 1fr));
    min-width: min(420px, calc(100vw - 48px));
    max-width: min(480px, calc(100vw - 40px));
  }
  /* Təbiət elmləri: 2-col mega (3+3). */
  .nav-dropdown--nested .nav-mega-grid--natural {
    grid-template-columns: repeat(2, minmax(14rem, 1fr));
    min-width: min(420px, calc(100vw - 48px));
    max-width: min(500px, calc(100vw - 40px));
  }
}

@keyframes navDropdownIn {
  from { opacity: 0; transform: translateY(-4px); }
  to { opacity: 1; transform: translateY(0); }
}

.nav-dropdown .menu-icon {
  --icon-from: var(--nav-blue);
  --icon-to: var(--blue-400);
  --icon-glow: var(--blue-soft);
  flex: 0 0 auto;
  width: 1.2rem;
  height: 1.2rem;
  display: inline-grid;
  place-items: center;
  border-radius: 0.35rem;
  background:
    linear-gradient(160deg, rgba(255, 255, 255, 0.92) 0%, rgba(255, 255, 255, 0.35) 42%, transparent 43%),
    linear-gradient(145deg, var(--icon-from), var(--icon-to));
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.55) inset,
    0 -1px 2px rgba(0, 0, 0, 0.14) inset,
    0 2px 6px color-mix(in srgb, var(--icon-glow) 50%, transparent);
  transform: none;
}
.nav-dropdown .menu-icon__svg {
  display: block;
  width: 9px;
  height: 9px;
  filter: drop-shadow(0 1px 0 rgba(255, 255, 255, 0.3)) drop-shadow(0 1px 1px rgba(0, 0, 0, 0.2));
  stroke: var(--surface);
}
.nav-dropdown-toggle .menu-icon {
  width: 1.2rem;
  height: 1.2rem;
}
.nav-dropdown-toggle .menu-icon__svg {
  width: 9px;
  height: 9px;
  stroke: var(--surface);
}
.menu-icon {
  --icon-from: var(--nav-blue);
  --icon-to: var(--blue-400);
  --icon-glow: var(--blue-soft);
  flex: 0 0 auto;
  width: 2.15rem;
  height: 2.15rem;
  display: inline-grid;
  place-items: center;
  border-radius: 0.7rem;
  background:
    linear-gradient(160deg, rgba(255, 255, 255, 0.92) 0%, rgba(255, 255, 255, 0.35) 42%, transparent 43%),
    linear-gradient(145deg, var(--icon-from), var(--icon-to));
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.55) inset,
    0 -2px 4px rgba(0, 0, 0, 0.18) inset,
    0 4px 10px color-mix(in srgb, var(--icon-glow) 55%, transparent),
    0 2px 0 rgba(0, 0, 0, 0.12);
  transform: perspective(120px) rotateX(12deg) translateY(0);
  transition: transform 160ms ease, box-shadow 160ms ease;
}
.menu-icon__svg {
  display: block;
  filter: drop-shadow(0 1px 0 rgba(255, 255, 255, 0.35)) drop-shadow(0 2px 2px rgba(0, 0, 0, 0.25));
  stroke: var(--surface);
}

@keyframes drop-in {
  to { opacity: 1; transform: translateY(0); }
}
@keyframes rise {
  from { opacity: 0; transform: translateY(14px); }
  to { opacity: 1; transform: translateY(0); }
}

@keyframes intro-fade {
  from { opacity: 0; transform: translateY(18px); }
  to { opacity: 1; transform: translateY(0); }
}
@keyframes intro-shine {
  from { transform: translateX(-30%) rotate(12deg); opacity: 0; }
  to { transform: translateX(0) rotate(12deg); opacity: 1; }
}
.intro {
  position: relative;
  isolation: isolate;
  max-width: var(--max);
  margin: 0 auto;
  padding: 1.1rem 1.25rem 0.35rem;
  overflow: hidden;
}
.intro__atmosphere {
  position: absolute;
  top: 1.1rem;
  left: 1.25rem;
  right: 1.25rem;
  bottom: 0.35rem;
  width: auto;
  height: auto;
  transform: none;
  border-radius: 28px;
  pointer-events: none;
  z-index: 0;
  overflow: hidden;
  background:
    radial-gradient(ellipse 70% 80% at 12% 20%, rgba(78, 180, 238, 0.28), transparent 60%),
    radial-gradient(ellipse 55% 70% at 88% 10%, rgba(0, 105, 180, 0.18), transparent 58%),
    radial-gradient(ellipse 40% 50% at 70% 85%, rgba(201, 155, 59, 0.14), transparent 55%),
    linear-gradient(135deg, rgba(217, 240, 255, 0.75), rgba(245, 251, 255, 0.35) 55%, rgba(255, 255, 255, 0.15));
  border: 1px solid rgba(158, 214, 245, 0.55);
  box-shadow: 0 24px 60px rgba(0, 78, 140, 0.1);
}
.intro__atmosphere::after {
  content: "";
  position: absolute;
  inset: -20% auto auto -10%;
  width: 55%;
  height: 140%;
  background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.45), transparent);
  animation: intro-shine 1.1s ease both 180ms;
}
.intro__content {
  position: relative;
  z-index: 1;
  display: grid;
  grid-template-columns: minmax(0, 1fr) minmax(18rem, 1.15fr);
  align-items: center;
  gap: clamp(0.75rem, 2.5vw, 1.75rem);
  width: 100%;
  max-width: none;
  margin: 0;
  padding: 0.55rem 1.35rem 0.45rem 1.6rem;
  border: 1px solid var(--line-14);
  border-radius: 28px;
  background:
    linear-gradient(160deg, rgba(255, 255, 255, 0.92) 0%, rgba(238, 248, 255, 0.88) 100%);
  backdrop-filter: blur(10px);
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.85) inset,
    0 18px 40px rgba(0, 78, 140, 0.12);
  animation: intro-fade 620ms var(--ease-outro) both;
}
.intro__copy {
  min-width: 0;
  text-align: center;
}
.intro__visual {
  position: relative;
  display: grid;
  place-items: center;
  place-self: center;
  justify-self: center;
  align-self: center;
  margin: 0 auto;
  width: min(100%, 26rem);
  aspect-ratio: 3 / 2;
  overflow: hidden;
  border-radius: 50%;
  background: #f5fbff;
  box-shadow: 0 10px 20px rgba(0, 78, 140, 0.16);
  animation: intro-fade 900ms var(--ease-outro) both 160ms;
}
.intro__visual::after {
  content: "";
  position: absolute;
  inset: 0;
  background: rgba(245, 251, 255, 0.12);
  z-index: 0;
  pointer-events: none;
}
.intro__visual img {
  position: relative;
  z-index: 1;
  width: 100%;
  height: 100%;
  display: block;
  object-fit: cover;
  object-position: center;
  background: transparent;
}
.intro__brand {
  font-family: var(--font-display);
  font-size: clamp(2.6rem, 6vw, 4.2rem);
  font-weight: 800;
  margin: 0 0 0.45rem;
  letter-spacing: -0.035em;
  line-height: 1.02;
  color: var(--ink);
  animation: intro-fade 760ms var(--ease-outro) both 120ms;
}
.intro__brand span {
  background: linear-gradient(135deg, var(--nav-blue) 10%, var(--blue-400) 55%, #2e9fd4 100%);
  -webkit-background-clip: text;
  background-clip: text;
  color: transparent;
}
.intro__lead {
  margin: 0;
  max-width: none;
  width: 100%;
  color: var(--ink-soft);
  font-family: var(--font-ui);
  font-size: clamp(0.84rem, 1.15vw, 0.92rem);
  line-height: 1.5;
  text-align: justify;
  text-justify: inter-word;
  hyphens: auto;
  animation: intro-fade 820ms var(--ease-outro) both 180ms;
}
.intro__source {
  display: flex;
  flex-direction: column;
  align-items: flex-start;
  gap: 0.55rem;
  margin: 0.7rem 0 0;
  font-family: var(--font-display);
  font-size: clamp(0.72rem, 1.05vw, 0.84rem);
  font-style: italic;
  line-height: 1.4;
  text-align: left;
  color: var(--blue-900);
  animation: intro-fade 900ms var(--ease-outro) both 240ms;
}
.intro__source-ornament {
  display: block;
  width: min(11.5rem, 56%);
  height: 0.72rem;
  align-self: center;
  background:
    radial-gradient(circle at 50% 50%, var(--gold-bright) 0 0.17rem, var(--gold) 0.18rem 0.22rem, transparent 0.24rem),
    linear-gradient(90deg, transparent, rgba(201, 155, 59, 0.15), var(--gold) 42%, var(--gold) 58%, rgba(201, 155, 59, 0.15), transparent);
  background-size: 100% 100%, 100% 1px;
  background-position: center, center;
  background-repeat: no-repeat;
}
.intro__source strong,
.intro__source-text {
  display: inline-block;
  max-width: min(46rem, 94%);
  font-weight: 400;
  font-style: italic;
  letter-spacing: 0.008em;
  padding: 0.32rem 0.85rem;
  border-radius: 1rem;
  color: var(--blue-900);
  background: linear-gradient(180deg, rgba(255, 240, 191, 0.58), rgba(223, 242, 255, 0.42));
  border: 1px solid rgba(201, 155, 59, 0.4);
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.78) inset,
    0 6px 16px rgba(0, 78, 140, 0.08);
}
@media (max-width: 900px) {
  .intro__content {
    grid-template-columns: 1fr;
    padding: 0.85rem 1.2rem 0.75rem;
    gap: 0.45rem;
  }
  .intro__lead {
    text-align: start;
    text-justify: auto;
    hyphens: manual;
  }
  .intro__visual {
    margin: 0;
    justify-self: center;
    width: min(92%, 22rem);
  }
  .intro__atmosphere {
    top: 1.1rem;
    left: 1.25rem;
    right: 1.25rem;
    bottom: 0.35rem;
    height: auto;
  }
}

.section {
  max-width: var(--max);
  margin: 0 auto;
  padding: 2rem 1.25rem 4.5rem;
}
.page-home #kateqoriyalar.home-browser {
  scroll-margin-top: calc(var(--header-h) + var(--breadcrumb-h) + 0.75rem);
}
.section__head {
  max-width: 40rem;
  margin-bottom: 1.5rem;
}
.section__head h2 {
  font-family: var(--font-display);
  font-size: clamp(1.7rem, 3vw, 2.3rem);
  margin: 0 0 0.45rem;
  letter-spacing: -0.02em;
}
.section__head p {
  margin: 0;
  color: var(--ink-soft);
  font-family: var(--font-ui);
  font-size: 1rem;
}

.visually-hidden {
  position: absolute !important;
  width: 1px;
  height: 1px;
  padding: 0;
  margin: -1px;
  overflow: hidden;
  clip: rect(0, 0, 0, 0);
  white-space: nowrap;
  border: 0;
}

.tools-bar {
  display: flex;
  flex-wrap: nowrap;
  align-items: flex-end;
  gap: 0.65rem 0.75rem;
  margin: 0 0 1.25rem;
  padding: 0.75rem 0.9rem;
  border: 1px solid var(--line);
  border-radius: var(--radius);
  background: rgba(255, 255, 255, 0.96);
  backdrop-filter: blur(8px);
  box-shadow: 0 4px 16px rgba(0, 78, 140, 0.08);
  font-family: var(--font-ui);
}
/* Category title spans layout; left edge matches .story-nav.sidebar. */
.page-category .category-hero {
  margin: 0;
  padding-left: 0;
  padding-right: 0;
}
.tools-bar__search {
  flex: 1 1 12rem;
  min-width: 9rem;
  max-width: 18rem;
}
.tools-bar__search input {
  width: 100%;
  min-height: 2.35rem;
  padding: 0.45rem 0.85rem;
  border: 1px solid var(--line-strong);
  border-radius: var(--radius-pill);
  background: var(--surface);
  color: var(--ink);
  font: inherit;
  font-size: 0.92rem;
}
.tools-bar__search input:focus {
  outline: 3px solid var(--ring);
  outline-offset: 1px;
  border-color: rgba(0, 105, 180, 0.45);
}
.tools-bar__field {
  display: inline-flex;
  flex-direction: column;
  flex: 0 0 auto;
  align-items: center;
  gap: 0.28rem;
  color: var(--ink-soft);
  font-size: 0.9rem;
  font-weight: 600;
}
.tools-bar__label {
  display: block;
  margin: 0;
  padding-inline: 0.1rem;
  width: 100%;
  font-family: var(--font-ui);
  font-size: 0.72rem;
  font-weight: 700;
  letter-spacing: 0.02em;
  text-transform: uppercase;
  color: rgba(52, 95, 134, 0.78);
  line-height: 1.2;
  text-align: center;
  white-space: nowrap;
  -webkit-font-smoothing: antialiased;
  text-rendering: optimizeLegibility;
  font-feature-settings: "kern" 1, "liga" 1;
}
:is(.page-home, .page-category) .tools-bar {
  flex-wrap: nowrap;
  align-items: flex-start;
  justify-content: flex-start;
  gap: var(--space-2);
  margin: 0 0 var(--space-5);
  padding: 0.55rem 0.65rem 0.6rem;
  border: 1px solid var(--line-14);
  border-radius: 1.35rem;
  background:
    linear-gradient(180deg, rgba(255, 255, 255, 0.94), rgba(238, 247, 252, 0.88));
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.9) inset,
    0 12px 32px rgba(0, 78, 140, 0.1);
  backdrop-filter: blur(14px) saturate(1.15);
}
/* Category tools stay one row on desktop; stack only on small screens. */
.tools-bar.tools-bar--dense {
  flex-wrap: nowrap;
  justify-content: flex-start;
  gap: 0.4rem;
  padding: 0.5rem 0.55rem 0.55rem;
}
:is(.page-home, .page-category) .tools-bar > * {
  flex: 0 0 auto;
}
.tools-bar.tools-bar--dense > * {
  flex: 0 0 auto;
  min-width: 0;
}
.tools-bar.tools-bar--dense > .tools-bar__search {
  flex: 0 1 17rem;
  max-width: 21rem;
  min-width: 10rem;
}
.tools-bar.tools-bar--dense > .tools-bar__batch {
  flex: 0 1 auto;
}
:is(.page-home, .page-category) .tools-bar__search {
  display: flex;
  flex-direction: column;
  align-items: stretch;
  gap: 0.4rem;
  flex: 1 1 9rem;
  min-width: 7.5rem;
  max-width: 14rem;
  align-self: flex-start;
  /* Match sibling fields' label row so the input top-aligns with controls */
  margin-top: calc(0.72rem * 1.2 + 0.28rem);
}
.tools-bar.tools-bar--dense > .tools-bar__search {
  margin-top: calc(0.65rem * 1.15 + 0.18rem);
}
:is(.page-home, .page-category) .tools-bar__search-field {
  position: relative;
  display: block;
  width: 100%;
}
:is(.page-home, .page-category) .tools-bar__search-field::before {
  content: "";
  position: absolute;
  left: 0.8rem;
  top: 50%;
  width: 0.8rem;
  height: 0.8rem;
  border: 2px solid rgba(0, 105, 180, 0.45);
  border-radius: 50%;
  transform: translateY(-55%);
  pointer-events: none;
  box-shadow: 0.24rem 0.24rem 0 -0.11rem rgba(0, 105, 180, 0.45);
}
:is(.page-home, .page-category) .tools-bar__search--active input {
  border-color: rgba(0, 105, 180, 0.5);
  background: var(--panel-blue);
  box-shadow: 0 0 0 4px var(--line-14);
}
.tools-bar__search-chip {
  display: none;
  align-items: center;
  gap: 0.4rem;
  min-width: 0;
  max-width: 100%;
  min-height: 1.7rem;
  padding: 0.12rem 0.18rem 0.12rem 0.55rem;
  border: 1px solid var(--nav-blue);
  border-radius: var(--radius-pill);
  background: var(--panel-blue);
  color: var(--nav-blue-deep);
  font-family: var(--font-ui);
  font-size: 0.78rem;
  font-weight: 600;
  line-height: 1.2;
}
.tools-bar__search-chip:not([hidden]) {
  display: inline-flex;
}
.tools-bar__search-chip-dot {
  flex: 0 0 auto;
  width: 0.42rem;
  height: 0.42rem;
  border-radius: 50%;
  background: var(--nav-blue);
}
.tools-bar__search-chip-text {
  min-width: 0;
  overflow: hidden;
  text-overflow: ellipsis;
  white-space: nowrap;
}
.tools-bar__search-clear {
  flex: 0 0 auto;
  min-width: 1.7rem;
  min-height: 1.7rem;
  padding: 0.12rem;
  line-height: 1;
}
:is(.page-home, .page-category) .tools-bar__search input {
  min-height: 2.25rem;
  padding: 0.4rem 0.75rem 0.4rem 2.05rem;
  border: 1px solid rgba(0, 105, 180, 0.16);
  background: rgba(255, 255, 255, 0.92);
  box-shadow: 0 1px 2px rgba(0, 78, 140, 0.04) inset;
  transition: border-color var(--duration) var(--ease-standard), box-shadow var(--duration) var(--ease-standard), background var(--duration) var(--ease-standard);
  font-size: 0.88rem;
}
:is(.page-home, .page-category) .tools-bar__search input:hover {
  border-color: rgba(0, 105, 180, 0.28);
}
:is(.page-home, .page-category) .tools-bar__search input:focus {
  outline: none;
  border-color: rgba(0, 105, 180, 0.5);
  background: #fff;
  box-shadow: 0 0 0 4px var(--line-14);
}
.tools-bar.tools-bar--dense .tools-bar__search input {
  min-height: 2.05rem;
  padding: 0.32rem 0.6rem 0.32rem 1.85rem;
  font-size: 0.82rem;
}
.tools-bar.tools-bar--dense .tools-bar__search-field::before {
  left: 0.65rem;
  width: 0.7rem;
  height: 0.7rem;
}
:is(.page-home, .page-category) .tools-bar__views {
  padding: 0.15rem;
  border: 1px solid var(--line-14);
  background: rgba(229, 244, 251, 0.7);
  box-shadow: 0 1px 2px rgba(0, 78, 140, 0.05) inset;
}
.tools-bar.tools-bar--dense .tools-bar__views {
  padding: 0.1rem;
}
:is(.page-home, .page-category) .tools-bar__view-btn {
  min-height: 1.95rem;
  padding: 0.25rem 0.65rem;
  font-size: 0.82rem;
  color: var(--ink-soft);
  transition: background var(--duration) var(--ease-standard), color var(--duration) var(--ease-standard), box-shadow var(--duration) var(--ease-standard), transform var(--duration) var(--ease-standard);
}
.tools-bar.tools-bar--dense .tools-bar__view-btn {
  min-height: 1.85rem;
  padding: 0.2rem 0.45rem;
  font-size: 0.78rem;
}
:is(.page-home, .page-category) .tools-bar__view-btn:hover {
  color: var(--nav-blue-deep);
}
:is(.page-home, .page-category) .tools-bar__view-btn[aria-pressed="true"] {
  background: linear-gradient(135deg, var(--nav-blue), var(--blue-400));
  color: #fff;
  box-shadow: 0 4px 12px rgba(0, 90, 154, 0.22);
}
:is(.page-home, .page-category) .tools-bar__view-btn--icon {
  min-width: 2.1rem;
  padding-inline: 0.42rem;
}
:is(.page-home, .page-category) .tools-bar__glyph {
  display: block;
  flex: 0 0 auto;
}
:is(.page-home, .page-category) .tools-bar__batch {
  display: inline-flex;
  flex-direction: column;
  align-items: stretch;
  gap: 0.28rem;
  flex: 0 0 auto;
  min-width: 0;
}
.tools-bar.tools-bar--dense .tools-bar__batch {
  gap: 0.2rem;
}
:is(.page-home, .page-category) .tools-bar__batch > .tools-bar__label {
  text-align: center;
  padding-inline: 0.35rem;
  color: rgba(0, 78, 140, 0.72);
}
:is(.page-home, .page-category) .tools-bar__batch-controls {
  display: inline-flex;
  align-items: center;
  flex-wrap: nowrap;
  gap: 0.35rem;
  padding: 0.2rem;
  border: 1px solid rgba(0, 105, 180, 0.16);
  border-radius: var(--radius-pill);
  background:
    linear-gradient(180deg, rgba(255, 255, 255, 0.98), rgba(232, 245, 252, 0.92));
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.95) inset,
    0 4px 14px rgba(0, 78, 140, 0.08);
}
.tools-bar.tools-bar--dense .tools-bar__batch-controls {
  gap: 0.28rem;
  padding: 0.16rem;
}
:is(.page-home, .page-category) .tools-bar__stepper,
:is(.page-home, .page-category) .tools-bar__pager {
  display: inline-flex;
  align-items: center;
  flex: 0 0 auto;
  gap: 0.12rem;
  padding: 0.1rem;
  border-radius: var(--radius-pill);
  background: rgba(229, 244, 251, 0.72);
  border: 1px solid rgba(0, 105, 180, 0.1);
}
:is(.page-home, .page-category) .tools-bar__batch-count {
  display: inline-flex;
  align-items: center;
  flex: 0 0 auto;
  margin: 0;
}
.tools-bar.tools-bar--dense .tools-bar__batch-input {
  width: 2.7rem;
  min-width: 2.5rem;
  min-height: 1.9rem;
  padding: 0.28rem 0.15rem;
  font-size: 0.8rem;
}
.tools-bar.tools-bar--dense .tools-bar__label {
  font-size: 0.68rem;
  letter-spacing: 0.02em;
}
.tools-bar.tools-bar--dense .tools-bar__batch [data-home-batch] {
  white-space: nowrap;
}
:is(.page-home, .page-category) .tools-bar__batch-input {
  width: 2.85rem;
  min-width: 2.65rem;
  min-height: 2.1rem;
  padding: 0.35rem 0.2rem;
  border: 1px solid rgba(0, 105, 180, 0.18);
  border-radius: var(--radius-pill);
  background: #fff;
  color: var(--nav-blue-deep);
  box-shadow:
    0 1px 2px rgba(0, 78, 140, 0.05) inset,
    0 0 0 1px rgba(255, 255, 255, 0.8);
  font-family: var(--font-ui);
  font-size: 0.88rem;
  font-weight: 800;
  letter-spacing: 0.01em;
  text-align: center;
  -moz-appearance: textfield;
  transition: border-color var(--duration) var(--ease-standard), box-shadow var(--duration) var(--ease-standard), background var(--duration) var(--ease-standard), color var(--duration) var(--ease-standard);
}
:is(.page-home, .page-category) .tools-bar__batch-input::-webkit-outer-spin-button,
:is(.page-home, .page-category) .tools-bar__batch-input::-webkit-inner-spin-button {
  -webkit-appearance: none;
  margin: 0;
}
:is(.page-home, .page-category) .tools-bar__batch-input:hover {
  border-color: rgba(0, 105, 180, 0.34);
}
:is(.page-home, .page-category) .tools-bar__batch-input:focus {
  outline: none;
  border-color: rgba(0, 105, 180, 0.55);
  box-shadow: 0 0 0 3px var(--line-14);
}
:is(.page-home, .page-category) .tools-bar__stepper-btn,
:is(.page-home, .page-category) .tools-bar__pager-btn,
:is(.page-home, .page-category) .tools-bar__icon-btn {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-width: 1.95rem;
  min-height: 1.95rem;
  padding: 0;
  border: 1px solid transparent;
  border-radius: var(--radius-pill);
  background: transparent;
  color: var(--nav-blue-deep);
  font-family: var(--font-ui);
  font-size: 1rem;
  font-weight: 700;
  line-height: 1;
  cursor: pointer;
  box-shadow: none;
}
.tools-bar.tools-bar--dense .tools-bar__stepper-btn,
.tools-bar.tools-bar--dense .tools-bar__pager-btn,
.tools-bar.tools-bar--dense .tools-bar__icon-btn {
  min-width: 1.85rem;
  min-height: 1.85rem;
}
:is(.page-home, .page-category) .tools-bar__stepper-btn:hover:not(:disabled),
:is(.page-home, .page-category) .tools-bar__pager-btn:hover:not(:disabled),
:is(.page-home, .page-category) .tools-bar__icon-btn:hover:not(:disabled) {
  background: rgba(255, 255, 255, 0.92);
  border-color: rgba(0, 105, 180, 0.16);
  box-shadow: 0 2px 8px rgba(0, 78, 140, 0.08);
}
:is(.page-home, .page-category) .tools-bar__stepper-btn:active:not(:disabled),
:is(.page-home, .page-category) .tools-bar__pager-btn:active:not(:disabled),
:is(.page-home, .page-category) .tools-bar__icon-btn:active:not(:disabled) {
  transform: translateY(1px);
}
:is(.page-home, .page-category) .tools-bar__batch [data-home-batch] {
  white-space: nowrap;
}
:is(.page-home, .page-category) .tools-bar__batch [data-home-batch]:disabled,
:is(.page-home, .page-category) [data-tools-play-visible]:disabled {
  opacity: 0.4;
  cursor: not-allowed;
  box-shadow: none;
}
:is(.page-home, .page-category) .tools-bar__batch-all {
  min-width: 2.1rem;
  min-height: 1.95rem;
  padding: 0.25rem 0.42rem;
}
:is(.page-home, .page-category) .tools-bar__batch-all.is-active,
:is(.page-home, .page-category) .tools-bar__batch-all[aria-pressed="true"] {
  background: linear-gradient(135deg, var(--nav-blue) 0%, var(--blue-400) 100%);
  border-color: transparent;
  color: #fff;
  box-shadow: 0 4px 12px rgba(0, 90, 154, 0.2);
}
:is(.page-home, .page-category) .tools-bar__batch-all.is-active:hover:not(:disabled),
:is(.page-home, .page-category) .tools-bar__batch-all[aria-pressed="true"]:hover:not(:disabled) {
  background: linear-gradient(135deg, var(--nav-blue-deep) 0%, var(--nav-blue) 100%);
  border-color: transparent;
  color: #fff;
  box-shadow: 0 6px 16px rgba(0, 78, 140, 0.24);
}
:is(.page-home, .page-category) .tools-bar__batch-range {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-width: 6.4rem;
  padding: 0.15rem 0.45rem;
  font-family: var(--font-ui);
  font-size: 0.72rem;
  font-weight: 700;
  letter-spacing: 0.01em;
  color: rgba(0, 60, 110, 0.78);
  line-height: 1.2;
  white-space: nowrap;
}
:is(.page-home, .page-category) .tools-bar__batch-range[hidden] {
  display: none !important;
}

/* Tools bar responsive: tablet — wrap before horizontal clip. */
@media (max-width: 1180px) {
  :is(.page-home, .page-category) .tools-bar,
  .tools-bar.tools-bar--dense {
    flex-wrap: wrap;
    align-items: flex-start;
  }
  :is(.page-home, .page-category) .tools-bar__search,
  .tools-bar.tools-bar--dense > .tools-bar__search {
    flex: 1 1 100%;
    max-width: none;
    min-width: 0;
    width: 100%;
  }
  .tools-bar.tools-bar--dense .tools-bar__batch-controls,
  .tools-bar.tools-bar--dense .tools-bar__pager,
  .tools-bar.tools-bar--dense .tools-bar__stepper {
    flex-wrap: wrap;
  }
}

/* Tools bar responsive: mobile */
@media (max-width: 760px) {
  .tools-bar,
  .page-category .tools-bar {
    display: grid;
    grid-template-columns: repeat(2, minmax(0, 1fr));
    align-items: end;
    justify-items: stretch;
    gap: 0.7rem 0.55rem;
    padding: 0.75rem;
    border-radius: 1.1rem;
  }
  :is(.page-home, .page-category) .tools-bar {
    display: grid;
    grid-template-columns: repeat(2, minmax(0, 1fr));
    align-items: start;
    justify-items: stretch;
    gap: 0.85rem;
    padding: 0.75rem;
    border-radius: 1.1rem;
  }
  .tools-bar__search,
  :is(.page-home, .page-category) .tools-bar__search {
    display: grid;
    grid-template-columns: minmax(0, 1fr);
    grid-column: 1 / -1;
    max-width: none;
    margin-top: 0;
  }
  .tools-bar__search-chip {
    grid-column: 1 / -1;
    width: 100%;
  }
  .tools-bar__field,
  :is(.page-home, .page-category) .tools-bar__field {
    width: 100%;
    align-items: stretch;
  }
  .tools-bar__label,
  :is(.page-home, .page-category) .tools-bar__label {
    text-align: left;
  }
  :is(.page-home, .page-category) .tools-bar__batch {
    grid-column: 1 / -1;
  }
  :is(.page-home, .page-category) .tools-bar__batch-controls {
    width: 100%;
    flex-wrap: wrap;
    border-radius: 1.1rem;
    justify-content: flex-start;
  }
  :is(.page-home, .page-category) .tools-bar__stepper,
  :is(.page-home, .page-category) .tools-bar__pager {
    flex: 1 1 auto;
  }
  :is(.page-home, .page-category) .tools-bar__batch-input {
    flex: 0 0 3.2rem;
  }
  :is(.page-home, .page-category) .tools-bar__batch-range {
    min-width: 5.5rem;
  }
  .tools-bar__views,
  :is(.page-home, .page-category) .tools-bar__views {
    width: 100%;
    justify-content: stretch;
  }
  .tools-bar__view-btn,
  :is(.page-home, .page-category) .tools-bar__view-btn {
    flex: 1 1 0;
    min-height: 2.75rem;
  }
  :is(.page-home, .page-category) .tools-bar__batch [data-home-batch],
  :is(.page-home, .page-category) .tools-bar__batch-input {
    min-height: 2.75rem;
  }
}

@media (pointer: coarse) {
  .global-search-toggle {
    min-height: 2.75rem;
  }
  :is(.page-home, .page-category) .tools-bar__batch [data-home-batch],
  :is(.page-home, .page-category) .tools-bar__batch-input,
  :is(.page-home, .page-category) .tools-bar__view-btn {
    min-height: 2.75rem;
  }
}

@media (max-width: 480px) {
  .tools-bar,
  :is(.page-home, .page-category) .tools-bar {
    grid-template-columns: 1fr;
  }
  .tools-bar__images-toggle,
  .tools-bar__texts-toggle,
  .tools-bar__pager,
  .tools-bar__stepper {
    width: 100%;
    min-width: 0;
  }
}
.story.story--text-hidden .story__text {
  display: none;
}
.tools-empty {
  margin: 1rem 0 0;
  padding: 1rem 1.1rem;
  border: 1px dashed var(--line-strong);
  border-radius: var(--radius-sm);
  background: var(--surface-muted);
  color: var(--ink-soft);
  font-family: var(--font-ui);
  text-align: center;
}
.page-category .tools-empty {
  max-width: var(--max-wide);
  margin-left: auto;
  margin-right: auto;
  width: calc(100% - 2.5rem);
}
[hidden] { display: none !important; }
[data-view="cards"][hidden],
[data-view="list"][hidden] { display: none !important; }
[data-view="cards"]:not([hidden]),
[data-view="list"]:not([hidden]) { display: block; }

.tools-bar__views {
  display: inline-flex;
  align-items: stretch;
  padding: 0.2rem;
  border: 1px solid rgba(0, 105, 180, 0.2);
  border-radius: var(--radius-pill);
  background: var(--panel-blue);
  gap: 0.15rem;
}
.tools-bar__view-btn {
  min-height: 2.2rem;
  padding: 0.35rem 0.95rem;
  border: 0;
  border-radius: var(--radius-pill);
  background: transparent;
  color: var(--ink-soft);
  font: inherit;
  font-size: 0.9rem;
  font-weight: 700;
  cursor: pointer;
}
.tools-bar__view-btn[aria-pressed="true"] {
  background: linear-gradient(135deg, var(--nav-blue), var(--blue-400));
  color: #fff;
  box-shadow: 0 4px 12px rgba(0, 90, 154, 0.22);
}
.home-browser .home-stories-layout {
  margin-top: 0.25rem;
  /* Avoid double inset vs .section padding so list aligns with tools. */
  padding-inline: 0;
}
.home-stories-main {
  min-width: 0;
  display: flex;
  flex-direction: column;
  gap: 0;
  overflow-anchor: none;
}
.home-stories-main .story-list {
  overflow-anchor: none;
}

/* DAAB home page-card parity */
.cat-grid {
  display: grid;
  grid-template-columns: repeat(4, minmax(0, 1fr));
  gap: 14px;
  align-items: stretch;
}
@media (max-width: 1180px) {
  .cat-grid { grid-template-columns: repeat(3, minmax(0, 1fr)); }
}
@media (max-width: 920px) {
  .cat-grid { grid-template-columns: repeat(2, minmax(0, 1fr)); }
}
@media (max-width: 620px) {
  .cat-grid { grid-template-columns: 1fr; }
}
@media (max-width: 1480px) {
  .brand {
    font-size: clamp(0.98rem, 1.2vw, 1.18rem);
  }
  .brand__logo {
    width: 36px;
    height: 36px;
  }
  .global-search-toggle__kbd {
    display: none;
  }
  .global-search-toggle {
    max-width: 7.25rem;
  }
  .primary-nav__link {
    font-size: 0.73rem;
    padding: 0.2rem 0.3rem 0.2rem 0.18rem;
    gap: 0.18rem;
  }
  .nav-dropdown > summary,
  .nav-dropdown > .nav-dropdown__summary {
    font-size: 0.73rem;
    padding: 0.2rem 0.34rem 0.2rem 0.18rem;
    gap: 0.18rem;
  }
  .primary-nav__link .menu-icon,
  .nav-dropdown > summary .menu-icon,
  .nav-dropdown > .nav-dropdown__summary .menu-icon {
    width: 1.05rem;
    height: 1.05rem;
  }
}
@media (max-width: 1400px) {
  .site-header__inner {
    display: grid;
    grid-template-columns: 2.5rem minmax(0, 1fr) auto;
    align-items: center;
    column-gap: 0.65rem;
    flex-wrap: unset;
    width: 100%;
    max-width: none;
    margin: 0;
    padding: 0.55rem 0.75rem;
    gap: 0;
  }
  .nav-toggle {
    display: inline-flex;
    grid-column: 1;
    grid-row: 1;
    justify-self: start;
  }
  .brand {
    grid-column: 2;
    grid-row: 1;
    justify-self: center;
    max-width: 100%;
    font-size: clamp(1.1rem, 4.2vw, 1.35rem);
  }
  .brand__logo {
    width: 42px;
    height: 42px;
  }
  .brand__name {
    overflow: hidden;
    text-overflow: ellipsis;
    white-space: nowrap;
  }
  .site-header__actions {
    grid-column: 3;
    grid-row: 1;
    margin-left: 0;
    justify-self: end;
  }
  .lang-switcher__name {
    max-width: 5.25rem;
  }
  .global-search-toggle__label,
  .global-search-toggle__kbd { display: none; }
  .global-search-toggle {
    min-width: 2.75rem;
    width: 2.75rem;
    min-height: 2.75rem;
    padding: 0.45rem;
    justify-content: center;
    box-shadow: none;
  }
  .global-search-toggle svg {
    width: 1.2rem;
    height: 1.2rem;
  }
  .primary-nav {
    display: none;
    grid-column: 1 / -1;
    grid-row: 2;
    position: absolute;
    left: 0;
    right: 0;
    top: 100%;
    transform: none;
    z-index: 45;
    padding: 0.75rem 1.25rem 1rem;
    background: linear-gradient(180deg, #f5fbff 0%, #e8f5fc 100%);
    border-bottom: 1px solid rgba(0, 105, 180, 0.18);
    box-shadow: 0 16px 28px rgba(0, 45, 82, 0.16);
    flex-direction: column;
    align-items: stretch;
    flex-wrap: nowrap;
    gap: 0.35rem;
  }
  .site-header.is-nav-open .primary-nav {
    display: flex;
    max-height: calc(100dvh - var(--header-h));
    overflow-x: hidden;
    overflow-y: auto;
    overscroll-behavior: contain;
    -webkit-overflow-scrolling: touch;
    animation: drop-in 180ms ease both;
  }
  .primary-nav__link {
    color: var(--ink);
    font-size: 0.95rem;
    padding: 0.55rem 0.7rem;
    border: 1px solid var(--line-14);
    background: rgba(255, 255, 255, 0.72);
    white-space: normal;
    gap: 0.45rem;
  }
  .primary-nav__link:hover {
    background: var(--surface);
    border-color: rgba(0, 105, 180, 0.28);
    color: var(--nav-blue-deep);
  }
  .nav-dropdown { width: 100%; }
  .nav-dropdown > summary,
  .nav-dropdown > .nav-dropdown__summary {
    display: inline-flex;
    width: 100%;
    box-sizing: border-box;
    color: var(--ink);
    font-size: 0.95rem;
    font-weight: 600;
    line-height: 1.25;
    padding: 0.55rem 0.7rem;
    border: 1px solid var(--line-14);
    border-radius: 0.65rem;
    background: rgba(255, 255, 255, 0.72);
    white-space: normal;
    gap: 0.45rem;
    justify-content: flex-start;
    align-items: center;
  }
  .nav-dropdown > summary:hover,
  .nav-dropdown > .nav-dropdown__summary:hover,
  .nav-dropdown[open] > summary,
  .nav-dropdown[open] > .nav-dropdown__summary {
    background: #fff;
    border-color: rgba(0, 105, 180, 0.28);
    color: var(--nav-blue-deep);
  }
  .nav-dropdown > summary::after,
  .nav-dropdown > .nav-dropdown__summary::after {
    content: "";
    display: block;
    margin-left: auto;
    flex: 0 0 auto;
    width: 0.45rem;
    height: 0.45rem;
    border-right: 2px solid currentColor;
    border-bottom: 2px solid currentColor;
    transform: rotate(45deg);
    transition: transform 160ms ease;
  }
  .nav-dropdown[open] > summary::after,
  .nav-dropdown[open] > .nav-dropdown__summary::after {
    transform: translateY(0.12rem) rotate(225deg);
  }
  /* Keep panels collapsed until the top-level item is opened. */
  .nav-dropdown--literature > .nav-dropdown-panel,
  .nav-dropdown--arts > .nav-dropdown-panel,
  .nav-dropdown--science > .nav-dropdown-panel {
    position: static;
    left: auto;
    top: auto;
    min-width: 0;
    max-width: none;
    width: 100%;
    margin: 0.35rem 0 0;
    box-shadow: none;
    border: 1px solid var(--line-14);
    animation: none;
  }
  .nav-dropdown--nested .nav-dropdown-panel--mega {
    position: static;
    width: 100%;
    max-width: none;
    margin-top: 0.25rem;
    box-shadow: none;
  }
  .nav-dropdown--nested .nav-mega-grid {
    grid-template-columns: 1fr;
  }
  .nav-dropdown-caret {
    transform: translateY(-60%) rotate(45deg);
  }
  .nav-dropdown--nested.is-mega-open > .nav-dropdown-toggle .nav-dropdown-caret {
    transform: translateY(-40%) rotate(225deg);
  }
  body.nav-open {
    overflow: hidden;
    overscroll-behavior: none;
  }
}

.cat-card.page-card {
  position: relative;
  isolation: isolate;
  box-sizing: border-box;
  display: grid;
  grid-template-columns: 44px minmax(0, 1fr);
  grid-template-rows: auto minmax(0, 1fr) auto;
  column-gap: 12px;
  row-gap: 0;
  min-height: 0;
  height: 100%;
  padding: 16px 16px 12px;
  overflow: hidden;
  border-radius: var(--radius-card);
  border: 1px solid var(--blue-soft);
  background: rgba(245, 251, 255, 0.96);
  box-shadow: var(--shadow);
  text-decoration: none;
  color: inherit;
  transition: transform 0.25s ease, border-color 0.25s ease, box-shadow 0.25s ease;
}
.cat-card.page-card::before {
  content: "";
  position: absolute;
  inset: 0;
  z-index: -1;
  opacity: 0;
  background: linear-gradient(135deg, rgba(46, 159, 212, 0.12), transparent 52%, rgba(200, 155, 56, 0.13));
  transition: opacity 0.25s ease;
}
.cat-card.page-card::after {
  content: "";
  position: absolute;
  right: -50px;
  top: -50px;
  z-index: -1;
  width: 130px;
  height: 130px;
  background: rgba(46, 159, 212, 0.09);
  border-radius: 50%;
  transition: transform 0.25s ease;
}
.cat-card.page-card:hover {
  transform: translateY(-7px);
  border-color: var(--blue-400);
  box-shadow: 0 20px 45px rgba(0, 78, 140, 0.2);
}
.cat-card.page-card:hover::before { opacity: 1; }
.cat-card.page-card:hover::after { transform: scale(1.25); }
.cat-card .card-icon-wrap {
  grid-column: 1;
  grid-row: 1 / 4;
  align-self: start;
  display: grid;
  place-items: center;
  width: 44px;
  height: 44px;
  margin: 0;
  background: linear-gradient(135deg, #d9f0ff, #fff);
  border: 1px solid var(--blue-soft);
  border-radius: 18px;
  box-shadow: inset 0 1px 0 #fff, 0 10px 22px rgba(0, 105, 180, 0.16);
}
.cat-card .card-icon-wrap .menu-icon {
  width: 2rem;
  height: 2rem;
  border-radius: 0.55rem;
  transform: none;
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.55) inset,
    0 -1px 3px rgba(0, 0, 0, 0.16) inset,
    0 3px 8px color-mix(in srgb, var(--icon-glow) 50%, transparent);
}
.cat-card .card-icon-wrap .menu-icon__svg {
  width: 16px;
  height: 16px;
}
.cat-card .card-body {
  display: contents;
}
.cat-card .card-title {
  grid-column: 2;
  grid-row: 1;
  margin: 0 0 8px;
  color: var(--nav-blue-deep);
  font-family: var(--font-display);
  font-size: 0.8125rem;
  font-weight: 800;
  line-height: 1.28;
  letter-spacing: -0.01em;
}
.cat-card .card-desc {
  grid-column: 2;
  grid-row: 2;
  align-self: start;
  margin: 0;
  color: var(--ink-soft);
  font-family: var(--font-ui);
  font-size: 0.90625rem;
  line-height: 1.48;
  display: -webkit-box;
  -webkit-box-orient: vertical;
  -webkit-line-clamp: 3;
  overflow: hidden;
}
.cat-card__meta {
  grid-column: 2;
  grid-row: 3;
  justify-self: start;
  margin-top: 10px;
  display: inline-flex;
  align-items: center;
  min-height: 1.55rem;
  padding: 0.15rem 0.6rem;
  border-radius: var(--radius-pill);
  border: 1px solid rgba(0, 105, 180, 0.16);
  background: rgba(255, 255, 255, 0.85);
  font-family: var(--font-ui);
  font-size: 0.75rem;
  font-weight: 700;
  color: var(--nav-blue);
  white-space: nowrap;
}
.cat-card__meta--soon {
  color: var(--ink-soft);
  border-color: rgba(8, 38, 59, 0.14);
  background: rgba(255, 255, 255, 0.72);
}
.content-placeholder {
  margin: 0.35rem 0 1.5rem;
  padding: 2.1rem 1.6rem 2rem;
  border-radius: 18px;
  border: 1px dashed rgba(0, 105, 180, 0.28);
  background: linear-gradient(165deg, rgba(229, 244, 251, 0.94), rgba(255, 255, 255, 0.96));
  text-align: center;
  box-shadow: 0 10px 28px rgba(8, 38, 59, 0.06);
}
.content-placeholder__kicker {
  display: inline-flex;
  align-items: center;
  margin: 0 0 0.7rem;
  padding: 0.15rem 0.7rem;
  border-radius: var(--radius-pill);
  border: 1px solid rgba(0, 105, 180, 0.18);
  background: rgba(255, 255, 255, 0.88);
  font-family: var(--font-ui);
  font-size: 0.75rem;
  font-weight: 700;
  letter-spacing: 0.04em;
  text-transform: uppercase;
  color: var(--nav-blue);
}
.content-placeholder__title {
  margin: 0 0 0.55rem;
  font-size: clamp(1.25rem, 2.4vw, 1.7rem);
  font-weight: 700;
  color: var(--ink);
}
.content-placeholder__lead {
  margin: 0 auto;
  max-width: 40rem;
  color: var(--ink-soft);
  font-family: var(--font-ui);
  font-size: 0.98rem;
  line-height: 1.55;
}
.page-category .category-layout--placeholder {
  grid-template-columns: minmax(0, 1fr);
}
.page-category .category-layout--placeholder > .category-hero,
.page-category .category-layout--placeholder > .content-placeholder {
  grid-column: 1;
}

.category-page { padding: 1.25rem 0 3rem; }
.category-hero {
  width: 100%;
  margin: 0;
  padding: 0.15rem 0 0;
  background: transparent;
  animation: intro-fade 520ms var(--ease-outro) both;
}
.category-hero h1 {
  font-family: var(--font-display);
  font-size: clamp(1.55rem, 3.2vw, 2.15rem);
  margin: 0 0 0.35rem;
  letter-spacing: -0.03em;
  line-height: 1.15;
  color: var(--ink);
  background: linear-gradient(135deg, var(--ink) 0%, var(--nav-blue-deep) 55%, var(--nav-blue) 100%);
  -webkit-background-clip: text;
  background-clip: text;
  color: transparent;
  overflow-wrap: anywhere;
  word-break: break-word;
}
.category-hero__lead {
  margin: 0;
  color: var(--ink-soft);
  font-family: var(--font-ui);
  max-width: 58ch;
  font-size: 0.95rem;
  line-height: 1.4;
  overflow-wrap: anywhere;
  word-break: break-word;
}
.category-hero .intro__source {
  margin-top: 0.85rem;
  max-width: min(46rem, 100%);
}
.category-hero .intro__source-ornament {
  align-self: flex-start;
  width: min(11.5rem, 70%);
}
.category-hero .intro__source-text {
  max-width: 100%;
}

.category-layout {
  max-width: var(--max-wide);
  margin: 0 auto;
  padding: 0 1.25rem 2rem;
  display: grid;
  grid-template-columns: minmax(240px, 300px) minmax(0, 1fr);
  gap: 1.75rem;
  align-items: start;
  overflow: visible;
}
/* Title left col + tools right col, same top row; panels below. */
.page-category .category-layout {
  column-gap: 28px;
  row-gap: 1.25rem;
  align-items: start;
}
.page-category .category-layout > .category-hero {
  grid-column: 1;
  grid-row: 1;
  width: 100%;
  min-width: 0;
  margin: 0;
  padding: 0.15rem 0 0;
  box-sizing: border-box;
  align-self: start;
}
.tools-bar.tools-bar--dense {
  grid-column: 2;
  grid-row: 1;
  width: 100%;
  max-width: 100%;
  min-width: 0;
  margin: 0;
  box-sizing: border-box;
  align-self: start;
}
.page-category .category-layout > .story-nav.sidebar {
  grid-column: 1;
  grid-row: 2;
}
.page-category .category-layout > .story-list {
  grid-column: 2;
  grid-row: 2;
  min-width: 0;
}
/* DAAB News sidebar widget parity (activities.html) */
.story-nav.sidebar {
  position: sticky;
  top: calc(var(--sticky-stack-bottom) + 12px);
  align-self: start;
  width: 100%;
  max-height: calc(100vh - var(--sticky-stack-bottom) - 28px);
  overflow: visible;
  z-index: 20;
  display: flex;
  flex-direction: column;
  gap: 24px;
  padding: 0;
  border: 0;
  background: transparent;
  box-shadow: none;
}
.story-nav .sidebar-widget {
  display: flex;
  flex-direction: column;
  width: 100%;
  max-height: calc(100vh - var(--sticky-stack-bottom) - 28px);
  overflow: hidden;
  background: rgba(255, 255, 255, 0.96);
  border: 1px solid rgba(0, 105, 180, 0.18);
  border-radius: 14px;
  box-shadow: 0 14px 28px rgba(0, 45, 82, 0.10);
}
.story-nav .widget-head {
  flex: 0 0 auto;
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 12px;
  padding: 14px 20px;
  background: linear-gradient(135deg, var(--ink) 0%, var(--nav-blue-deep) 58%, #2e9fd4 100%);
  border-bottom: 2px solid var(--gold-bright);
  color: #fff;
  font-family: var(--font-display);
  font-size: 0.7875rem;
  font-weight: 800;
  letter-spacing: 0.01em;
}
.story-nav .widget-head > span {
  display: inline-flex;
  align-items: center;
  gap: 8px;
}
.story-nav .events-menu-toggle {
  display: none;
  flex-direction: column;
  justify-content: center;
  gap: 4px;
  width: 44px;
  height: 44px;
  padding: 6px 8px;
  background: rgba(255, 255, 255, 0.14);
  border: 1px solid rgba(255, 255, 255, 0.35);
  border-radius: 8px;
  cursor: pointer;
}
.story-nav .events-menu-toggle span {
  display: block;
  width: 100%;
  height: 2px;
  background: #fff;
  border-radius: var(--radius-pill);
}
.story-nav .widget-body {
  flex: 1 1 auto;
  min-height: 0;
  padding: 16px 20px;
  overflow-y: auto;
  overflow-x: hidden;
  overscroll-behavior: contain;
  -webkit-overflow-scrolling: touch;
  scrollbar-width: thin;
  scrollbar-color: #2e9fd4 #e8eef5;
}
.story-nav .widget-body::-webkit-scrollbar { width: 6px; }
.story-nav .widget-body::-webkit-scrollbar-track { background: #e8eef5; border-radius: 3px; }
.story-nav .widget-body::-webkit-scrollbar-thumb { background: #2e9fd4; border-radius: 3px; }
.story-nav .widget-body::-webkit-scrollbar-thumb:hover { background: #1a6fa8; }
.story-nav .timeline-list {
  list-style: none;
  margin: 0;
  padding: 0;
}
.story-nav .timeline-list li {
  display: flex;
  gap: 12px;
  padding: 6px 0;
  border-bottom: 1px solid rgba(0, 45, 82, 0.08);
  font-family: var(--font-ui);
  font-size: 0.9375rem;
  line-height: 1.3;
  color: #345d76;
  transition: color 0.18s ease;
}
.story-nav .timeline-list li:last-child { border-bottom: none; }
.story-nav .timeline-list li:hover { color: var(--nav-blue); }
.story-nav .timeline-list a {
  flex: 1 1 auto;
  min-width: 0;
  display: block;
  position: relative;
  margin: -2px -6px;
  padding: 2px 6px 2px 1.15em;
  color: inherit;
  background: transparent;
  border-radius: 10px;
  font-family: var(--font-ui);
  font-size: 0.9375rem;
  font-weight: 500;
  line-height: 1.3;
  text-decoration: none !important;
  border: 0;
  transition: background-color 0.15s ease, color 0.15s ease;
}
.story-nav .timeline-list a::before {
  content: "";
  position: absolute;
  left: 0.22em;
  top: calc(2px + 0.52em);
  width: 0.26em;
  height: 0.26em;
  border-radius: 50%;
  background: currentColor;
  opacity: 0.5;
  pointer-events: none;
}
.story-nav .timeline-list a:hover,
.story-nav .timeline-list a:focus-visible,
.story-nav .timeline-list a.is-active,
.story-nav .timeline-list a.tl-active {
  color: var(--nav-blue) !important;
  background: var(--surface-muted) !important;
  text-decoration: none !important;
}
.story-nav .timeline-list a.is-active,
.story-nav .timeline-list a.tl-active {
  font-weight: 700;
}

.story-list {
  min-width: 0;
  display: flex;
  flex-direction: column;
  gap: 0;
}
/* DAAB News feed card parity (activities.html) */
.story.news-card {
  scroll-margin-top: calc(var(--sticky-stack-bottom) + 28px);
  margin: 0 0 26px;
  background: rgba(255, 255, 255, 0.96);
  border: 1px solid rgba(0, 105, 180, 0.18);
  border-radius: var(--radius-card);
  box-shadow: var(--shadow);
  overflow: hidden;
  transition: transform 0.22s ease, box-shadow 0.22s ease, border-color 0.22s ease;
}
.story.news-card:hover {
  transform: translateY(-5px);
  box-shadow: 0 20px 45px rgba(0, 78, 140, 0.2);
  border-color: rgba(78, 180, 238, 0.42);
}
.story.news-card.story--playing {
  border-color: rgba(212, 160, 23, 0.72);
  box-shadow:
    0 0 0 2px rgba(240, 199, 94, 0.45),
    0 16px 36px rgba(0, 78, 140, 0.16);
}
.story .card-header {
  padding: 20px 28px 14px;
  background: linear-gradient(135deg, var(--blue-900) 0%, var(--nav-blue) 58%, var(--blue-400) 100%);
  border-bottom: 2px solid var(--gold-bright);
  text-align: center;
}
.story .card-title,
.story__title {
  margin: 0;
  font-family: var(--font-display);
  font-size: clamp(14px, 1.54vw, 19.6px);
  font-weight: 800;
  letter-spacing: -0.015em;
  line-height: 1.38;
  color: #fff;
  background: none;
  padding: 0;
  border-radius: 0;
  box-shadow: none;
}
.story .card-body {
  padding: 14px 28px 26px;
  line-height: 1.45;
}
.story .card-body > :first-child {
  margin-top: 0;
}
.story__panel {
  width: 100%;
  margin: 0;
  padding: 0.85rem 1.05rem 1.1rem;
  border-radius: 14px;
  overflow: hidden;
  border: 1px solid var(--line-14);
  background: #fff;
  box-shadow: 0 4px 16px rgba(0, 0, 0, 0.12);
  box-sizing: border-box;
}
.story__text,
.story .card-text {
  width: 100%;
  max-width: none;
  margin: 0;
  padding: 0;
  border: 0;
  background: transparent;
  box-shadow: none;
  border-radius: 0;
  color: var(--ink-soft);
  font-family: var(--font-body);
  font-size: 1.02rem;
  line-height: 1.55;
  text-align: justify;
  text-justify: inter-word;
  hyphens: auto;
  box-sizing: border-box;
  cursor: zoom-in;
}
.story__text:focus-visible,
.story .card-text:focus-visible {
  outline: 3px solid var(--ring);
  outline-offset: 3px;
  border-radius: 8px;
}
.story__text p,
.story .card-text p {
  margin: 0 0 0.75rem;
  font-size: inherit;
  line-height: inherit;
  text-align: justify;
  text-justify: inter-word;
  hyphens: auto;
}
.story__text p:last-child,
.story .card-text p:last-child {
  margin-bottom: 0;
}
.story__text .story__source,
.story .card-text .story__source {
  display: inline-block;
  max-width: min(46rem, 100%);
  margin: 0.85rem 0 0;
  padding: 0.32rem 0.85rem;
  border-radius: 1rem;
  color: var(--blue-900);
  font-family: var(--font-display);
  font-size: clamp(0.72rem, 1.05vw, 0.84rem);
  font-style: italic;
  font-weight: 400;
  letter-spacing: 0.008em;
  line-height: 1.4;
  text-align: left;
  text-justify: auto;
  hyphens: manual;
  background: linear-gradient(180deg, rgba(255, 240, 191, 0.58), rgba(223, 242, 255, 0.42));
  border: 1px solid rgba(201, 155, 59, 0.4);
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.78) inset,
    0 6px 16px rgba(0, 78, 140, 0.08);
}
.story__text .story__moral,
.story .card-text .story__moral {
  clear: both;
  margin: 1.65rem 0 0;
  padding: 0.85rem 1rem 0.9rem 1.05rem;
  background: linear-gradient(135deg, #fff8e4 0%, var(--gold-soft) 55%, #ffe9b8 100%);
  border: 1px solid rgba(201, 155, 59, 0.42);
  border-left: 4px solid var(--gold);
  border-radius: var(--radius-sm);
  color: var(--ink);
  font-family: var(--font-display);
  font-size: 1.02em;
  font-weight: 700;
  line-height: 1.45;
  text-align: left;
  text-justify: auto;
  hyphens: manual;
  box-shadow: 0 4px 14px rgba(201, 155, 59, 0.12);
}
@media (max-width: 760px) {
  .story__text .story__moral,
  .story .card-text .story__moral {
    margin-top: 1.35rem;
    padding: 0.75rem 0.85rem 0.8rem 0.9rem;
    font-size: 1.05em;
    line-height: 1.42;
  }
}
.story__figure {
  margin: 14px 0 0;
  border-radius: 14px;
  overflow: hidden;
  border: 1px solid var(--line-14);
  background: #fff;
  padding: 8px;
  box-shadow: 0 4px 16px rgba(0, 0, 0, 0.12);
  transition: opacity 160ms ease, max-height 220ms ease, margin 160ms ease, padding 160ms ease;
}
.story__figure-open {
  display: block;
  width: 100%;
  margin: 0;
  padding: 0;
  border: 0;
  background: transparent;
  cursor: zoom-in;
  border-radius: 10px;
}
.story__figure-open:focus-visible {
  outline: 3px solid var(--ring);
  outline-offset: 2px;
}
.story__figure img {
  width: 100%;
  height: auto;
  max-width: 100%;
  object-fit: contain;
  object-position: center;
  display: block;
  border-radius: 10px;
  pointer-events: none;
}
.story.story--figure-hidden .story__figure {
  display: none;
}

.illustration-lightbox {
  position: fixed;
  inset: 0;
  z-index: 12000;
  display: flex;
  align-items: center;
  justify-content: center;
  padding: max(0.75rem, env(safe-area-inset-top)) max(0.75rem, env(safe-area-inset-right)) max(0.75rem, env(safe-area-inset-bottom)) max(0.75rem, env(safe-area-inset-left));
  background: rgba(8, 28, 48, 0.78);
  backdrop-filter: blur(6px);
  -webkit-backdrop-filter: blur(6px);
}
.illustration-lightbox[hidden] {
  display: none !important;
}
.illustration-lightbox__dialog {
  position: relative;
  display: flex;
  flex-direction: column;
  align-items: center;
  gap: 0.65rem;
  width: min(1100px, 100%);
  max-height: min(92vh, 100%);
  margin: 0;
  padding: 0;
  border: 0;
  background: transparent;
}
.illustration-lightbox__close {
  position: absolute;
  top: -0.15rem;
  right: -0.15rem;
  z-index: 2;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  width: 2.6rem;
  height: 2.6rem;
  border: 1px solid rgba(255, 255, 255, 0.35);
  border-radius: var(--radius-pill);
  background: rgba(15, 45, 72, 0.88);
  color: #fff;
  font-family: var(--font-ui);
  font-size: 1.45rem;
  font-weight: 500;
  line-height: 1;
  cursor: pointer;
  box-shadow: 0 8px 24px rgba(0, 0, 0, 0.28);
}
.illustration-lightbox__close:hover {
  background: var(--nav-blue);
}
.illustration-lightbox__close:focus-visible {
  outline: 3px solid var(--gold-soft);
  outline-offset: 2px;
}
.illustration-lightbox__frame {
  width: 100%;
  max-height: min(84vh, 100%);
  overflow: auto;
  border-radius: 16px;
  background: rgba(255, 255, 255, 0.98);
  border: 1px solid rgba(255, 255, 255, 0.45);
  box-shadow: 0 24px 64px rgba(0, 20, 40, 0.45);
  padding: 0.65rem;
}
.illustration-lightbox__image {
  display: block;
  width: 100%;
  height: auto;
  max-height: min(80vh, 100%);
  object-fit: contain;
  margin: 0 auto;
  border-radius: 10px;
}
.illustration-lightbox__caption {
  margin: 0;
  max-width: 42rem;
  padding: 0 0.5rem;
  color: rgba(255, 255, 255, 0.92);
  font-family: var(--font-ui);
  font-size: 0.95rem;
  font-weight: 600;
  text-align: center;
  text-shadow: 0 1px 8px rgba(0, 0, 0, 0.35);
}
body.illustration-lightbox-open,
body.text-lightbox-open {
  overflow: hidden;
}
@media (max-width: 720px) {
  .illustration-lightbox__close {
    top: 0.35rem;
    right: 0.35rem;
  }
  .illustration-lightbox__frame {
    padding: 0.4rem;
    border-radius: 12px;
  }
}

.text-lightbox {
  position: fixed;
  inset: 0;
  z-index: 12000;
  display: flex;
  align-items: center;
  justify-content: center;
  padding: max(0.75rem, env(safe-area-inset-top)) max(0.75rem, env(safe-area-inset-right)) max(0.75rem, env(safe-area-inset-bottom)) max(0.75rem, env(safe-area-inset-left));
  background: rgba(8, 28, 48, 0.78);
  backdrop-filter: blur(6px);
  -webkit-backdrop-filter: blur(6px);
}
.text-lightbox[hidden] {
  display: none !important;
}
.text-lightbox__dialog {
  position: relative;
  display: flex;
  flex-direction: column;
  width: min(820px, 100%);
  max-height: min(90vh, 100%);
  margin: 0;
  border: 0;
  border-radius: 18px;
  overflow: hidden;
  background: rgba(255, 255, 255, 0.98);
  box-shadow: 0 24px 64px rgba(0, 20, 40, 0.45);
}
.text-lightbox__close {
  position: absolute;
  top: 0.7rem;
  right: 0.7rem;
  z-index: 2;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  width: 2.6rem;
  height: 2.6rem;
  border: 1px solid rgba(255, 255, 255, 0.35);
  border-radius: var(--radius-pill);
  background: rgba(15, 45, 72, 0.88);
  color: #fff;
  font-family: var(--font-ui);
  font-size: 1.45rem;
  font-weight: 500;
  line-height: 1;
  cursor: pointer;
  box-shadow: 0 8px 24px rgba(0, 0, 0, 0.28);
}
.text-lightbox__close:hover {
  background: var(--nav-blue);
}
.text-lightbox__close:focus-visible {
  outline: 3px solid var(--gold-soft);
  outline-offset: 2px;
}
.text-lightbox__header {
  flex: 0 0 auto;
  padding: 1.15rem 3.4rem 0.95rem 1.35rem;
  background: linear-gradient(135deg, var(--blue-900) 0%, var(--nav-blue) 58%, var(--blue-400) 100%);
  border-bottom: 2px solid var(--gold-bright);
}
.text-lightbox__title {
  margin: 0;
  color: #fff;
  font-family: var(--font-display);
  font-size: clamp(1.15rem, 2.4vw, 1.55rem);
  font-weight: 800;
  line-height: 1.35;
  text-align: center;
}
.text-lightbox__tts {
  flex: 0 0 auto;
  display: flex;
  flex-wrap: wrap;
  align-items: center;
  gap: 0.55rem 0.75rem;
  padding: 0.75rem 1.25rem;
  border-bottom: 1px solid rgba(0, 105, 180, 0.12);
  background: linear-gradient(180deg, #f4fbff 0%, #eef7fc 100%);
}
.text-lightbox__tts .story-tts {
  flex: 0 0 auto;
}
.text-lightbox__tts .story-tts__note {
  flex: 1 1 12rem;
  margin: 0;
  max-width: none;
}
.text-lightbox__body {
  flex: 1 1 auto;
  min-height: 0;
  overflow: auto;
  padding: 1.35rem 1.45rem 1.6rem;
  color: var(--ink);
  font-family: var(--font-body);
  font-size: clamp(1.12rem, 1.7vw, 1.28rem);
  line-height: 1.65;
  text-align: justify;
  text-justify: inter-word;
  hyphens: auto;
  -webkit-overflow-scrolling: touch;
}
.text-lightbox__body p {
  margin: 0 0 0.95rem;
}
.text-lightbox__body p:last-child {
  margin-bottom: 0;
}
.text-lightbox__body .story__source {
  display: inline-block;
  max-width: min(46rem, 100%);
  margin: 0.85rem 0 0;
  padding: 0.32rem 0.85rem;
  border-radius: 1rem;
  color: var(--blue-900);
  font-family: var(--font-display);
  font-size: clamp(0.72rem, 1.05vw, 0.84rem);
  font-style: italic;
  font-weight: 400;
  letter-spacing: 0.008em;
  line-height: 1.4;
  text-align: left;
  background: linear-gradient(180deg, rgba(255, 240, 191, 0.58), rgba(223, 242, 255, 0.42));
  border: 1px solid rgba(201, 155, 59, 0.4);
  box-shadow:
    0 1px 0 rgba(255, 255, 255, 0.78) inset,
    0 6px 16px rgba(0, 78, 140, 0.08);
}
.text-lightbox__body .story__moral {
  clear: both;
  margin: 1.65rem 0 0;
  padding: 0.95rem 1.1rem 1rem;
  background: linear-gradient(135deg, #fff8e4 0%, var(--gold-soft) 55%, #ffe9b8 100%);
  border: 1px solid rgba(201, 155, 59, 0.42);
  border-left: 4px solid var(--gold);
  border-radius: var(--radius-sm);
  color: var(--ink);
  font-family: var(--font-display);
  font-size: 1.02em;
  font-weight: 700;
  line-height: 1.45;
  text-align: left;
  box-shadow: 0 4px 14px rgba(201, 155, 59, 0.12);
}
@media (max-width: 720px) {
  .text-lightbox__dialog {
    border-radius: 14px;
    max-height: min(94vh, 100%);
  }
  .text-lightbox__header {
    padding: 1rem 3.2rem 0.85rem 1.05rem;
  }
  .text-lightbox__body {
    padding: 1.05rem 1.05rem 1.25rem;
  }
}
.story__actions {
  float: right;
  clear: right;
  display: flex;
  flex-direction: row;
  flex-wrap: wrap;
  align-items: flex-end;
  justify-content: flex-end;
  gap: 0.45rem;
  width: max-content;
  max-width: min(100%, 36rem);
  margin: 0 0 0.55rem 1rem;
}
.story__action-group,
.text-lightbox__tts .story__action-group {
  display: inline-flex;
  flex-direction: column;
  align-items: center;
  gap: 0.28rem;
  flex: 0 0 auto;
}
.story__actions .tools-bar__label,
.text-lightbox__tts .tools-bar__label {
  text-align: center;
}
@media (max-width: 760px) {
  .story__actions {
    float: none;
    clear: both;
    width: 100%;
    max-width: none;
    margin: 0 0 0.75rem;
    justify-content: stretch;
  }
  .story__actions .tools-bar__views,
  .text-lightbox__tts .tools-bar__views {
    width: auto;
  }
  :is(.page-home, .page-category) .story__actions .tools-bar__view-btn,
  :is(.page-home, .page-category) .text-lightbox__tts .tools-bar__view-btn {
    flex: 0 0 auto;
  }
  :is(.page-home, .page-category) .tools-bar__view-btn {
    min-height: 2.75rem;
    padding: 0.4rem 0.7rem;
  }
  :is(.page-home, .page-category) .tools-bar__batch-input {
    min-height: 2.75rem;
  }
}
.story__panel::after {
  content: "";
  display: table;
  clear: both;
}
.story.story--text-hidden .story__actions {
  float: none;
  clear: both;
  width: 100%;
  max-width: none;
  margin: 0;
}
.story__actions .tools-bar__views,
.text-lightbox__tts .tools-bar__views {
  flex: 0 0 auto;
  width: auto;
  gap: 0.45rem;
}
:is(.page-home, .page-category) .story__actions .tools-bar__view-btn,
:is(.page-home, .page-category) .text-lightbox__tts .tools-bar__view-btn {
  flex: 0 0 auto;
}
.story-tts {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  gap: 0;
  width: auto;
  flex: 0 0 auto;
  white-space: nowrap;
}
.story-tts .tools-bar__glyph {
  display: block;
  flex: 0 0 auto;
}
.story-tts__note {
  margin: 0;
  flex: 1 0 100%;
  width: 100%;
  text-align: right;
  font-family: var(--font-ui);
  font-size: 0.82rem;
  line-height: 1.35;
  color: var(--ink-soft);
}
.story-tts__note[hidden] {
  display: none !important;
}

.audio-player {
  position: fixed;
  left: 0;
  right: 0;
  bottom: 0;
  z-index: 13000;
  padding: 0.55rem 0.75rem calc(0.55rem + env(safe-area-inset-bottom));
  background: linear-gradient(180deg, rgba(8, 38, 59, 0.96) 0%, rgba(0, 78, 140, 0.98) 100%);
  border-top: 2px solid rgba(240, 199, 94, 0.85);
  box-shadow: 0 -12px 36px rgba(0, 20, 40, 0.28);
  color: #fff;
  font-family: var(--font-ui);
}
.audio-player[hidden] {
  display: none !important;
}
.audio-player__inner {
  max-width: var(--max-wide);
  margin: 0 auto;
  display: grid;
  grid-template-columns: minmax(0, 1.1fr) minmax(0, 1.4fr) auto;
  gap: 0.55rem 1rem;
  align-items: center;
}
.audio-player__meta {
  min-width: 0;
}
.audio-player__title {
  margin: 0;
  font-size: 0.95rem;
  font-weight: 700;
  line-height: 1.3;
  white-space: nowrap;
  overflow: hidden;
  text-overflow: ellipsis;
}
.audio-player__progress {
  display: grid;
  grid-template-columns: auto 1fr auto;
  gap: 0.45rem;
  align-items: center;
  min-width: 0;
}
.audio-player__time {
  font-size: 0.75rem;
  font-variant-numeric: tabular-nums;
  opacity: 0.92;
  min-width: 2.4rem;
}
.audio-player__time--duration {
  text-align: right;
}
.audio-player__seek,
.audio-player__volume {
  -webkit-appearance: none;
  appearance: none;
  width: 100%;
  height: 0.35rem;
  border-radius: var(--radius-pill);
  background: rgba(255, 255, 255, 0.28);
  outline: none;
  cursor: pointer;
}
.audio-player__seek::-webkit-slider-thumb,
.audio-player__volume::-webkit-slider-thumb {
  -webkit-appearance: none;
  appearance: none;
  width: 0.95rem;
  height: 0.95rem;
  border-radius: 50%;
  background: var(--gold-bright);
  border: 2px solid #fff;
  box-shadow: 0 2px 8px rgba(0, 0, 0, 0.25);
  cursor: pointer;
}
.audio-player__seek::-moz-range-thumb,
.audio-player__volume::-moz-range-thumb {
  width: 0.95rem;
  height: 0.95rem;
  border-radius: 50%;
  background: var(--gold-bright);
  border: 2px solid #fff;
  box-shadow: 0 2px 8px rgba(0, 0, 0, 0.25);
  cursor: pointer;
}
.audio-player__controls {
  display: inline-flex;
  align-items: center;
  justify-content: flex-end;
  flex-wrap: wrap;
  gap: 0.35rem;
}
.audio-player__btn {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-width: 2.35rem;
  min-height: 2.35rem;
  padding: 0.35rem 0.55rem;
  border: 1px solid rgba(255, 255, 255, 0.28);
  border-radius: var(--radius-pill);
  background: rgba(255, 255, 255, 0.12);
  color: #fff;
  font-family: var(--font-ui);
  font-size: 0.82rem;
  font-weight: 700;
  cursor: pointer;
  transition: background 140ms ease, transform 140ms ease, border-color 140ms ease;
}
.audio-player__btn:hover {
  background: rgba(255, 255, 255, 0.22);
  transform: translateY(-1px);
}
.audio-player__btn:focus-visible {
  outline: 3px solid var(--gold-soft);
  outline-offset: 2px;
}
.audio-player__btn--play {
  min-width: 2.7rem;
  min-height: 2.7rem;
  background: linear-gradient(135deg, var(--gold-grad-from) 0%, var(--gold-grad-mid) 48%, var(--gold-bright) 100%);
  border-color: rgba(255, 255, 255, 0.45);
  color: #123;
  text-shadow: none;
}
.audio-player__btn--play:hover {
  filter: brightness(1.05);
  background: linear-gradient(135deg, var(--gold-grad-from) 0%, var(--gold-grad-mid) 48%, var(--gold-bright) 100%);
}
.audio-player__btn--story[hidden] {
  display: none !important;
}
.audio-player__btn--story:disabled {
  opacity: 0.4;
  cursor: not-allowed;
  transform: none;
}
.audio-player__btn--close {
  min-width: 2.2rem;
  background: rgba(0, 0, 0, 0.22);
}
.audio-player__speed {
  display: inline-flex;
  align-items: center;
  flex-wrap: wrap;
  gap: 0.28rem;
  padding: 0.2rem 0.35rem;
  border: 1px solid rgba(255, 255, 255, 0.22);
  border-radius: var(--radius-pill);
  background: rgba(0, 0, 0, 0.18);
}
.audio-player__speed-label {
  margin: 0 0.15rem 0 0.25rem;
  font-size: 0.72rem;
  font-weight: 700;
  letter-spacing: 0.02em;
  text-transform: uppercase;
  opacity: 0.9;
  white-space: nowrap;
}
.audio-player__speed-btn {
  min-width: 2.55rem;
  min-height: 1.9rem;
  padding: 0.2rem 0.4rem;
  border: 1px solid transparent;
  border-radius: var(--radius-pill);
  background: transparent;
  color: rgba(255, 255, 255, 0.92);
  font-family: var(--font-ui);
  font-size: 0.78rem;
  font-weight: 700;
  font-variant-numeric: tabular-nums;
  cursor: pointer;
}
.audio-player__speed-btn:hover {
  background: rgba(255, 255, 255, 0.14);
}
.audio-player__speed-btn[aria-pressed="true"] {
  background: linear-gradient(135deg, var(--gold-grad-from) 0%, var(--gold-grad-mid) 55%, var(--gold-bright) 100%);
  border-color: rgba(255, 255, 255, 0.35);
  color: #123;
}
.audio-player__speed-btn:focus-visible {
  outline: 3px solid var(--gold-soft);
  outline-offset: 1px;
}
.audio-player__volume-wrap {
  display: inline-flex;
  align-items: center;
  gap: 0.3rem;
  min-width: 7rem;
}
.audio-player__volume {
  width: 5.5rem;
}
.audio-player audio {
  display: none;
}
body.audio-player-open {
  padding-bottom: calc(var(--audio-player-h, 6.5rem) + 0.5rem);
}
body.audio-player-open .back-to-top {
  bottom: calc(var(--audio-player-h, 6.5rem) + 0.85rem);
}
body.audio-player-open .text-lightbox {
  padding-bottom: calc(var(--audio-player-h, 6.5rem) + env(safe-area-inset-bottom));
}
@media (max-width: 1100px) {
  .audio-player__inner {
    grid-template-columns: 1fr;
    gap: 0.45rem;
  }
  .audio-player__controls {
    justify-content: center;
  }
  .audio-player__title {
    text-align: center;
    white-space: normal;
    display: -webkit-box;
    -webkit-line-clamp: 2;
    -webkit-box-orient: vertical;
  }
  body.audio-player-open {
    padding-bottom: calc(var(--audio-player-h, 9.5rem) + 0.5rem);
  }
  body.audio-player-open .back-to-top {
    bottom: calc(var(--audio-player-h, 9.5rem) + 0.85rem);
  }
  body.audio-player-open .text-lightbox {
    padding-bottom: calc(var(--audio-player-h, 9.5rem) + env(safe-area-inset-bottom));
  }
}
@media (max-width: 520px) {
  .audio-player__speed {
    width: 100%;
    justify-content: center;
    border-radius: 14px;
  }
  .audio-player__volume-wrap {
    min-width: 0;
    width: 100%;
    justify-content: center;
  }
  .audio-player__volume {
    flex: 1 1 auto;
    max-width: 10rem;
  }
}

.story__content {
  display: flow-root;
}
.story__content .story__text {
  display: block;
}
@media (max-width: 760px) {
  .story .card-header,
  .story .card-body {
    padding-left: 18px;
    padding-right: 18px;
  }
  .story .card-title,
  .story__title {
    font-size: clamp(14px, 4.2vw, 17px);
  }
  .story__panel {
    padding: 0.75rem 0.85rem 0.95rem;
    border-radius: 12px;
  }
  .story__text,
  .story .card-text {
    padding: 0;
    border-radius: 0;
  }
}

@media (max-width: 1060px) {
  .category-layout {
    display: flex;
    flex-direction: column;
    gap: 22px;
  }
  .page-category .category-layout > .category-hero {
    grid-column: auto;
    grid-row: auto;
    order: -2;
    margin-bottom: 0;
    width: 100%;
    max-width: none;
  }
  .tools-bar.tools-bar--dense {
    grid-column: auto;
    grid-row: auto;
    order: -1;
    width: 100%;
    max-width: none;
    margin-bottom: 0;
  }
  .page-category .category-layout > .story-nav.sidebar,
  .page-category .category-layout > .story-list {
    grid-column: auto;
    grid-row: auto;
  }
  /* Keep title left-aligned with sidebar stack on narrow screens too. */
  .story-nav.sidebar {
    position: static;
    top: auto;
    max-height: none;
    order: 0;
    z-index: 30;
  }
  .story-list { order: 1; }
  .home-stories-main { order: 1; }
  .story-nav .sidebar-widget {
    max-height: none;
    overflow: hidden;
    border-radius: 18px;
  }
  .story-nav .widget-head {
    min-height: 58px;
  }
  .story-nav .events-menu-toggle {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    flex: 0 0 auto;
    width: 44px;
    height: 42px;
    min-width: 44px;
  }
  .story-nav .sidebar-widget:not(.events-open) .widget-body {
    max-height: 0;
    padding-top: 0;
    padding-bottom: 0;
    overflow: hidden;
  }
  .story-nav .sidebar-widget.events-open .widget-body {
    max-height: min(62vh, 520px);
    padding: 12px 16px;
    overflow-y: auto;
    overscroll-behavior: contain;
    -webkit-overflow-scrolling: touch;
  }
  .story-nav .timeline-list a {
    min-height: 44px;
    padding: 10px 12px 10px calc(12px + 0.85em);
    margin: -6px -8px;
    line-height: 1.35;
  }
  .story-nav .timeline-list a::before {
    left: 12px;
    top: calc(10px + 0.545em);
  }
}
@media (max-width: 720px) {
  .story-nav .widget-head {
    padding: 12px 16px;
    font-size: 16px;
  }
  .story-nav .timeline-list li {
    gap: 10px;
    font-size: 14px;
  }
}

.back-to-top {
  position: fixed;
  right: 1.25rem;
  bottom: 1.25rem;
  z-index: 50;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  width: 3rem;
  height: 3rem;
  padding: 0;
  border-radius: var(--radius-pill);
  border: 1px solid rgba(0, 105, 180, 0.28);
  background: var(--nav-blue);
  color: #fff;
  text-decoration: none;
  box-shadow: 0 8px 24px rgba(0, 78, 140, 0.18);
  transition: transform 160ms ease, background 160ms ease;
}
.back-to-top::before {
  content: "";
  width: 0.7rem;
  height: 0.7rem;
  border-top: 2.5px solid currentColor;
  border-left: 2.5px solid currentColor;
  transform: translateY(0.18rem) rotate(45deg);
}
.back-to-top:hover {
  transform: translateY(-2px);
  background: var(--nav-blue-deep);
  color: #fff;
}
@media (max-width: 620px) {
  .back-to-top {
    right: 0.85rem;
    bottom: 0.85rem;
    width: 2.75rem;
    height: 2.75rem;
  }
}

.footer-pro {
  position: relative;
  overflow: hidden;
  margin-top: auto;
  color: #fff;
  background-color: #274380;
  background-image: url("diaspor-footer-bg.jpg");
  background-repeat: no-repeat;
  background-position: top center;
  background-size: 100% auto;
  border-top: 1px solid rgba(255, 255, 255, 0.18);
  font-family: var(--font-ui);
}
.footer-inner {
  position: relative;
  max-width: var(--max-wide);
  margin: 0 auto;
  padding: 28px 24px;
}
.footer-grid {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 14px;
  align-items: stretch;
}
.footer-col {
  min-height: 7.5rem;
  padding: 18px;
  text-align: center;
  background: rgba(255, 255, 255, 0.1);
  border: 1px solid rgba(255, 255, 255, 0.18);
  border-radius: var(--radius-sm);
  box-shadow: none;
  backdrop-filter: blur(10px);
  -webkit-backdrop-filter: blur(10px);
}
.footer-col--about {
  display: flex;
  align-items: center;
  text-align: justify;
}
.footer-about {
  margin: 0;
  width: 100%;
  color: rgba(255, 255, 255, 0.92);
  font-size: 0.78rem;
  font-weight: 500;
  line-height: 1.55;
  text-align: justify;
  text-justify: inter-word;
  hyphens: auto;
}
.footer-col--brand {
  display: flex;
  align-items: center;
  justify-content: center;
}
.footer-logo {
  display: flex;
  flex-direction: column;
  align-items: center;
  justify-content: center;
  gap: 0.7rem;
  text-decoration: none;
  color: #fff;
}
.footer-logo:hover {
  color: #fff;
  opacity: 0.94;
}
.footer-logo__img {
  width: 72px;
  height: 72px;
  object-fit: contain;
}
.footer-logo__name {
  font-family: var(--font-display);
  font-size: 1.2rem;
  font-weight: 800;
  letter-spacing: -0.01em;
  white-space: nowrap;
}
.footer-col--contact {
  display: flex;
  flex-direction: column;
  align-items: center;
  justify-content: center;
  text-align: center;
  gap: 0.7rem;
}
.footer-contact__title {
  margin: 0;
  color: #fff;
  font-family: var(--font-display);
  font-size: 0.95rem;
  font-weight: 800;
  letter-spacing: 0.01em;
}
.footer-contact {
  list-style: none;
  margin: 0;
  padding: 0;
  display: flex;
  flex-direction: row;
  align-items: flex-start;
  justify-content: center;
  gap: 1.1rem;
  width: 100%;
}
.footer-contact li,
.footer-contact__link {
  display: flex;
  flex-direction: column;
  align-items: center;
  gap: 0.4rem;
  min-width: 3.6rem;
}
.footer-contact__link {
  color: inherit;
  text-decoration: none;
}
.footer-contact__link:hover {
  color: inherit;
  opacity: 1;
}
.footer-contact__link:hover .menu-icon {
  transform: perspective(120px) rotateX(12deg) translateY(-2px);
}
.footer-col--contact .menu-icon {
  width: 2.15rem;
  height: 2.15rem;
  border-radius: 0.7rem;
}
.footer-col--contact .menu-icon__svg {
  width: 16px;
  height: 16px;
}
.footer-contact__label {
  color: rgba(255, 255, 255, 0.72);
  font-size: 0.68rem;
  font-weight: 700;
  letter-spacing: 0.04em;
  text-transform: uppercase;
}
.footer-contact__url {
  color: rgba(255, 255, 255, 0.92);
  font-size: 0.8rem;
  font-weight: 600;
  text-decoration: none;
}
.footer-contact__url:hover {
  color: #fff;
  text-decoration: underline;
  text-underline-offset: 0.18em;
}
.footer-bottom {
  display: flex;
  flex-direction: column;
  align-items: center;
  justify-content: center;
  gap: 8px;
  padding: 12px 20px;
  color: rgba(255, 255, 255, 0.68);
  background: transparent;
  border-top: 1px solid rgba(255, 255, 255, 0.14);
  font-size: 0.75rem;
  text-align: center;
}
.footer-copy {
  width: 100%;
  margin: 0;
  text-align: center;
}
@media (max-width: 900px) {
  .footer-grid {
    grid-template-columns: 1fr;
  }
  .footer-inner {
    padding: 24px 16px;
  }
  .footer-col {
    min-height: 5.5rem;
  }
  .footer-col:empty {
    display: none;
  }
}

@media (prefers-reduced-motion: reduce) {
  *, *::before, *::after {
    animation: none !important;
    transition: none !important;
    scroll-behavior: auto !important;
  }
}
"""

JS = r"""
(() => {
  const I18N = window.__BIRINCI_I18N__ || { lang: "az", ui: {}, js: {} };
  const LOCALE_TAG = I18N.lang || document.documentElement.lang || "az";
  const tUi = (key, fallback) =>
    (I18N.ui && I18N.ui[key]) || fallback || key;
  const tJs = (key, fallback) =>
    (I18N.js && I18N.js[key]) || fallback || key;

  const syncSearchFilterUi = (q, total) => {
    const wrap = document.querySelector(".tools-bar__search");
    if (!wrap) return;
    const chip = wrap.querySelector("[data-search-filter]");
    const textEl = wrap.querySelector("[data-search-filter-text]");
    const raw = String(q || "").trim();
    const active = raw.length > 0;
    wrap.classList.toggle("tools-bar__search--active", active);
    if (!chip) return;
    if (!active) {
      chip.hidden = true;
      if (textEl) textEl.textContent = "";
      return;
    }
    chip.hidden = false;
    if (textEl) {
      const label = tUi("search_filter_label", "Axtarış");
      const count = tUi("search_results_count", "{n} nəticə").replace(/\{n\}/g, String(total));
      textEl.textContent = `${label}: ${raw} · ${count}`;
    }
  };

  const bindSearchFilterClear = (searchInput) => {
    const wrap = searchInput && searchInput.closest(".tools-bar__search");
    const btn = wrap && wrap.querySelector("[data-search-filter-clear]");
    if (!btn || !searchInput || btn.dataset.bound === "1") return;
    btn.dataset.bound = "1";
    btn.addEventListener("click", () => {
      searchInput.value = "";
      searchInput.dispatchEvent(new Event("input", { bubbles: true }));
      searchInput.focus();
    });
  };
  const STORY_ICONS = window.__BIRINCI_STORY_ICONS__ || {
    text: "",
    "text-off": "",
    eye: "",
    "eye-off": "",
    listen: "",
    stop: "",
  };
  const setStoryModePressed = (root, attr, visible) => {
    if (!root) return;
    root.querySelectorAll("[" + attr + "]").forEach((btn) => {
      const mode = btn.getAttribute(attr);
      const pressed = visible ? mode === "show" : mode === "hide";
      btn.setAttribute("aria-pressed", pressed ? "true" : "false");
    });
  };

  const initLangSwitcher = () => {
    const root = document.querySelector(".lang-switcher");
    const toggle = root && root.querySelector(".lang-switcher__toggle");
    const menu = root && root.querySelector(".lang-switcher__menu");
    if (!root || !toggle || !menu) return;

    const supportsPopover = typeof menu.showPopover === "function";
    if (supportsPopover && menu.getAttribute("popover") !== "manual") {
      menu.setAttribute("popover", "manual");
    }
    const finePointerQuery = window.matchMedia("(hover: hover) and (pointer: fine)");
    const coarsePointerQuery = window.matchMedia("(pointer: coarse)");
    const canHoverLang = () => finePointerQuery.matches && !coarsePointerQuery.matches;
    let hideTimer = 0;
    let ignoreOutside = false;
    let outsideBound = false;

    const isPopoverOpen = () => supportsPopover && menu.matches(":popover-open");
    const isOpen = () => root.classList.contains("is-open") || isPopoverOpen();

    const placeMenu = () => {
      const rect = toggle.getBoundingClientRect();
      const gap = 6;
      const top = Math.round(rect.bottom + gap);
      const right = Math.round(Math.max(8, window.innerWidth - rect.right));
      menu.style.position = "fixed";
      menu.style.inset = "auto";
      menu.style.margin = "0";
      menu.style.top = `${top}px`;
      menu.style.right = `${right}px`;
      menu.style.left = "auto";
      menu.style.bottom = "auto";
    };

    const onOutsidePointer = (event) => {
      if (ignoreOutside || !isOpen()) return;
      const target = event.target;
      if (root.contains(target) || menu.contains(target)) return;
      closeMenu();
    };

    const bindOutside = () => {
      if (outsideBound) return;
      outsideBound = true;
      document.addEventListener("pointerdown", onOutsidePointer, true);
    };

    const unbindOutside = () => {
      if (!outsideBound) return;
      outsideBound = false;
      document.removeEventListener("pointerdown", onOutsidePointer, true);
    };

    const openMenu = () => {
      window.clearTimeout(hideTimer);
      menu.hidden = false;
      placeMenu();
      root.classList.add("is-open");
      toggle.setAttribute("aria-expanded", "true");
      if (supportsPopover) {
        try {
          menu.showPopover();
        } catch (_) {}
      }
      ignoreOutside = true;
      window.setTimeout(() => {
        ignoreOutside = false;
        bindOutside();
      }, 0);
    };

    const closeMenu = () => {
      window.clearTimeout(hideTimer);
      root.classList.remove("is-open");
      toggle.setAttribute("aria-expanded", "false");
      if (supportsPopover) {
        try {
          menu.hidePopover();
        } catch (_) {}
      }
      menu.hidden = true;
      unbindOutside();
    };

    const scheduleClose = () => {
      window.clearTimeout(hideTimer);
      hideTimer = window.setTimeout(closeMenu, 160);
    };

    toggle.addEventListener("click", (event) => {
      event.preventDefault();
      event.stopPropagation();
      if (isOpen()) closeMenu();
      else openMenu();
    });

    // iOS synthesizes mouseenter on tap; only hover-open for a real mouse.
    root.addEventListener("pointerenter", (event) => {
      if (event.pointerType === "mouse" && canHoverLang()) openMenu();
    });
    root.addEventListener("pointerleave", (event) => {
      if (event.pointerType === "mouse" && canHoverLang()) scheduleClose();
    });
    menu.addEventListener("pointerenter", (event) => {
      if (event.pointerType === "mouse" && canHoverLang()) window.clearTimeout(hideTimer);
    });
    menu.addEventListener("pointerleave", (event) => {
      if (event.pointerType === "mouse" && canHoverLang()) scheduleClose();
    });

    document.addEventListener("keydown", (event) => {
      if (event.key === "Escape" && isOpen()) closeMenu();
    });
    window.addEventListener("resize", () => {
      if (isOpen()) placeMenu();
    }, { passive: true });
    window.addEventListener("scroll", () => {
      if (isOpen()) placeMenu();
    }, { passive: true, capture: true });

    root.addEventListener("click", (event) => {
      const link = event.target.closest("[data-lang]");
      if (!link) return;
      try {
        localStorage.setItem("birinci-lang", link.getAttribute("data-lang") || "");
      } catch (_) {}
    });
  };
  initLangSwitcher();
  const header = document.querySelector(".site-header");
  const dropdowns = Array.from(document.querySelectorAll(".nav-dropdown"));
  const navToggle = document.getElementById("nav-toggle");
  const mobileNavQuery = window.matchMedia("(max-width: 1400px)");
  const finePointerQuery = window.matchMedia("(hover: hover) and (pointer: fine)");
  const canHoverNav = () => finePointerQuery.matches && !mobileNavQuery.matches;

  const syncStickyChrome = () => {
    const root = document.documentElement;
    if (header) {
      root.style.setProperty("--header-h", `${Math.ceil(header.getBoundingClientRect().height)}px`);
    }
    const crumbs = document.querySelector(".breadcrumbs");
    if (crumbs) {
      root.style.setProperty("--breadcrumb-h", `${Math.ceil(crumbs.getBoundingClientRect().height)}px`);
    }
  };
  if (typeof ResizeObserver !== "undefined") {
    const stickyRo = new ResizeObserver(() => syncStickyChrome());
    if (header) stickyRo.observe(header);
    const crumbsEl = document.querySelector(".breadcrumbs");
    if (crumbsEl) stickyRo.observe(crumbsEl);
  }
  window.addEventListener("resize", syncStickyChrome, { passive: true });
  syncStickyChrome();

  const resetMobileNavSections = () => {
    dropdowns.forEach((dropdown) => {
      dropdown.open = false;
      dropdown.classList.remove("is-hover-open");
    });
    document.querySelectorAll(".nav-dropdown--nested.is-mega-open").forEach((group) => {
      group.classList.remove("is-mega-open");
      const btn = group.querySelector("[data-nav-mega-toggle]");
      if (btn) btn.setAttribute("aria-expanded", "false");
    });
  };

  const closeMobileNav = () => {
    if (!header || !navToggle) return;
    header.classList.remove("is-nav-open");
    document.body.classList.remove("nav-open");
    navToggle.setAttribute("aria-expanded", "false");
    navToggle.setAttribute("aria-label", tUi("open_menu", "Menyunu aç"));
    resetMobileNavSections();
  };

  const openMobileNav = () => {
    if (!header || !navToggle || !dropdowns.length) return;
    resetMobileNavSections();
    header.classList.add("is-nav-open");
    document.body.classList.add("nav-open");
    navToggle.setAttribute("aria-expanded", "true");
    navToggle.setAttribute("aria-label", tUi("close_menu", "Menyunu bağla"));
  };

  if (navToggle && header && dropdowns.length) {
    navToggle.addEventListener("click", (event) => {
      event.stopPropagation();
      if (header.classList.contains("is-nav-open")) closeMobileNav();
      else openMobileNav();
    });
    mobileNavQuery.addEventListener("change", (event) => {
      if (!event.matches) {
        closeMobileNav();
        dropdowns.forEach((dropdown) => {
          dropdown.open = false;
          dropdown.classList.remove("is-hover-open");
        });
      }
    });
  }

  const nestedGroups = Array.from(
    document.querySelectorAll(".nav-dropdown--nested.nav-dropdown--has-mega")
  );

  const setMegaOpen = (target, open) => {
    const scope = target ? target.closest(".nav-dropdown") : null;
    nestedGroups.forEach((group) => {
      if (scope && !scope.contains(group)) return;
      const shouldOpen = !!open && group === target;
      group.classList.toggle("is-mega-open", shouldOpen);
      const btn = group.querySelector("[data-nav-mega-toggle]");
      if (btn) btn.setAttribute("aria-expanded", shouldOpen ? "true" : "false");
    });
  };

  const closeMegasIn = (dropdown) => {
    if (!dropdown) return;
    nestedGroups.forEach((group) => {
      if (!dropdown.contains(group)) return;
      group.classList.remove("is-mega-open");
      const btn = group.querySelector("[data-nav-mega-toggle]");
      if (btn) btn.setAttribute("aria-expanded", "false");
    });
  };

  const setDropdownOpen = (dropdown, open) => {
    if (!dropdown) return;
    dropdown.open = !!open;
    dropdown.classList.toggle("is-hover-open", !!open);
    if (!open) closeMegasIn(dropdown);
  };

  const closeAllDropdowns = () => {
    dropdowns.forEach((dropdown) => setDropdownOpen(dropdown, false));
  };

  nestedGroups.forEach((group) => {
    const megaToggle = group.querySelector("[data-nav-mega-toggle]");
    if (megaToggle) {
      megaToggle.addEventListener("click", (event) => {
        event.preventDefault();
        event.stopPropagation();
        const next = !group.classList.contains("is-mega-open");
        setMegaOpen(group, next);
      });
    }
    group.addEventListener("mouseenter", () => {
      if (canHoverNav()) setMegaOpen(group, true);
    });
    group.addEventListener("mouseleave", () => {
      if (canHoverNav()) setMegaOpen(group, false);
    });
  });

  dropdowns.forEach((dropdown) => {
    dropdown.addEventListener("mouseenter", () => {
      if (!canHoverNav()) return;
      dropdowns.forEach((other) => {
        if (other !== dropdown) setDropdownOpen(other, false);
      });
      setDropdownOpen(dropdown, true);
    });
    dropdown.addEventListener("mouseleave", () => {
      if (canHoverNav()) setDropdownOpen(dropdown, false);
    });
    dropdown.addEventListener("toggle", () => {
      if (!dropdown.open) {
        dropdown.classList.remove("is-hover-open");
        closeMegasIn(dropdown);
        return;
      }
      if (mobileNavQuery.matches) {
        dropdowns.forEach((other) => {
          if (other !== dropdown) setDropdownOpen(other, false);
        });
      }
    });
    dropdown.querySelectorAll("a").forEach((link) => {
      link.addEventListener("click", (event) => {
        if (link.getAttribute("aria-disabled") === "true") {
          event.preventDefault();
          return;
        }
        setDropdownOpen(dropdown, false);
        closeMobileNav();
      });
    });
  });

  if (dropdowns.length) {
    document.addEventListener("click", (event) => {
      if (mobileNavQuery.matches) {
        if (!header || !header.classList.contains("is-nav-open")) return;
        if (header.contains(event.target)) return;
        closeMobileNav();
        return;
      }
      const inside = dropdowns.some((dropdown) => dropdown.contains(event.target));
      if (!inside) closeAllDropdowns();
    });
    document.querySelectorAll(".primary-nav__link").forEach((link) => {
      link.addEventListener("click", (event) => {
        if (link.getAttribute("aria-disabled") === "true") {
          event.preventDefault();
        }
        closeMobileNav();
      });
    });
    document.addEventListener("keydown", (event) => {
      if (event.key !== "Escape") return;
      closeAllDropdowns();
      closeMobileNav();
    });
  }

  const backToTop = document.getElementById("back-to-top");
  if (backToTop) {
    backToTop.addEventListener("click", (event) => {
      event.preventDefault();
      event.stopPropagation();
      const html = document.documentElement;
      html.classList.add("no-smooth-scroll");
      html.scrollTop = 0;
      if (document.body) document.body.scrollTop = 0;
      window.scrollTo(0, 0);
      history.replaceState(null, "", window.location.pathname + window.location.search);
      requestAnimationFrame(() => {
        html.classList.remove("no-smooth-scroll");
      });
    });
  }

  const initGlobalSearch = () => {
    const root = document.getElementById("global-search");
    const toggle = document.getElementById("global-search-toggle");
    const input = document.getElementById("global-search-input");
    const results = document.getElementById("global-search-results");
    const status = document.getElementById("global-search-status");
    if (!root || !toggle || !input || !results) return;

    let index = null;
    let loading = null;
    let lastQuery = "";

    const closeSearch = () => {
      root.hidden = true;
      toggle.setAttribute("aria-expanded", "false");
      document.body.classList.remove("global-search-open");
    };

    const openSearch = () => {
      root.hidden = false;
      toggle.setAttribute("aria-expanded", "true");
      document.body.classList.add("global-search-open");
      window.setTimeout(() => input.focus(), 20);
      ensureIndex();
    };

    const ensureIndex = () => {
      if (index || loading) return loading;
      if (Array.isArray(window.__BIRINCI_SEARCH__)) {
        index = window.__BIRINCI_SEARCH__;
        if (status) status.textContent = `${index.length} hekayə`;
        if (lastQuery) render(lastQuery);
        return Promise.resolve(index);
      }
      const url = root.getAttribute("data-search-index");
      if (!url) return null;
      if (status) status.textContent = "İndeks yüklənir…";
      loading = new Promise((resolve, reject) => {
        const script = document.createElement("script");
        script.src = url;
        script.async = true;
        script.onload = () => {
          if (Array.isArray(window.__BIRINCI_SEARCH__)) resolve(window.__BIRINCI_SEARCH__);
          else reject(new Error("empty-index"));
        };
        script.onerror = () => reject(new Error("script-error"));
        document.head.appendChild(script);
      })
        .then((rows) => {
          index = rows || [];
          if (status) status.textContent = lastQuery ? status.textContent : `${index.length} hekayə`;
          if (lastQuery) render(lastQuery);
        })
        .catch(() => {
          index = [];
          if (status) {
            status.textContent =
              "Axtarış indeksi yüklənmədi. Saytı http://localhost:8765/az/ ünvanından açın.";
          }
        })
        .finally(() => {
          loading = null;
        });
      return loading;
    };

    const render = (query) => {
      lastQuery = query;
      const q = query.trim().toLocaleLowerCase(LOCALE_TAG);
      results.innerHTML = "";
      if (!q) {
        if (status) status.textContent = index ? `${index.length} hekayə` : "";
        return;
      }
      if (!index) {
        if (status) status.textContent = "İndeks yüklənir…";
        return;
      }
      const matches = index.filter((row) => row.hay.includes(q)).slice(0, 40);
      if (status) {
        status.textContent = matches.length
          ? `${matches.length} nəticə`
          : "Uyğun hekayə tapılmadı.";
      }
      const inCategories = window.location.pathname.includes("/categories/");
      const homeListBase = inCategories ? "../index.html" : "index.html";
      matches.forEach((row) => {
        const a = document.createElement("a");
        a.className = "global-search__item";
        a.href = `${homeListBase}?view=list#${encodeURIComponent(row.stem)}`;
        a.innerHTML =
          `<span class="global-search__item-title"></span>` +
          `<span class="global-search__item-meta"></span>`;
        a.querySelector(".global-search__item-title").textContent = row.title;
        a.querySelector(".global-search__item-meta").textContent = row.category;
        a.addEventListener("click", closeSearch);
        results.appendChild(a);
      });
    };

    const kbdHint = toggle.querySelector(".global-search-toggle__kbd");
    if (kbdHint && /Mac|iPhone|iPad|iPod/.test(navigator.platform || navigator.userAgent || "")) {
      kbdHint.textContent = "⌘K";
      toggle.title = "Axtar (⌘K)";
      toggle.setAttribute("aria-label", "Qlobal axtarış, Command+K");
    }

    toggle.addEventListener("click", () => {
      if (root.hidden) openSearch();
      else closeSearch();
    });
    root.querySelectorAll("[data-global-search-close]").forEach((el) => {
      el.addEventListener("click", closeSearch);
    });
    input.addEventListener("input", () => render(input.value));
    document.addEventListener("keydown", (event) => {
      if (event.key === "Escape" && !root.hidden) {
        closeSearch();
        toggle.focus();
      }
      if ((event.ctrlKey || event.metaKey) && event.key.toLowerCase() === "k") {
        event.preventDefault();
        openSearch();
      }
    });
  };

  initGlobalSearch();

  const localeCompareAz = (a, b) =>
    String(a || "").localeCompare(String(b || ""), LOCALE_TAG, { sensitivity: "base" });

  const initCategoryTools = () => {
    if (!document.body.classList.contains("page-category")) return;
    const bar = document.querySelector('[data-tools="category"]');
    const list = document.querySelector("[data-tools-list]");
    const empty = document.querySelector("[data-tools-empty]");
    if (!bar || !list) return;

    const searchInput = bar.querySelector("[data-tools-search]");
    if (!searchInput) return;
    bindSearchFilterClear(searchInput);

    const imagesToggle = bar.querySelector("[data-tools-images]");
    const imagesBtns = Array.from(bar.querySelectorAll("[data-images-mode]"));
    const textsToggle = bar.querySelector("[data-tools-texts]");
    const textsBtns = Array.from(bar.querySelectorAll("[data-texts-mode]"));
    const batchSizeInput = bar.querySelector("[data-home-batch-size]");
    const batchDecBtn = bar.querySelector('[data-home-batch="dec"]');
    const batchIncBtn = bar.querySelector('[data-home-batch="inc"]');
    const batchPrevBtn = bar.querySelector('[data-home-batch="prev"]');
    const batchNextBtn = bar.querySelector('[data-home-batch="next"]');
    const batchRandomBtn = bar.querySelector('[data-home-batch="random"]');
    const batchAllBtn = bar.querySelector('[data-home-batch="all"]');
    const batchRangeEl = bar.querySelector("[data-home-batch-range]");
    const navList = document.querySelector("[data-tools-nav]");
    const countEl = document.querySelector("[data-tools-count]");
    const batchSizeStorageKey = "birinci-category-batch-size";
    const batchAllStorageKey = "birinci-category-batch-all";
    // One-shot migration from the pre-pager page-size key; removed after read/persist.
    const legacyPageSizeStorageKey = "birinci-category-page-size";

    const allStories = Array.from(list.querySelectorAll(".story"));
    allStories.sort((a, b) => localeCompareAz(a.dataset.title, b.dataset.title));
    allStories.forEach((story) => list.appendChild(story));
    if (navList) {
      const navItems = Array.from(navList.querySelectorAll("li[data-stem]"));
      navItems.sort((a, b) => localeCompareAz(a.dataset.title, b.dataset.title));
      navItems.forEach((item) => navList.appendChild(item));
    }

    let filtered = [];
    let batchSize = 12;
    let windowStart = 0;
    let randomStems = null;
    let allMode = false;
    let pendingStem = null;

    const batchCap = () => {
      const n =
        (filtered && filtered.length) ||
        (allStories && allStories.length) ||
        0;
      return Math.max(1, n);
    };

    const inputRaw = () =>
      batchSizeInput ? String(batchSizeInput.value || "").trim() : "";

    const readBatchSize = () => {
      const raw = inputRaw();
      if (!raw) return batchSize || 12;
      const n = Number(raw);
      return Number.isFinite(n) && n > 0 ? Math.floor(n) : batchSize || 12;
    };

    const persistBatchSize = () => {
      try {
        localStorage.setItem(batchSizeStorageKey, String(batchSize));
        localStorage.removeItem(legacyPageSizeStorageKey);
      } catch (_) {}
    };

    const persistAllMode = () => {
      try {
        if (allMode) localStorage.setItem(batchAllStorageKey, "1");
        else localStorage.removeItem(batchAllStorageKey);
      } catch (_) {}
    };

    const syncBatchUi = (visibleCount = 0) => {
      const total = (filtered && filtered.length) || 0;
      const cap = batchCap();
      if (batchSizeInput) {
        batchSizeInput.min = "1";
        batchSizeInput.max = String(cap);
        batchSizeInput.value = String(batchSize);
      }
      const inRandom = !!(randomStems && randomStems.length);
      const atStart = !inRandom && windowStart <= 0;
      const atEnd = allMode || total === 0 || inRandom || windowStart + batchSize >= total;
      if (batchDecBtn) batchDecBtn.disabled = total === 0 || batchSize <= 1;
      if (batchIncBtn) batchIncBtn.disabled = total === 0 || batchSize >= cap;
      if (batchPrevBtn) batchPrevBtn.disabled = total === 0 || allMode || (!inRandom && atStart);
      if (batchNextBtn) batchNextBtn.disabled = total === 0 || allMode || inRandom || atEnd;
      if (batchRandomBtn) batchRandomBtn.disabled = total === 0;
      bar.querySelectorAll("[data-tools-play-visible]").forEach((btn) => {
        btn.disabled = total === 0;
      });
      if (typeof window.__birinciSyncPlayVisibleUi === "function") {
        window.__birinciSyncPlayVisibleUi();
      }
      if (batchAllBtn) {
        const showingAll = allMode && !inRandom && total > 0;
        batchAllBtn.disabled = total === 0;
        batchAllBtn.classList.toggle("is-active", showingAll);
        batchAllBtn.setAttribute("aria-pressed", showingAll ? "true" : "false");
        batchAllBtn.removeAttribute("aria-disabled");
      }
      if (batchRangeEl) {
        if (total === 0) {
          batchRangeEl.hidden = true;
          batchRangeEl.textContent = "";
        } else {
          batchRangeEl.hidden = false;
          batchRangeEl.removeAttribute("hidden");
          if (inRandom) {
            batchRangeEl.textContent = `${tUi("batch_random", "Təsadüfi")} · ${visibleCount} / ${total}`;
          } else if (allMode) {
            batchRangeEl.textContent = `1–${total} / ${total}`;
          } else {
            const from = windowStart + 1;
            const to = Math.max(from, windowStart + visibleCount);
            batchRangeEl.textContent = `${from}–${to} / ${total}`;
          }
        }
      }
    };

    const applyPageSize = (n, { persist = true, render = false } = {}) => {
      const cap = batchCap();
      let size = Number(n);
      if (!Number.isFinite(size) || size < 1) size = batchSize || 12;
      size = Math.min(Math.floor(size), cap);
      if (size < 1) size = 1;
      batchSize = size;
      allMode = false;
      randomStems = null;
      windowStart = 0;
      if (batchSizeInput) batchSizeInput.value = String(batchSize);
      if (persist) {
        persistBatchSize();
        persistAllMode();
      }
      if (render) {
        pendingStem = null;
        renderList();
      } else {
        syncBatchUi(0);
      }
    };

    const commitBatchSize = ({ persist = true, render = false } = {}) => {
      applyPageSize(readBatchSize(), { persist, render });
    };

    const applyStoredBatchSize = () => {
      let stored = "";
      let storedAll = false;
      try {
        storedAll = localStorage.getItem(batchAllStorageKey) === "1";
        stored = localStorage.getItem(batchSizeStorageKey) || "";
        if (!stored) {
          const legacy = localStorage.getItem(legacyPageSizeStorageKey) || "";
          if (legacy && legacy !== "all") stored = legacy;
          else if (legacy === "all") storedAll = true;
        }
      } catch (_) {}
      const n = Number(stored);
      batchSize = Number.isFinite(n) && n > 0 ? Math.floor(n) : 12;
      allMode = !!storedAll;
      if (batchSizeInput) batchSizeInput.value = String(batchSize);
      syncBatchUi(0);
    };

    applyStoredBatchSize();

    const pickRandomStems = (count) => {
      const total = filtered.length;
      if (!total) return [];
      const n = Math.min(Math.max(1, count), total);
      const idxs = Array.from({ length: total }, (_, i) => i);
      for (let i = idxs.length - 1; i > 0; i -= 1) {
        const j = Math.floor(Math.random() * (i + 1));
        const tmp = idxs[i];
        idxs[i] = idxs[j];
        idxs[j] = tmp;
      }
      return idxs.slice(0, n).map((i) => filtered[i].dataset.stem);
    };

    const applyImagesState = (collapsed) => {
      document.body.classList.toggle("images-collapsed", collapsed);
      imagesBtns.forEach((btn) => {
        const mode = btn.getAttribute("data-images-mode");
        const pressed = collapsed ? mode === "hide" : mode === "show";
        btn.setAttribute("aria-pressed", pressed ? "true" : "false");
      });
      if (typeof window.__birinciSetAllStoryFigures === "function") {
        window.__birinciSetAllStoryFigures(!collapsed);
      } else {
        document.querySelectorAll("article.story").forEach((story) => {
          story.classList.toggle("story--figure-hidden", collapsed);
          setStoryModePressed(story, "data-images-mode", !collapsed);
        });
      }
      try {
        localStorage.setItem("birinci-images-collapsed", collapsed ? "1" : "0");
      } catch (_) {}
    };

    if (imagesToggle && imagesBtns.length) {
      let collapsed = false;
      try {
        collapsed = localStorage.getItem("birinci-images-collapsed") === "1";
      } catch (_) {}
      applyImagesState(collapsed);
      imagesBtns.forEach((btn) => {
        btn.addEventListener("click", () => {
          applyImagesState(btn.getAttribute("data-images-mode") === "hide");
        });
      });
    }

    const applyTextsState = (collapsed) => {
      document.body.classList.toggle("texts-collapsed", collapsed);
      textsBtns.forEach((btn) => {
        const mode = btn.getAttribute("data-texts-mode");
        const pressed = collapsed ? mode === "hide" : mode === "show";
        btn.setAttribute("aria-pressed", pressed ? "true" : "false");
      });
      if (typeof window.__birinciSetAllStoryTexts === "function") {
        window.__birinciSetAllStoryTexts(!collapsed);
      } else {
        document.querySelectorAll("article.story").forEach((story) => {
          story.classList.toggle("story--text-hidden", collapsed);
          setStoryModePressed(story, "data-texts-mode", !collapsed);
        });
      }
      try {
        localStorage.setItem("birinci-texts-collapsed", collapsed ? "1" : "0");
      } catch (_) {}
    };

    if (textsToggle && textsBtns.length) {
      let textsCollapsed = false;
      try {
        textsCollapsed = localStorage.getItem("birinci-texts-collapsed") === "1";
      } catch (_) {}
      applyTextsState(textsCollapsed);
      textsBtns.forEach((btn) => {
        btn.addEventListener("click", () => {
          applyTextsState(btn.getAttribute("data-texts-mode") === "hide");
        });
      });
    }

    const escapeHtml = (value) =>
      String(value || "")
        .replace(/&/g, "&amp;")
        .replace(/</g, "&lt;")
        .replace(/>/g, "&gt;")
        .replace(/"/g, "&quot;")
        .replace(/'/g, "&#39;");

    const refreshSidebarNav = (visibleStories) => {
      if (navList) {
        navList.innerHTML = visibleStories
          .map(
            (s) =>
              `<li data-stem="${escapeHtml(s.dataset.stem)}" data-title="${escapeHtml(s.dataset.title)}"><a href="#${escapeHtml(
                s.dataset.stem
              )}">${escapeHtml(s.dataset.title)}</a></li>`
          )
          .join("");
      }
      const layout = document.querySelector(".category-layout");
      if (layout && typeof window.__birinciBindStorySidebar === "function") {
        window.__birinciBindStorySidebar(layout);
      } else if (layout && layout.__birinciSidebar) {
        layout.__birinciSidebar.refresh();
      }
    };

    const scrollToolsIntoView = () => {
      try {
        bar.scrollIntoView({ block: "nearest", behavior: "auto" });
      } catch (_) {}
    };

    const renderList = ({ resetWindow = false } = {}) => {
      if (typeof window.__birinciStopStoryTts === "function") window.__birinciStopStoryTts();
      const q = searchInput.value.trim().toLocaleLowerCase(LOCALE_TAG);
      filtered = allStories.filter((story) => {
        const textEl = story.querySelector(".story__text");
        const hay = `${story.dataset.title || ""} ${textEl ? textEl.textContent : ""}`.toLocaleLowerCase(LOCALE_TAG);
        return !q || hay.includes(q);
      });

      const total = filtered.length;
      syncSearchFilterUi(searchInput.value.trim(), total);
      const cap = Math.max(1, total || 1);
      let n = readBatchSize();
      if (!Number.isFinite(n) || n < 1) n = batchSize || 12;
      if (n > cap) n = cap;
      batchSize = n;
      if (batchSizeInput && String(batchSizeInput.value) !== String(n)) {
        batchSizeInput.value = String(n);
      }
      if (!allMode) {
        try {
          localStorage.setItem(batchSizeStorageKey, String(n));
        } catch (_) {}
      }

      if (resetWindow) {
        windowStart = 0;
        randomStems = null;
      }

      if (pendingStem) {
        const idx = filtered.findIndex(
          (story) => story.dataset.stem === pendingStem || story.dataset.stem === String(pendingStem)
        );
        if (idx >= 0) {
          randomStems = null;
          if (!allMode) {
            windowStart = Math.floor(idx / batchSize) * batchSize;
          }
        }
      }

      let visibleStories = [];
      if (total === 0) {
        windowStart = 0;
        randomStems = null;
        visibleStories = [];
      } else if (randomStems && randomStems.length) {
        allMode = false;
        const byStem = new Map(filtered.map((story) => [story.dataset.stem, story]));
        visibleStories = randomStems.map((stem) => byStem.get(stem)).filter(Boolean);
        if (!visibleStories.length) {
          randomStems = null;
          windowStart = 0;
          visibleStories = filtered.slice(0, batchSize);
        }
      } else if (allMode) {
        randomStems = null;
        windowStart = 0;
        visibleStories = filtered.slice();
      } else {
        randomStems = null;
        const maxStart = Math.max(0, total - 1);
        if (windowStart > maxStart) windowStart = Math.floor(maxStart / batchSize) * batchSize;
        if (windowStart < 0) windowStart = 0;
        visibleStories = filtered.slice(windowStart, windowStart + batchSize);
      }

      const visibleSet = new Set(visibleStories.map((s) => s.dataset.stem));
      allStories.forEach((story) => {
        story.hidden = !visibleSet.has(story.dataset.stem);
      });
      visibleStories.forEach((story) => list.appendChild(story));
      allStories
        .filter((story) => !visibleSet.has(story.dataset.stem))
        .forEach((story) => list.appendChild(story));

      if (typeof window.__birinciSetAllStoryFigures === "function") {
        window.__birinciSetAllStoryFigures(!document.body.classList.contains("images-collapsed"));
      }
      if (typeof window.__birinciSetAllStoryTexts === "function") {
        window.__birinciSetAllStoryTexts(!document.body.classList.contains("texts-collapsed"));
      }
      refreshSidebarNav(visibleStories);
      if (countEl) countEl.textContent = String(total);
      if (empty) empty.hidden = total !== 0;
      if (typeof window.__birinciClearListenQueue === "function") {
        window.__birinciClearListenQueue({ keepTrack: true });
      }
      syncBatchUi(visibleStories.length);
      persistAllMode();
      if (pendingStem) {
        const el = document.getElementById(pendingStem);
        if (el) {
          window.requestAnimationFrame(() => {
            el.scrollIntoView({ block: "start", behavior: "auto" });
          });
        }
        pendingStem = null;
      }
    };

    searchInput.addEventListener("input", () => {
      pendingStem = null;
      renderList({ resetWindow: true });
    });
    if (batchSizeInput) {
      batchSizeInput.addEventListener("change", () => {
        commitBatchSize({ persist: true, render: true });
      });
      batchSizeInput.addEventListener("blur", () => {
        commitBatchSize({ persist: true, render: true });
      });
      batchSizeInput.addEventListener("keydown", (event) => {
        if (event.key !== "Enter") return;
        event.preventDefault();
        commitBatchSize({ persist: true, render: true });
        batchSizeInput.blur();
      });
    }
    const runBatchAction = (action) => {
      const total = filtered.length;
      if (!total) {
        randomStems = null;
        windowStart = 0;
        persistBatchSize();
        persistAllMode();
        renderList();
        return;
      }
      if (action === "all") {
        if (allMode && !randomStems) {
          allMode = false;
        } else {
          allMode = true;
          randomStems = null;
        }
        windowStart = 0;
        persistBatchSize();
        persistAllMode();
        pendingStem = null;
        renderList();
        scrollToolsIntoView();
        return;
      }
      if (action === "dec") {
        applyPageSize(batchSize - 1, { persist: true, render: true });
        scrollToolsIntoView();
        return;
      }
      if (action === "inc") {
        applyPageSize(batchSize + 1, { persist: true, render: true });
        scrollToolsIntoView();
        return;
      }
      if (action === "prev") {
        allMode = false;
        if (randomStems) {
          randomStems = null;
          windowStart = 0;
        } else {
          windowStart = Math.max(0, windowStart - batchSize);
        }
      } else if (action === "next") {
        allMode = false;
        randomStems = null;
        if (windowStart + batchSize < total) {
          windowStart += batchSize;
        }
      } else if (action === "random") {
        allMode = false;
        randomStems = pickRandomStems(batchSize);
      } else {
        return;
      }
      persistBatchSize();
      persistAllMode();
      pendingStem = null;
      renderList();
      scrollToolsIntoView();
    };
    if (batchDecBtn) batchDecBtn.addEventListener("click", () => runBatchAction("dec"));
    if (batchIncBtn) batchIncBtn.addEventListener("click", () => runBatchAction("inc"));
    if (batchPrevBtn) batchPrevBtn.addEventListener("click", () => runBatchAction("prev"));
    if (batchNextBtn) batchNextBtn.addEventListener("click", () => runBatchAction("next"));
    if (batchRandomBtn) batchRandomBtn.addEventListener("click", () => runBatchAction("random"));
    if (batchAllBtn) batchAllBtn.addEventListener("click", () => runBatchAction("all"));

    try {
      const hash = decodeURIComponent((window.location.hash || "").replace(/^#/, ""));
      if (hash) pendingStem = hash;
    } catch (_) {}

    renderList();
  };

  initCategoryTools();

  /**
   * DAAB News-style sidebar: sticky TOC, scroll-spy, mobile accordion.
   * (No dual-panel scroll sync — it fought page scroll.)
   */
  const bindStorySidebarLayout = (layout) => {
    if (!layout) return null;
    if (layout.__birinciSidebar) {
      layout.__birinciSidebar.refresh();
      return layout.__birinciSidebar;
    }

    const nav = layout.querySelector(".story-nav.sidebar");
    if (!nav) return null;
    const widget = nav.querySelector(".sidebar-widget");
    const toggle = nav.querySelector(".events-menu-toggle");
    const mobileQuery = window.matchMedia("(max-width: 1060px)");

    let links = [];
    let cards = [];

    const closeMenu = () => {
      if (!widget || !toggle) return;
      widget.classList.remove("events-open");
      toggle.setAttribute("aria-expanded", "false");
    };
    const toggleMenu = () => {
      if (!widget || !toggle) return;
      const open = widget.classList.toggle("events-open");
      toggle.setAttribute("aria-expanded", open ? "true" : "false");
    };

    const setActive = (activeLink) => {
      links.forEach((link) => {
        const on = link === activeLink;
        link.classList.toggle("is-active", on);
        link.classList.toggle("tl-active", on);
      });
    };

    const updateActive = () => {
      if (!cards.length) {
        setActive(null);
        return;
      }
      const mid = window.scrollY + window.innerHeight * 0.35;
      let active = null;
      for (let i = cards.length - 1; i >= 0; i -= 1) {
        const top = cards[i].el.getBoundingClientRect().top + window.scrollY;
        if (top <= mid) {
          active = cards[i].link;
          break;
        }
      }
      setActive(active);
    };

    const refresh = () => {
      links = Array.from(nav.querySelectorAll('.timeline-list a[href^="#"]'));
      cards = links
        .map((link) => {
          const raw = (link.getAttribute("href") || "").slice(1);
          let id = raw;
          try {
            id = decodeURIComponent(raw);
          } catch (_) {}
          const el = document.getElementById(id);
          return el ? { link, el, id } : null;
        })
        .filter(Boolean);
      updateActive();
    };

    if (toggle) {
      toggle.addEventListener("click", (event) => {
        event.stopPropagation();
        toggleMenu();
      });
    }
    document.addEventListener("click", (event) => {
      if (!mobileQuery.matches || !widget || !widget.classList.contains("events-open")) return;
      if (widget.contains(event.target)) return;
      closeMenu();
    });
    document.addEventListener("keydown", (event) => {
      if (event.key === "Escape") closeMenu();
    });
    mobileQuery.addEventListener("change", () => {
      if (!mobileQuery.matches) closeMenu();
    });

    nav.addEventListener("click", (event) => {
      const link = event.target.closest('a[href^="#"]');
      if (!link || !nav.contains(link)) return;
      const raw = (link.getAttribute("href") || "").slice(1);
      let id = raw;
      try {
        id = decodeURIComponent(raw);
      } catch (_) {}
      const target = document.getElementById(id);
      if (!target) return;
      event.preventDefault();
      setActive(link);
      const html = document.documentElement;
      const prevBehavior = html.style.scrollBehavior;
      html.style.scrollBehavior = "auto";
      target.scrollIntoView({ block: "start", behavior: "auto" });
      html.style.scrollBehavior = prevBehavior;
      try {
        history.pushState(null, "", `#${id}`);
      } catch (_) {}
      if (mobileQuery.matches) closeMenu();
    });

    window.addEventListener("scroll", updateActive, { passive: true });
    window.addEventListener("resize", updateActive, { passive: true });

    const api = { refresh, closeMenu, updateActive };
    layout.__birinciSidebar = api;
    refresh();
    return api;
  };

  window.__birinciBindStorySidebar = bindStorySidebarLayout;

  const initHomeViews = () => {
    if (!document.body.classList.contains("page-home")) return;
    const bar = document.querySelector('[data-tools="home"]');
    const cardsPanel = document.querySelector('[data-view="cards"]');
    const listPanel = document.querySelector('[data-view="list"]');
    if (!bar || !cardsPanel || !listPanel) return;

    const searchInput = bar.querySelector("[data-tools-search]");
    if (!searchInput) return;
    bindSearchFilterClear(searchInput);

    const cardsList = cardsPanel.querySelector("[data-tools-list]");
    const cardsEmpty = cardsPanel.querySelector("[data-tools-empty]");
    const storiesList = listPanel.querySelector("[data-stories-list]");
    const listEmpty = listPanel.querySelector("[data-home-list-empty]");
    const navList = listPanel.querySelector("[data-home-nav]");
    const imagesToggle = bar.querySelector("[data-tools-images]");
    const imagesBtns = Array.from(bar.querySelectorAll("[data-images-mode]"));
    const textsToggle = bar.querySelector("[data-tools-texts]");
    const textsBtns = Array.from(bar.querySelectorAll("[data-texts-mode]"));
    const batchSizeInput = bar.querySelector("[data-home-batch-size]");
    const batchDecBtn = bar.querySelector('[data-home-batch="dec"]');
    const batchIncBtn = bar.querySelector('[data-home-batch="inc"]');
    const batchPrevBtn = bar.querySelector('[data-home-batch="prev"]');
    const batchNextBtn = bar.querySelector('[data-home-batch="next"]');
    const batchRandomBtn = bar.querySelector('[data-home-batch="random"]');
    const batchAllBtn = bar.querySelector('[data-home-batch="all"]');
    const batchRangeEl = bar.querySelector("[data-home-batch-range]");
    const viewBtns = Array.from(bar.querySelectorAll("[data-home-view]"));
    const listOnly = Array.from(bar.querySelectorAll("[data-home-list-only]"));
    const storiesUrl = listPanel.getAttribute("data-stories-url") || "data/stories.json";
    const assetVersion = listPanel.getAttribute("data-asset-version") || "";
    const viewStorageKey = "birinci-home-view";
    const batchSizeStorageKey = "birinci-home-batch-size";
    const batchAllStorageKey = "birinci-home-batch-all";
    const legacyPageSizeStorageKey = "birinci-home-page-size";
    // One-shot migration from the pre-pager page-size key; removed after read/persist.

    let view = "cards";
    let allStories = null;
    let filtered = [];
    let loading = null;
    let pendingStem = null;
    let batchSize = 12;
    let windowStart = 0;
    let randomStems = null;
    let allMode = false;

    const batchCap = () => {
      const n =
        (filtered && filtered.length) ||
        (allStories && allStories.length) ||
        0;
      return Math.max(1, n);
    };

    const inputRaw = () =>
      batchSizeInput ? String(batchSizeInput.value || "").trim() : "";

    const readBatchSize = () => {
      const raw = inputRaw();
      if (!raw) return batchSize || 12;
      const n = Number(raw);
      return Number.isFinite(n) && n > 0 ? Math.floor(n) : batchSize || 12;
    };

    const persistBatchSize = () => {
      try {
        localStorage.setItem(batchSizeStorageKey, String(batchSize));
        localStorage.removeItem(legacyPageSizeStorageKey);
      } catch (_) {}
    };

    const persistAllMode = () => {
      try {
        if (allMode) localStorage.setItem(batchAllStorageKey, "1");
        else localStorage.removeItem(batchAllStorageKey);
      } catch (_) {}
    };

    const syncBatchUi = (visibleCount = 0) => {
      const total = (filtered && filtered.length) || 0;
      const cap = batchCap();
      if (batchSizeInput) {
        batchSizeInput.min = "1";
        batchSizeInput.max = String(cap);
        batchSizeInput.value = String(batchSize);
      }
      const inRandom = !!(randomStems && randomStems.length);
      const atStart = !inRandom && windowStart <= 0;
      const atEnd = allMode || total === 0 || inRandom || windowStart + batchSize >= total;
      if (batchDecBtn) batchDecBtn.disabled = total === 0 || batchSize <= 1;
      if (batchIncBtn) batchIncBtn.disabled = total === 0 || batchSize >= cap;
      if (batchPrevBtn) batchPrevBtn.disabled = total === 0 || allMode || (!inRandom && atStart);
      if (batchNextBtn) batchNextBtn.disabled = total === 0 || allMode || inRandom || atEnd;
      if (batchRandomBtn) batchRandomBtn.disabled = total === 0;
      bar.querySelectorAll("[data-tools-play-visible]").forEach((btn) => {
        btn.disabled = total === 0;
      });
      if (typeof window.__birinciSyncPlayVisibleUi === "function") {
        window.__birinciSyncPlayVisibleUi();
      }
      if (batchAllBtn) {
        const showingAll = allMode && !inRandom && total > 0;
        batchAllBtn.disabled = total === 0;
        batchAllBtn.classList.toggle("is-active", showingAll);
        batchAllBtn.setAttribute("aria-pressed", showingAll ? "true" : "false");
        batchAllBtn.removeAttribute("aria-disabled");
      }
      if (batchRangeEl) {
        if (total === 0) {
          batchRangeEl.hidden = true;
          batchRangeEl.textContent = "";
        } else {
          batchRangeEl.hidden = false;
          batchRangeEl.removeAttribute("hidden");
          if (inRandom) {
            batchRangeEl.textContent = `${tUi("batch_random", "Təsadüfi")} · ${visibleCount} / ${total}`;
          } else if (allMode) {
            batchRangeEl.textContent = `1–${total} / ${total}`;
          } else {
            const from = windowStart + 1;
            const to = Math.max(from, windowStart + visibleCount);
            batchRangeEl.textContent = `${from}–${to} / ${total}`;
          }
        }
      }
    };

    const applyPageSize = (n, { persist = true, render = false } = {}) => {
      const cap = batchCap();
      let size = Number(n);
      if (!Number.isFinite(size) || size < 1) size = batchSize || 12;
      size = Math.min(Math.floor(size), cap);
      if (size < 1) size = 1;
      batchSize = size;
      allMode = false;
      randomStems = null;
      windowStart = 0;
      if (batchSizeInput) batchSizeInput.value = String(batchSize);
      if (persist) {
        persistBatchSize();
        persistAllMode();
      }
      if (render && view === "list") {
        pendingStem = null;
        renderList();
      } else {
        syncBatchUi(0);
      }
    };

    const commitBatchSize = ({ persist = true, render = false } = {}) => {
      applyPageSize(readBatchSize(), { persist, render });
    };

    const applyStoredBatchSize = () => {
      let stored = "";
      let storedAll = false;
      try {
        storedAll = localStorage.getItem(batchAllStorageKey) === "1";
        stored = localStorage.getItem(batchSizeStorageKey) || "";
        if (!stored) {
          const legacy = localStorage.getItem(legacyPageSizeStorageKey) || "";
          if (legacy && legacy !== "all") stored = legacy;
          else if (legacy === "all") storedAll = true;
        }
      } catch (_) {}
      const n = Number(stored);
      batchSize = Number.isFinite(n) && n > 0 ? Math.floor(n) : 12;
      allMode = !!storedAll;
      if (batchSizeInput) batchSizeInput.value = String(batchSize);
      syncBatchUi(0);
    };

    applyStoredBatchSize();

    const pickRandomStems = (count) => {
      const total = filtered.length;
      if (!total) return [];
      const n = Math.min(Math.max(1, count), total);
      const idxs = Array.from({ length: total }, (_, i) => i);
      for (let i = idxs.length - 1; i > 0; i -= 1) {
        const j = Math.floor(Math.random() * (i + 1));
        const tmp = idxs[i];
        idxs[i] = idxs[j];
        idxs[j] = tmp;
      }
      return idxs.slice(0, n).map((i) => filtered[i].stem);
    };

    const escapeHtml = (value) =>
      String(value || "")
        .replace(/&/g, "&amp;")
        .replace(/</g, "&lt;")
        .replace(/>/g, "&gt;")
        .replace(/"/g, "&quot;")
        .replace(/'/g, "&#39;");

    const readUrlState = () => {
      const params = new URLSearchParams(window.location.search);
      const hash = decodeURIComponent((window.location.hash || "").replace(/^#/, ""));
      return {
        view: params.get("view"),
        q: params.get("q"),
        stem: hash || null,
      };
    };

    const writeUrlState = () => {
      try {
        const params = new URLSearchParams();
        if (view === "list") params.set("view", "list");
        const q = searchInput.value.trim();
        if (view === "list" && q) params.set("q", q);
        const url = new URL(window.location.href);
        url.search = params.toString();
        url.hash = pendingStem ? pendingStem : "";
        history.replaceState(null, "", `${url.pathname}${url.search}${url.hash}`);
      } catch (_) {
        /* file:// or sandboxed histories must not block view switching */
      }
    };

    const applyImagesState = (collapsed) => {
      document.body.classList.toggle("images-collapsed", collapsed);
      imagesBtns.forEach((btn) => {
        const mode = btn.getAttribute("data-images-mode");
        const pressed = collapsed ? mode === "hide" : mode === "show";
        btn.setAttribute("aria-pressed", pressed ? "true" : "false");
      });
      if (typeof window.__birinciSetAllStoryFigures === "function") {
        window.__birinciSetAllStoryFigures(!collapsed);
      } else {
        document.querySelectorAll("article.story").forEach((story) => {
          story.classList.toggle("story--figure-hidden", collapsed);
          setStoryModePressed(story, "data-images-mode", !collapsed);
        });
      }
      try {
        localStorage.setItem("birinci-images-collapsed", collapsed ? "1" : "0");
      } catch (_) {}
    };

    if (imagesToggle && imagesBtns.length) {
      let collapsed = false;
      try {
        collapsed = localStorage.getItem("birinci-images-collapsed") === "1";
      } catch (_) {}
      applyImagesState(collapsed);
      imagesBtns.forEach((btn) => {
        btn.addEventListener("click", () => {
          applyImagesState(btn.getAttribute("data-images-mode") === "hide");
        });
      });
    }

    const applyTextsState = (collapsed) => {
      document.body.classList.toggle("texts-collapsed", collapsed);
      textsBtns.forEach((btn) => {
        const mode = btn.getAttribute("data-texts-mode");
        const pressed = collapsed ? mode === "hide" : mode === "show";
        btn.setAttribute("aria-pressed", pressed ? "true" : "false");
      });
      if (typeof window.__birinciSetAllStoryTexts === "function") {
        window.__birinciSetAllStoryTexts(!collapsed);
      } else {
        document.querySelectorAll("article.story").forEach((story) => {
          story.classList.toggle("story--text-hidden", collapsed);
          setStoryModePressed(story, "data-texts-mode", !collapsed);
        });
      }
      try {
        localStorage.setItem("birinci-texts-collapsed", collapsed ? "1" : "0");
      } catch (_) {}
    };

    if (textsToggle && textsBtns.length) {
      let textsCollapsed = false;
      try {
        textsCollapsed = localStorage.getItem("birinci-texts-collapsed") === "1";
      } catch (_) {}
      applyTextsState(textsCollapsed);
      textsBtns.forEach((btn) => {
        btn.addEventListener("click", () => {
          applyTextsState(btn.getAttribute("data-texts-mode") === "hide");
        });
      });
    }

    const applyCards = () => {
      if (!cardsList) return;
      const q = searchInput.value.trim().toLocaleLowerCase(LOCALE_TAG);
      const items = Array.from(cardsList.querySelectorAll(".cat-card"));
      items.sort((a, b) => localeCompareAz(a.dataset.title, b.dataset.title));
      items.forEach((item) => cardsList.appendChild(item));
      let visible = 0;
      items.forEach((item) => {
        const hay = `${item.dataset.title || ""} ${item.dataset.blurb || ""}`.toLocaleLowerCase(LOCALE_TAG);
        const show = !q || hay.includes(q);
        item.hidden = !show;
        if (show) visible += 1;
      });
      if (cardsEmpty) cardsEmpty.hidden = visible !== 0;
      syncSearchFilterUi(searchInput.value.trim(), visible);
    };

    const flattenStories = (catalog) => {
      const rows = [];
      (catalog.categories || []).forEach((cat) => {
        (cat.stories || []).forEach((story) => {
          rows.push({
            stem: story.stem,
            title: story.title,
            paragraphs: story.paragraphs || [],
            categoryTitle: cat.title,
            categorySlug: cat.slug,
            hasAudio: !!story.hasAudio,
            hasImage: !!story.hasImage,
            hay: `${story.title || ""} ${(story.paragraphs || []).join(" ")}`.toLocaleLowerCase(LOCALE_TAG),
          });
        });
      });
      return rows;
    };

    const storiesScriptUrl =
      listPanel.getAttribute("data-stories-script") || `assets/stories-data.js?v=${assetVersion}`;

    const loadCatalogViaScript = () => {
      if (window.__BIRINCI_STORIES__) return Promise.resolve(window.__BIRINCI_STORIES__);
      return new Promise((resolve, reject) => {
        const script = document.createElement("script");
        script.src = storiesScriptUrl;
        script.async = true;
        script.onload = () => {
          if (window.__BIRINCI_STORIES__) resolve(window.__BIRINCI_STORIES__);
          else reject(new Error("empty-stories"));
        };
        script.onerror = () => reject(new Error("script-error"));
        document.head.appendChild(script);
      });
    };

    const ensureStories = () => {
      if (allStories) return Promise.resolve(allStories);
      if (loading) return loading;
      loading = Promise.resolve()
        .then(() => {
          if (window.__BIRINCI_STORIES__) return window.__BIRINCI_STORIES__;
          return loadCatalogViaScript().catch(() =>
            fetch(storiesUrl).then((res) => {
              if (!res.ok) throw new Error("fetch-failed");
              return res.json();
            })
          );
        })
        .then((catalog) => {
          window.__BIRINCI_STORIES__ = catalog;
          allStories = flattenStories(catalog);
          return allStories;
        })
        .catch(() => {
          allStories = [];
          return allStories;
        })
        .finally(() => {
          loading = null;
        });
      return loading;
    };

    const paragraphsHtml = (paragraphs, stem) => {
      if (!paragraphs.length) return "";
      const last = paragraphs.length - 1;
      const foldAzI = (s) => String(s || "").replace(/[İIı]/g, "i");
      const srcRe = /(internet\s+sources|internet\s+mənb|internet\s+kaynak|открыт\w*\s+источник|интернет|(?:source|mənbə|kaynak|источник|булак|булагы)\s*:)/i;
      const moralRe = /^(ibrət|ibret|moral|мораль|үлгү)\s*:/i;
      const authorSrcStems = { "everyone-has-work-to-do": 1, "weeds-must-be-pulled-from-the-root": 1, "the-silent-corridor": 1 };
      const authorSrc = !!(stem && authorSrcStems[stem]);
      const lastIsSrc = last >= 0 && (authorSrc || srcRe.test(foldAzI(paragraphs[last] || "")));
      const srcLabel = (I18N.ui && I18N.ui.story_source) || "";
      let moralI = -1;
      for (let j = lastIsSrc ? last - 1 : last; j >= 0; j--) {
        if (moralRe.test(foldAzI(String(paragraphs[j] || "").trim()))) {
          moralI = j;
          break;
        }
      }
      if (moralI < 0) moralI = lastIsSrc && last >= 1 ? last - 1 : last;
      return paragraphs
        .map((p, i) => {
          const isSrc = lastIsSrc && i === last;
          const cls = isSrc ? "story__source" : i === moralI ? "story__moral" : "";
          const text = isSrc && srcLabel && !authorSrc ? srcLabel : p;
          return `<p${cls ? ` class="${cls}"` : ""}>${escapeHtml(text)}</p>`;
        })
        .join("");
    };

    const storyArticleHtml = (story) => {
      const audioAttr = story.hasAudio
        ? ` data-audio="audio/${escapeHtml(story.stem)}.mp3?v=${escapeHtml(assetVersion)}"`
        : "";
      const audioLabel = escapeHtml(tUi("story_audio_label", "Səs"));
      const imageLabel = escapeHtml(tUi("story_image_label", "Şəkil"));
      const textLabel = escapeHtml(tUi("story_text_label", "Mətn"));
      const figureToggle = story.hasImage
        ? `
          <div class="story__action-group">
            <span class="tools-bar__label">${imageLabel}</span>
            <div class="tools-bar__views" role="group" aria-label="${imageLabel}">
            <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-images-mode="show" aria-pressed="true" aria-controls="figure-${escapeHtml(story.stem)}" title="${escapeHtml(tUi("show_image", "Şəkli göstər"))}" aria-label="${escapeHtml(tUi("show_image", "Şəkli göstər"))}">
              ${STORY_ICONS.eye}
            </button>
            <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-images-mode="hide" aria-pressed="false" aria-controls="figure-${escapeHtml(story.stem)}" title="${escapeHtml(tUi("hide_image", "Şəkli gizlət"))}" aria-label="${escapeHtml(tUi("hide_image", "Şəkli gizlət"))}">
              ${STORY_ICONS["eye-off"]}
            </button>
            </div>
          </div>`
        : "";
      const figureHtml = story.hasImage
        ? `
    <figure class="story__figure" id="figure-${escapeHtml(story.stem)}">
      <button type="button" class="story__figure-open" aria-label="${escapeHtml(story.title)} şəklini böyüt">
        <img src="illustrations/${escapeHtml(story.stem)}.webp" alt="${escapeHtml(story.title)} illüstrasiyası" loading="lazy" width="1536" height="1024" />
      </button>
    </figure>`
        : "";
      return `
<article class="story news-card" id="${escapeHtml(story.stem)}" data-stem="${escapeHtml(story.stem)}" data-title="${escapeHtml(story.title)}"${audioAttr}>
  <div class="card-header">
    <h2 class="card-title story__title">${escapeHtml(story.title)}</h2>
  </div>
  <div class="card-body">
    <div class="story__content">
      <div class="story__panel">
        <div class="story__actions">
          <div class="story__action-group">
            <span class="tools-bar__label">${audioLabel}</span>
            <div class="tools-bar__views" role="group" aria-label="${audioLabel}">
            <button type="button" class="story-tts tools-bar__view-btn tools-bar__view-btn--icon" data-story-tts data-tts-mode="listen" aria-pressed="false" title="${escapeHtml(tUi("listen", "Mətni dinlə"))}" aria-label="${escapeHtml(tUi("listen", "Mətni dinlə"))}">
              ${STORY_ICONS.listen}
            </button>
            <button type="button" class="story-tts tools-bar__view-btn tools-bar__view-btn--icon" data-story-tts data-tts-mode="stop" aria-pressed="true" title="${escapeHtml(tUi("stop", "Dayandır"))}" aria-label="${escapeHtml(tUi("stop", "Dayandır"))}">
              ${STORY_ICONS.stop}
            </button>
            </div>
          </div>
          ${figureToggle}
          <div class="story__action-group">
            <span class="tools-bar__label">${textLabel}</span>
            <div class="tools-bar__views" role="group" aria-label="${textLabel}">
            <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-texts-mode="show" aria-pressed="true" aria-controls="text-${escapeHtml(story.stem)}" title="${escapeHtml(tUi("show_text", "Mətni göstər"))}" aria-label="${escapeHtml(tUi("show_text", "Mətni göstər"))}">
              ${STORY_ICONS.text}
            </button>
            <button type="button" class="tools-bar__view-btn tools-bar__view-btn--icon" data-texts-mode="hide" aria-pressed="false" aria-controls="text-${escapeHtml(story.stem)}" title="${escapeHtml(tUi("hide_text", "Mətni gizlət"))}" aria-label="${escapeHtml(tUi("hide_text", "Mətni gizlət"))}">
              ${STORY_ICONS["text-off"]}
            </button>
            </div>
          </div>
          <p class="story-tts__note" data-story-tts-note hidden></p>
        </div>
        <div class="story__text card-text" id="text-${escapeHtml(story.stem)}">
          ${paragraphsHtml(story.paragraphs, story.stem)}
        </div>
      </div>
    </div>
    ${figureHtml}
  </div>
</article>`.trim();
    };

    const refreshSidebarNav = (items) => {
      if (navList) {
        navList.innerHTML = items
          .map(
            (s) =>
              `<li data-stem="${escapeHtml(s.stem)}" data-title="${escapeHtml(s.title)}"><a href="#${escapeHtml(
                s.stem
              )}">${escapeHtml(s.title)}</a></li>`
          )
          .join("");
      }
      if (typeof window.__birinciBindStorySidebar === "function") {
        const layout = listPanel.querySelector(".category-layout");
        if (layout) window.__birinciBindStorySidebar(layout);
      }
    };

    const bindHomeNav = () => {
      const layout = listPanel.querySelector(".category-layout");
      if (!layout) return;
      if (typeof window.__birinciBindStorySidebar === "function") {
        window.__birinciBindStorySidebar(layout);
      }
    };

    const renderList = ({ resetWindow = false } = {}) => {
      if (!storiesList) return;
      if (typeof window.__birinciStopStoryTts === "function") window.__birinciStopStoryTts();
      const q = searchInput.value.trim().toLocaleLowerCase(LOCALE_TAG);
      filtered = (allStories || []).filter((story) => !q || story.hay.includes(q));
      filtered.sort((a, b) => localeCompareAz(a.title, b.title));

      const total = filtered.length;
      syncSearchFilterUi(searchInput.value.trim(), total);
      const cap = Math.max(1, total || 1);
      let n = readBatchSize();
      if (!Number.isFinite(n) || n < 1) n = batchSize || 12;
      if (n > cap) n = cap;
      batchSize = n;
      if (batchSizeInput && String(batchSizeInput.value) !== String(n)) {
        batchSizeInput.value = String(n);
      }
      if (!allMode) {
        try {
          localStorage.setItem(batchSizeStorageKey, String(n));
        } catch (_) {}
      }

      if (resetWindow) {
        windowStart = 0;
        randomStems = null;
      }

      if (pendingStem) {
        const idx = filtered.findIndex(
          (story) => story.stem === pendingStem || story.stem === String(pendingStem)
        );
        if (idx >= 0) {
          randomStems = null;
          if (!allMode) {
            windowStart = Math.floor(idx / batchSize) * batchSize;
          }
        }
      }

      let visibleStories = [];
      if (total === 0) {
        windowStart = 0;
        randomStems = null;
        visibleStories = [];
      } else if (randomStems && randomStems.length) {
        allMode = false;
        const byStem = new Map(filtered.map((story) => [story.stem, story]));
        visibleStories = randomStems.map((stem) => byStem.get(stem)).filter(Boolean);
        if (!visibleStories.length) {
          randomStems = null;
          windowStart = 0;
          visibleStories = filtered.slice(0, batchSize);
        }
      } else if (allMode) {
        randomStems = null;
        windowStart = 0;
        visibleStories = filtered.slice();
      } else {
        randomStems = null;
        const maxStart = Math.max(0, total - 1);
        if (windowStart > maxStart) windowStart = Math.floor(maxStart / batchSize) * batchSize;
        if (windowStart < 0) windowStart = 0;
        visibleStories = filtered.slice(windowStart, windowStart + batchSize);
      }

      storiesList.innerHTML = visibleStories.map(storyArticleHtml).join("");
      if (typeof window.__birinciSetAllStoryFigures === "function") {
        window.__birinciSetAllStoryFigures(!document.body.classList.contains("images-collapsed"));
      }
      if (typeof window.__birinciSetAllStoryTexts === "function") {
        window.__birinciSetAllStoryTexts(!document.body.classList.contains("texts-collapsed"));
      }
      refreshSidebarNav(visibleStories);
      if (listEmpty) listEmpty.hidden = total !== 0;
      if (typeof window.__birinciClearListenQueue === "function") {
        window.__birinciClearListenQueue({ keepTrack: true });
      }
      syncBatchUi(visibleStories.length);
      persistAllMode();
      writeUrlState();
      if (pendingStem) {
        const el = document.getElementById(pendingStem);
        if (el) {
          window.requestAnimationFrame(() => {
            el.scrollIntoView({ block: "start", behavior: "auto" });
          });
        }
        pendingStem = null;
      }
    };

    const setHidden = (el, hide) => {
      if (!el) return;
      el.hidden = !!hide;
      if (hide) el.setAttribute("hidden", "");
      else el.removeAttribute("hidden");
    };

    const setView = (nextView, { persist = true } = {}) => {
      const prevView = view;
      view = nextView === "list" ? "list" : "cards";
      window.__birinciHomeView = view;
      // Panels first — never gated on fetch / history / TTS.
      setHidden(cardsPanel, view !== "cards");
      setHidden(listPanel, view !== "list");
      viewBtns.forEach((btn) => {
        btn.setAttribute("aria-pressed", btn.getAttribute("data-home-view") === view ? "true" : "false");
      });
      listOnly.forEach((el) => {
        setHidden(el, view !== "list");
      });
      if (persist) {
        try {
          localStorage.setItem(viewStorageKey, view);
        } catch (_) {}
      }
      if (view === "cards") {
        if (typeof window.__birinciStopStoryTts === "function") window.__birinciStopStoryTts();
        applyCards();
        writeUrlState();
        return;
      }
      writeUrlState();
      try {
        bindHomeNav();
      } catch (_) {}
      const scrollToTools = () => {
        if (typeof window.__birinciScrollHomeTools === "function") {
          window.__birinciScrollHomeTools();
          return;
        }
        window.scrollTo(0, 0);
      };
      ensureStories()
        .then(() => {
          renderList();
          if (prevView !== "list") scrollToTools();
        })
        .catch(() => {
          if (listEmpty) listEmpty.hidden = false;
          if (prevView !== "list") scrollToTools();
        });
      if (prevView !== "list") scrollToTools();
    };

    const onViewButton = (event) => {
      const btn = event.target.closest("[data-home-view]");
      if (!btn || !bar.contains(btn)) return;
      const next = btn.getAttribute("data-home-view");
      if (next !== "list" && next !== "cards") return;
      // Re-apply even if already selected so a stalled list can recover.
      pendingStem = null;
      setView(next);
    };
    viewBtns.forEach((btn) => {
      btn.addEventListener("click", (event) => {
        event.preventDefault();
        event.stopPropagation();
        onViewButton(event);
      });
    });
    bar.addEventListener("click", onViewButton);

    searchInput.addEventListener("input", () => {
      if (view === "cards") {
        applyCards();
        return;
      }
      pendingStem = null;
      renderList({ resetWindow: true });
    });
    if (batchSizeInput) {
      batchSizeInput.addEventListener("change", () => {
        commitBatchSize({ persist: true, render: true });
      });
      batchSizeInput.addEventListener("blur", () => {
        commitBatchSize({ persist: true, render: true });
      });
      batchSizeInput.addEventListener("keydown", (event) => {
        if (event.key !== "Enter") return;
        event.preventDefault();
        commitBatchSize({ persist: true, render: true });
        batchSizeInput.blur();
      });
    }
    const runBatchAction = (action) => {
      if (view !== "list") return;
      const total = filtered.length;
      const scrollHomeTools = () => {
        if (typeof window.__birinciScrollHomeTools === "function") {
          window.__birinciScrollHomeTools();
        }
      };
      if (!total) {
        randomStems = null;
        windowStart = 0;
        persistBatchSize();
        persistAllMode();
        renderList();
        return;
      }
      if (action === "all") {
        if (allMode && !randomStems) {
          allMode = false;
        } else {
          allMode = true;
          randomStems = null;
        }
        windowStart = 0;
        persistBatchSize();
        persistAllMode();
        pendingStem = null;
        renderList();
        scrollHomeTools();
        return;
      }
      if (action === "dec") {
        applyPageSize(batchSize - 1, { persist: true, render: true });
        scrollHomeTools();
        return;
      }
      if (action === "inc") {
        applyPageSize(batchSize + 1, { persist: true, render: true });
        scrollHomeTools();
        return;
      }
      if (action === "prev") {
        allMode = false;
        if (randomStems) {
          randomStems = null;
          windowStart = 0;
        } else {
          windowStart = Math.max(0, windowStart - batchSize);
        }
      } else if (action === "next") {
        allMode = false;
        randomStems = null;
        if (windowStart + batchSize < total) {
          windowStart += batchSize;
        }
      } else if (action === "random") {
        allMode = false;
        randomStems = pickRandomStems(batchSize);
      } else {
        return;
      }
      persistBatchSize();
      persistAllMode();
      pendingStem = null;
      renderList();
      scrollHomeTools();
    };
    if (batchDecBtn) batchDecBtn.addEventListener("click", () => runBatchAction("dec"));
    if (batchIncBtn) batchIncBtn.addEventListener("click", () => runBatchAction("inc"));
    if (batchPrevBtn) {
      batchPrevBtn.addEventListener("click", () => runBatchAction("prev"));
    }
    if (batchNextBtn) {
      batchNextBtn.addEventListener("click", () => runBatchAction("next"));
    }
    if (batchRandomBtn) {
      batchRandomBtn.addEventListener("click", () => runBatchAction("random"));
    }
    if (batchAllBtn) {
      batchAllBtn.addEventListener("click", () => runBatchAction("all"));
    }

    const urlState = readUrlState();
    let initialView = "cards";
    // Prefer view already chosen by the inline bootstrap (avoids reset race).
    if (window.__birinciHomeView === "list" || window.__birinciHomeView === "cards") {
      initialView = window.__birinciHomeView;
    } else if (urlState.view === "list" || urlState.view === "cards") {
      initialView = urlState.view;
    } else {
      try {
        const stored = localStorage.getItem(viewStorageKey);
        if (stored === "list" || stored === "cards") initialView = stored;
      } catch (_) {}
    }
    if (urlState.stem) {
      initialView = "list";
      pendingStem = urlState.stem;
    }
    if (urlState.q) searchInput.value = urlState.q;

    try {
      setView(initialView, { persist: false });
    } catch (_) {
      setHidden(cardsPanel, initialView !== "cards");
      setHidden(listPanel, initialView !== "list");
    }
  };

  const initStoryTts = () => {
    const unsupportedMessage =
      "Hörmətli oxucu, təəssüf ki, bu cihazda və ya brauzerdə səsə çevirmə (TTS) xidməti mövcud deyil. Zəhmət olmasa hekayəni oxuyaraq davam edin.";
    const noVoiceMessage =
      "Hörmətli oxucu, bu cihazda Azərbaycan nitq səsi tapılmadı.";
    const failedMessage =
      "Hörmətli oxucu, hazırda səsə çevirməni başlatmaq mümkün olmadı. Zəhmət olmasa bir az sonra yenidən cəhd edin və ya hekayəni oxuyun.";
    const audioFailedMessage =
      "Hörmətli oxucu, səs faylını oxumaq mümkün olmadı. Zəhmət olmasa bir az sonra yenidən cəhd edin.";

    const SPEED_STEPS = [0.75, 1, 1.25, 1.5, 1.75, 2];
    const SPEED_KEY = "birinci-audio-rate";
    const VOLUME_KEY = "birinci-audio-volume";
    const MUTE_KEY = "birinci-audio-muted";

    let activeBtn = null;
    let activeStem = "";
    let utterance = null;
    let audioPlayer = null;
    let suppressError = false;
    let ignoreClicksUntil = 0;
    let startGuardUntil = 0;
    let playerShell = null;
    let playerEls = null;
    let seeking = false;
    let playbackRate = 1;
    let savedVolume = 1;
    let savedMuted = false;
    let objectUrl = "";
    let activeSourceKey = "";
    let loadToken = 0;
    let speakToken = 0;
    let fetchController = null;
    let queueActive = false;
    let queueStems = [];
    let queueIndex = 0;

    const setLabel = (btn, text) => {
      if (!btn) return;
      btn.setAttribute("aria-label", text);
      btn.setAttribute("title", text);
    };

    const showNote = (btn, message) => {
      const root =
        (btn && btn.closest(".story__actions, .text-lightbox__tts")) ||
        (btn && btn.parentElement);
      const note = root && root.querySelector("[data-story-tts-note]");
      if (!note) return;
      note.hidden = !message;
      note.textContent = message || "";
    };

    const resolveStory = (btn) => {
      if (!btn) return null;
      const nested = btn.closest("article.story");
      if (nested) return nested;
      const stem = (btn.getAttribute("data-story-stem") || "").trim();
      if (!stem) return null;
      return (
        document.getElementById(stem) ||
        document.querySelector(`article.story[data-stem="${stem}"]`)
      );
    };

    const stemFor = (btn) => {
      if (!btn) return "";
      const fromAttr = (btn.getAttribute("data-story-stem") || "").trim();
      if (fromAttr) return fromAttr;
      const story = resolveStory(btn);
      return ((story && (story.dataset.stem || story.id)) || "").trim();
    };

    const titleFor = (btn, story) => {
      const fromStory = ((story && story.dataset.title) || "").trim();
      if (fromStory) return fromStory;
      const titleNode =
        story &&
        (story.querySelector(".story__title, .card-title") || story.querySelector("h2"));
      if (titleNode) return titleNode.textContent.trim();
      return "Hekayə";
    };

    const escapeStem = (stem) =>
      window.CSS && typeof window.CSS.escape === "function"
        ? window.CSS.escape(stem)
        : stem.replace(/\\/g, "\\\\").replace(/"/g, '\\"');

    const buttonsForStem = (stem, btn) => {
      const buttons = new Set();
      if (btn) buttons.add(btn);
      if (!stem) return buttons;
      const esc = escapeStem(stem);
      document
        .querySelectorAll(
          `[data-story-tts][data-story-stem="${esc}"], article.story[data-stem="${esc}"] [data-story-tts], article.story#${esc} [data-story-tts]`
        )
        .forEach((el) => buttons.add(el));
      return buttons;
    };

    const syncTtsPairUi = (btn, state) => {
      const stem = stemFor(btn) || activeStem;
      const roots = new Set();
      buttonsForStem(stem, btn).forEach((el) => {
        const root = el.closest(".tools-bar__views") || el.parentElement;
        if (root) roots.add(root);
      });
      const listenOn = state === "playing" || state === "paused";
      roots.forEach((root) => {
        root.querySelectorAll("[data-tts-mode]").forEach((el) => {
          const mode = el.getAttribute("data-tts-mode");
          const pressed = listenOn ? mode === "listen" : mode === "stop";
          el.setAttribute("aria-pressed", pressed ? "true" : "false");
          el.setAttribute("data-tts-state", state);
        });
      });
    };

    const syncPlayingUi = (btn, playing) => {
      syncTtsPairUi(btn, playing ? "playing" : "idle");
    };

    const syncPausedUi = (btn) => {
      syncTtsPairUi(btn, "paused");
    };

    const formatTime = (seconds) => {
      if (!Number.isFinite(seconds) || seconds < 0) return "0:00";
      const total = Math.floor(seconds);
      const m = Math.floor(total / 60);
      const s = total % 60;
      return `${m}:${String(s).padStart(2, "0")}`;
    };

    const readPrefs = () => {
      try {
        const rate = Number(localStorage.getItem(SPEED_KEY));
        if (SPEED_STEPS.includes(rate)) playbackRate = rate;
      } catch (_) {}
      try {
        const vol = Number(localStorage.getItem(VOLUME_KEY));
        if (Number.isFinite(vol) && vol >= 0 && vol <= 1) savedVolume = vol;
      } catch (_) {}
      try {
        savedMuted = localStorage.getItem(MUTE_KEY) === "1";
      } catch (_) {}
    };

    const writePrefs = () => {
      try {
        localStorage.setItem(SPEED_KEY, String(playbackRate));
        localStorage.setItem(VOLUME_KEY, String(savedVolume));
        localStorage.setItem(MUTE_KEY, savedMuted ? "1" : "0");
      } catch (_) {}
    };

    const updatePlayButton = (playing) => {
      if (!playerEls || !playerEls.playBtn) return;
      playerEls.playBtn.setAttribute("aria-label", playing ? "Fasilə" : "Oynat");
      playerEls.playBtn.setAttribute("aria-pressed", playing ? "true" : "false");
      playerEls.playBtn.innerHTML = playing
        ? '<svg viewBox="0 0 24 24" width="18" height="18" aria-hidden="true"><rect x="6" y="5" width="4" height="14" rx="1" fill="currentColor"></rect><rect x="14" y="5" width="4" height="14" rx="1" fill="currentColor"></rect></svg>'
        : '<svg viewBox="0 0 24 24" width="18" height="18" aria-hidden="true"><polygon points="8 5 20 12 8 19" fill="currentColor"></polygon></svg>';
    };

    const updateSpeedLabel = () => {
      if (!playerEls || !playerEls.speedBtns) return;
      playerEls.speedBtns.forEach((btn) => {
        const rate = Number(btn.getAttribute("data-speed"));
        const active = rate === playbackRate;
        btn.setAttribute("aria-pressed", active ? "true" : "false");
      });
    };

    const updateMuteUi = () => {
      if (!playerEls) return;
      if (playerEls.muteBtn) {
        playerEls.muteBtn.setAttribute("aria-pressed", savedMuted ? "true" : "false");
        playerEls.muteBtn.setAttribute("aria-label", savedMuted ? "Səsi aç" : "Səssiz");
        playerEls.muteBtn.innerHTML = savedMuted
          ? '<svg viewBox="0 0 24 24" width="16" height="16" fill="none" stroke="currentColor" stroke-width="2.1" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"></polygon><line x1="23" y1="9" x2="17" y2="15"></line><line x1="17" y1="9" x2="23" y2="15"></line></svg>'
          : '<svg viewBox="0 0 24 24" width="16" height="16" fill="none" stroke="currentColor" stroke-width="2.1" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"></polygon><path d="M15.54 8.46a5 5 0 0 1 0 7.07"></path><path d="M19.07 4.93a10 10 0 0 1 0 14.14"></path></svg>';
      }
      if (playerEls.volume) {
        playerEls.volume.value = String(savedMuted ? 0 : savedVolume);
      }
    };

    const updateProgressUi = () => {
      if (!audioPlayer || !playerEls) return;
      const duration = Number.isFinite(audioPlayer.duration) ? audioPlayer.duration : 0;
      const current = Number.isFinite(audioPlayer.currentTime) ? audioPlayer.currentTime : 0;
      if (playerEls.current) playerEls.current.textContent = formatTime(current);
      if (playerEls.duration) playerEls.duration.textContent = formatTime(duration);
      if (playerEls.seek && !seeking) {
        playerEls.seek.max = String(duration || 0);
        playerEls.seek.value = String(current || 0);
      }
    };

    const applyAudioSettings = () => {
      if (!audioPlayer) return;
      audioPlayer.playbackRate = playbackRate;
      audioPlayer.volume = savedVolume;
      audioPlayer.muted = savedMuted;
      updateSpeedLabel();
      updateMuteUi();
    };

    const syncAudioPlayerInset = () => {
      if (!playerShell || playerShell.hidden) {
        document.documentElement.style.removeProperty("--audio-player-h");
        return;
      }
      const h = Math.ceil(playerShell.getBoundingClientRect().height || 0);
      if (h > 0) {
        document.documentElement.style.setProperty("--audio-player-h", `${h}px`);
      }
    };

    const hidePlayerShell = () => {
      if (!playerShell) return;
      playerShell.hidden = true;
      playerShell.setAttribute("hidden", "");
      document.body.classList.remove("audio-player-open");
      syncAudioPlayerInset();
    };

    const showPlayerShell = () => {
      if (!playerShell) return;
      playerShell.hidden = false;
      playerShell.removeAttribute("hidden");
      document.body.classList.add("audio-player-open");
      syncAudioPlayerInset();
      window.requestAnimationFrame(syncAudioPlayerInset);
    };

    const clearActive = () => {
      if (activeBtn) syncPlayingUi(activeBtn, false);
      activeBtn = null;
      activeStem = "";
      utterance = null;
    };

    const revokeObjectUrl = () => {
      if (!objectUrl) return;
      try {
        URL.revokeObjectURL(objectUrl);
      } catch (_) {}
      objectUrl = "";
    };

    const stopAudioElement = ({ clearSrc = true } = {}) => {
      if (fetchController) {
        try {
          fetchController.abort();
        } catch (_) {}
        fetchController = null;
      }
      if (!audioPlayer) return;
      try {
        audioPlayer.pause();
      } catch (_) {}
      if (clearSrc) {
        try {
          audioPlayer.removeAttribute("src");
          audioPlayer.load();
        } catch (_) {}
        revokeObjectUrl();
        activeSourceKey = "";
      }
    };

    const stopCurrentMedia = () => {
      suppressError = true;
      loadToken += 1;
      speakToken += 1;
      stopAudioElement({ clearSrc: true });
      if (window.speechSynthesis) window.speechSynthesis.cancel();
      window.setTimeout(() => {
        suppressError = false;
      }, 120);
    };

    const clearQueue = ({ keepTrack = true } = {}) => {
      queueActive = false;
      queueStems = [];
      queueIndex = 0;
      if (!keepTrack) highlightPlaying(null);
      updateQueueChrome();
      syncPlayVisibleButton();
    };

    const idleProbeForStem = (stem) => {
      if (!stem) return null;
      const esc = escapeStem(stem);
      return document.querySelector(
        `[data-story-tts][data-story-stem="${esc}"], article.story[data-stem="${esc}"] [data-story-tts], article.story#${esc} [data-story-tts]`
      );
    };

    const closePlayer = () => {
      const btn = activeBtn;
      const stem = activeStem;
      clearQueue({ keepTrack: false });
      hidePlayerShell();
      stopCurrentMedia();
      const probe = btn || idleProbeForStem(stem);
      if (probe) syncTtsPairUi(probe, "idle");
      activeBtn = null;
      activeStem = "";
      utterance = null;
      updatePlayButton(false);
      updateProgressUi();
      syncPlayVisibleButton();
    };

    const stopSpeech = () => {
      closePlayer();
    };

    window.__birinciStopStoryTts = stopSpeech;
    window.__birinciIgnoreStoryTtsClicks = (ms) => {
      ignoreClicksUntil = Date.now() + Math.max(0, Number(ms) || 0);
    };
    window.__birinciSyncStoryTtsUi = (stem, playing) => {
      if (!stem) return;
      const probe = document.querySelector(
        `[data-story-tts][data-story-stem="${escapeStem(stem)}"], article.story[data-stem="${escapeStem(stem)}"] [data-story-tts], article.story#${escapeStem(stem)} [data-story-tts]`
      );
      if (probe) {
        if (playing) syncPlayingUi(probe, true);
        else if (window.__birinciIsStoryAudioActive(stem)) syncPausedUi(probe);
        else syncPlayingUi(probe, false);
      }
    };
    window.__birinciIsStoryAudioActive = (stem) =>
      !!(stem && activeStem === stem && playerShell && !playerShell.hidden);

    const isActivelyPlaying = () => {
      if (audioPlayer && !audioPlayer.paused && !audioPlayer.ended) return true;
      if (window.speechSynthesis && window.speechSynthesis.speaking) return true;
      return false;
    };

    const isSameStoryActive = (btn) => {
      if (!btn) return false;
      const stem = stemFor(btn);
      if (stem && activeStem) return stem === activeStem;
      if (activeBtn === btn) return true;
      if (activeBtn) {
        const a = activeBtn.closest(".tools-bar__views");
        const b = btn.closest(".tools-bar__views");
        if (a && b && a === b) return true;
      }
      return false;
    };

    const collectVisibleStems = () =>
      Array.from(document.querySelectorAll("article.story"))
        .filter((el) => !el.hidden && !el.closest("[hidden]"))
        .map((el) => (el.dataset.stem || el.id || "").trim())
        .filter(Boolean);

    const storyElForStem = (stem) => {
      if (!stem) return null;
      const esc = escapeStem(stem);
      return (
        document.getElementById(stem) ||
        document.querySelector(`article.story[data-stem="${esc}"]`)
      );
    };

    const listenBtnForStem = (stem) => {
      const story = storyElForStem(stem);
      return (
        (story && story.querySelector('[data-tts-mode="listen"]')) ||
        (story && story.querySelector("[data-story-tts]")) ||
        null
      );
    };

    const isPausedPlayback = () => {
      if (audioPlayer && audioPlayer.src && audioPlayer.paused && !audioPlayer.ended) return true;
      if (window.speechSynthesis && window.speechSynthesis.paused) return true;
      return false;
    };

    const highlightPlaying = (story, { scroll = false } = {}) => {
      document.querySelectorAll("article.story.story--playing").forEach((el) => {
        el.classList.remove("story--playing");
      });
      if (!story) return;
      story.classList.add("story--playing");
      if (scroll) {
        try {
          story.scrollIntoView({ block: "start", behavior: "smooth" });
        } catch (_) {}
      }
    };

    const syncPlayVisibleButton = () => {
      const stems = collectVisibleStems();
      const listenOn = queueActive && (isActivelyPlaying() || isPausedPlayback());
      const base = tUi("listen_page", "Səhifəni dinlə");
      const stop = tUi("stop", "Dayandır");
      const suffix = tUi("stories_count_suffix", "hekayə");
      const listenLabel = stems.length ? `${base} · ${stems.length} ${suffix}` : base;
      document.querySelectorAll("[data-tools-play-visible]").forEach((btn) => {
        btn.disabled = stems.length === 0;
        const mode = btn.getAttribute("data-tts-mode") || "listen";
        const pressed = listenOn ? mode === "listen" : mode === "stop";
        btn.setAttribute("aria-pressed", pressed ? "true" : "false");
        if (mode === "listen") {
          btn.title = listenLabel;
          btn.setAttribute("aria-label", listenLabel);
        } else {
          btn.title = stop;
          btn.setAttribute("aria-label", stop);
        }
      });
    };

    const updateQueueChrome = (title) => {
      if (!playerEls) return;
      const hasQueue = queueActive && queueStems.length > 0;
      if (playerShell) playerShell.classList.toggle("audio-player--queue", hasQueue);
      if (playerEls.storyPrev) {
        playerEls.storyPrev.hidden = !hasQueue;
        playerEls.storyPrev.disabled = !hasQueue || queueIndex <= 0;
      }
      if (playerEls.storyNext) {
        playerEls.storyNext.hidden = !hasQueue;
        playerEls.storyNext.disabled = !hasQueue || queueIndex >= queueStems.length - 1;
      }
      if (playerEls.title) {
        if (title) {
          playerEls.title.textContent = hasQueue
            ? `${queueIndex + 1} / ${queueStems.length}  ·  ${title}`
            : title;
        } else if (!hasQueue) {
          playerEls.title.textContent = String(playerEls.title.textContent || "").replace(
            /^\d+\s*\/\s*\d+\s*·\s*/,
            ""
          );
        }
      }
    };

    const sameVisibleQueue = () => {
      if (!queueActive) return false;
      const now = collectVisibleStems();
      return (
        now.length === queueStems.length && now.every((stem, i) => stem === queueStems[i])
      );
    };

    const pauseCurrent = () => {
      if (audioPlayer && audioPlayer.src && !audioPlayer.paused && !audioPlayer.ended) {
        audioPlayer.pause();
        return;
      }
      if (window.speechSynthesis && window.speechSynthesis.speaking && !window.speechSynthesis.paused) {
        window.speechSynthesis.pause();
        if (activeBtn) syncPausedUi(activeBtn);
        updatePlayButton(false);
      }
    };

    const resumeCurrent = () => {
      if (audioPlayer && audioPlayer.src && audioPlayer.paused) {
        const start = audioPlayer.play();
        if (start && typeof start.catch === "function") {
          start.catch(() => showNote(activeBtn, audioFailedMessage));
        }
        return;
      }
      if (window.speechSynthesis && window.speechSynthesis.paused) {
        window.speechSynthesis.resume();
        if (activeBtn) markPlaying(activeBtn, true);
      }
    };

    const loadVoices = () =>
      new Promise((resolve) => {
        if (!window.speechSynthesis) {
          resolve([]);
          return;
        }
        const current = () => window.speechSynthesis.getVoices() || [];
        const now = current();
        if (now.length) {
          resolve(now);
          return;
        }
        let done = false;
        const finish = () => {
          if (done) return;
          done = true;
          window.speechSynthesis.onvoiceschanged = null;
          resolve(current());
        };
        window.speechSynthesis.onvoiceschanged = finish;
        window.setTimeout(finish, 800);
      });

    const warmVoices = () => {
      if (!window.speechSynthesis) return;
      try {
        window.speechSynthesis.getVoices();
      } catch (_) {}
    };
    document.addEventListener("pointerdown", warmVoices, { once: true, passive: true });

    const pickVoice = (voices) => {
      const lang = String(LOCALE_TAG || "az").toLowerCase();
      const nameRe = {
        az: /azərbaycan|azerbaijani/i,
        en: /english/i,
        ru: /russian|русск/i,
        tr: /turkish|türk/i,
        ky: /kyrgyz|kirghiz|кыргыз/i,
      }[lang];
      const byLang = voices.find((v) => (v.lang || "").toLowerCase().startsWith(lang));
      const byName = nameRe ? voices.find((v) => nameRe.test(v.name || "")) : null;
      const turkicFallback =
        lang === "az" || lang === "ky"
          ? voices.find((v) => (v.lang || "").toLowerCase().startsWith("tr")) ||
            voices.find((v) => /turkish|türk/i.test(v.name || ""))
          : null;
      return byLang || byName || turkicFallback || null;
    };

    const textForSpeech = (story) => {
      const textEl = story && story.querySelector(".story__text");
      const title = ((story && story.dataset.title) || "").trim();
      const paras = textEl
        ? Array.from(textEl.querySelectorAll("p"))
            .map((p) => p.textContent.replace(/\s+/g, " ").trim())
            .filter(Boolean)
        : [];
      let body = paras.join(" ");
      body = body
        .replace(/[\u00AD\u200B-\u200D\uFEFF]/g, "")
        .replace(/[«»„“”]/g, "")
        .replace(/[‘’']/g, "")
        .replace(/[—–-]+\s*/g, "")
        .replace(/\s+/g, " ")
        .trim();
      if (!body) return title;
      if (title && body.toLocaleLowerCase(LOCALE_TAG).startsWith(title.toLocaleLowerCase(LOCALE_TAG))) {
        return body;
      }
      return title ? `${title}. ${body}` : body;
    };

    const resolveAudioUrl = (src) => {
      try {
        return new URL(src, document.baseURI).href;
      } catch (_) {
        return src;
      }
    };

    const markPlaying = (btn, playing = true) => {
      activeBtn = btn;
      activeStem = stemFor(btn);
      startGuardUntil = Date.now() + 450;
      syncPlayingUi(btn, playing);
      showNote(btn, "");
      updatePlayButton(playing);
    };

    const ensurePlayer = () => {
      if (playerShell && playerEls && audioPlayer) return playerEls;
      readPrefs();
      playerShell = document.createElement("div");
      playerShell.className = "audio-player";
      playerShell.hidden = true;
      playerShell.setAttribute("hidden", "");
      playerShell.setAttribute("role", "region");
      playerShell.setAttribute("aria-label", "Səs pleyeri");
      playerShell.innerHTML = `
        <div class="audio-player__inner">
          <div class="audio-player__meta">
            <p class="audio-player__title" data-audio-title>Hekayə</p>
          </div>
          <div class="audio-player__progress">
            <span class="audio-player__time" data-audio-current>0:00</span>
            <input class="audio-player__seek" data-audio-seek type="range" min="0" max="0" value="0" step="0.1" aria-label="İrəliləmə" />
            <span class="audio-player__time audio-player__time--duration" data-audio-duration>0:00</span>
          </div>
          <div class="audio-player__controls">
            <button type="button" class="audio-player__btn audio-player__btn--story" data-audio-story-prev hidden aria-label="${tUi("queue_prev", "Əvvəlki hekayə")}">⏮</button>
            <button type="button" class="audio-player__btn" data-audio-skip-back aria-label="15 saniyə geriyə">−15</button>
            <button type="button" class="audio-player__btn audio-player__btn--play" data-audio-play aria-label="Oynat" aria-pressed="false"></button>
            <button type="button" class="audio-player__btn" data-audio-skip-fwd aria-label="15 saniyə irəli">+15</button>
            <button type="button" class="audio-player__btn audio-player__btn--story" data-audio-story-next hidden aria-label="${tUi("queue_next", "Növbəti hekayə")}">⏭</button>
            <div class="audio-player__speed" data-audio-speed role="group" aria-label="Sürət">
              <span class="audio-player__speed-label">Sürət</span>
              <button type="button" class="audio-player__speed-btn" data-speed="0.75" aria-pressed="false">0.75×</button>
              <button type="button" class="audio-player__speed-btn" data-speed="1" aria-pressed="true">1×</button>
              <button type="button" class="audio-player__speed-btn" data-speed="1.25" aria-pressed="false">1.25×</button>
              <button type="button" class="audio-player__speed-btn" data-speed="1.5" aria-pressed="false">1.5×</button>
              <button type="button" class="audio-player__speed-btn" data-speed="1.75" aria-pressed="false">1.75×</button>
              <button type="button" class="audio-player__speed-btn" data-speed="2" aria-pressed="false">2×</button>
            </div>
            <div class="audio-player__volume-wrap">
              <button type="button" class="audio-player__btn" data-audio-mute aria-label="Səssiz" aria-pressed="false"></button>
              <input class="audio-player__volume" data-audio-volume type="range" min="0" max="1" value="1" step="0.01" aria-label="Səs səviyyəsi" />
            </div>
            <button type="button" class="audio-player__btn audio-player__btn--close" data-audio-close aria-label="Pleyeri bağla">&times;</button>
          </div>
        </div>
        <audio data-audio-el preload="auto" playsinline webkit-playsinline></audio>
      `.trim();
      document.body.appendChild(playerShell);

      audioPlayer = playerShell.querySelector("[data-audio-el]");
      playerEls = {
        title: playerShell.querySelector("[data-audio-title]"),
        current: playerShell.querySelector("[data-audio-current]"),
        duration: playerShell.querySelector("[data-audio-duration]"),
        seek: playerShell.querySelector("[data-audio-seek]"),
        playBtn: playerShell.querySelector("[data-audio-play]"),
        storyPrev: playerShell.querySelector("[data-audio-story-prev]"),
        storyNext: playerShell.querySelector("[data-audio-story-next]"),
        skipBack: playerShell.querySelector("[data-audio-skip-back]"),
        skipFwd: playerShell.querySelector("[data-audio-skip-fwd]"),
        speedGroup: playerShell.querySelector("[data-audio-speed]"),
        speedBtns: Array.from(playerShell.querySelectorAll("[data-speed]")),
        muteBtn: playerShell.querySelector("[data-audio-mute]"),
        volume: playerShell.querySelector("[data-audio-volume]"),
        closeBtn: playerShell.querySelector("[data-audio-close]"),
      };

      updatePlayButton(false);
      updateSpeedLabel();
      updateMuteUi();
      applyAudioSettings();

      if (typeof ResizeObserver !== "undefined") {
        const audioRo = new ResizeObserver(() => syncAudioPlayerInset());
        audioRo.observe(playerShell);
      }
      window.addEventListener("resize", syncAudioPlayerInset, { passive: true });
      if (window.visualViewport) {
        window.visualViewport.addEventListener("resize", syncAudioPlayerInset, { passive: true });
      }

      playerEls.playBtn.addEventListener("click", () => {
        if (utterance && window.speechSynthesis) {
          if (window.speechSynthesis.speaking && !window.speechSynthesis.paused) pauseCurrent();
          else if (window.speechSynthesis.paused) resumeCurrent();
          return;
        }
        if (!audioPlayer || !audioPlayer.src) return;
        if (audioPlayer.paused) {
          const start = audioPlayer.play();
          if (start && typeof start.catch === "function") {
            start.catch(() => showNote(activeBtn, audioFailedMessage));
          }
        } else {
          audioPlayer.pause();
        }
      });
      if (playerEls.storyPrev) {
        playerEls.storyPrev.addEventListener("click", () => {
          if (typeof window.__birinciQueuePrev === "function") window.__birinciQueuePrev();
        });
      }
      if (playerEls.storyNext) {
        playerEls.storyNext.addEventListener("click", () => {
          if (typeof window.__birinciQueueNext === "function") window.__birinciQueueNext();
        });
      }

      playerEls.skipBack.addEventListener("click", () => {
        if (!audioPlayer) return;
        audioPlayer.currentTime = Math.max(0, (audioPlayer.currentTime || 0) - 15);
        updateProgressUi();
      });

      playerEls.skipFwd.addEventListener("click", () => {
        if (!audioPlayer) return;
        const duration = Number.isFinite(audioPlayer.duration) ? audioPlayer.duration : 0;
        const next = (audioPlayer.currentTime || 0) + 15;
        audioPlayer.currentTime = duration ? Math.min(duration, next) : next;
        updateProgressUi();
      });

      playerEls.speedBtns.forEach((btn) => {
        btn.addEventListener("click", () => {
          const rate = Number(btn.getAttribute("data-speed"));
          if (!SPEED_STEPS.includes(rate)) return;
          playbackRate = rate;
          if (audioPlayer) audioPlayer.playbackRate = playbackRate;
          updateSpeedLabel();
          writePrefs();
        });
      });

      playerEls.muteBtn.addEventListener("click", () => {
        savedMuted = !savedMuted;
        if (audioPlayer) audioPlayer.muted = savedMuted;
        updateMuteUi();
        writePrefs();
      });

      playerEls.volume.addEventListener("input", () => {
        const next = Number(playerEls.volume.value);
        savedVolume = Number.isFinite(next) ? Math.min(1, Math.max(0, next)) : 1;
        savedMuted = savedVolume === 0;
        if (audioPlayer) {
          audioPlayer.volume = savedVolume;
          audioPlayer.muted = savedMuted;
        }
        updateMuteUi();
        writePrefs();
      });

      const onSeekInput = () => {
        seeking = true;
        if (playerEls.current) {
          playerEls.current.textContent = formatTime(Number(playerEls.seek.value) || 0);
        }
      };
      const onSeekCommit = () => {
        if (!audioPlayer) {
          seeking = false;
          return;
        }
        audioPlayer.currentTime = Number(playerEls.seek.value) || 0;
        seeking = false;
        updateProgressUi();
      };
      playerEls.seek.addEventListener("input", onSeekInput);
      playerEls.seek.addEventListener("change", onSeekCommit);
      playerEls.seek.addEventListener("pointerup", onSeekCommit);
      playerEls.seek.addEventListener("touchend", onSeekCommit);

      playerEls.closeBtn.addEventListener("click", () => closePlayer());

      audioPlayer.addEventListener("timeupdate", updateProgressUi);
      audioPlayer.addEventListener("loadedmetadata", updateProgressUi);
      audioPlayer.addEventListener("durationchange", updateProgressUi);
      audioPlayer.addEventListener("play", () => {
        if (activeBtn) markPlaying(activeBtn, true);
        else updatePlayButton(true);
        syncPlayVisibleButton();
      });
      audioPlayer.addEventListener("pause", () => {
        if (playerShell && playerShell.hidden) {
          updatePlayButton(false);
          syncPlayVisibleButton();
          return;
        }
        if (audioPlayer && !audioPlayer.ended && activeBtn) syncPausedUi(activeBtn);
        updatePlayButton(false);
        syncPlayVisibleButton();
      });
      audioPlayer.addEventListener("ended", () => {
        updatePlayButton(false);
        updateProgressUi();
        if (queueActive) {
          if (typeof window.__birinciQueueAdvance === "function") window.__birinciQueueAdvance();
          return;
        }
        if (playerShell && playerShell.hidden) {
          syncPlayVisibleButton();
          return;
        }
        if (activeBtn) syncPausedUi(activeBtn);
        syncPlayVisibleButton();
      });
      audioPlayer.addEventListener("error", () => {
        if (suppressError) return;
        const btn = activeBtn;
        if (queueActive) {
          if (btn) showNote(btn, audioFailedMessage);
          if (typeof window.__birinciQueueAdvance === "function") window.__birinciQueueAdvance();
          return;
        }
        closePlayer();
        if (btn) showNote(btn, audioFailedMessage);
      });

      return playerEls;
    };

    const startPlayback = (btn) => {
      applyAudioSettings();
      updateProgressUi();
      markPlaying(btn, true);
      const start = audioPlayer.play();
      if (start && typeof start.catch === "function") {
        start.catch(() => {
          if (queueActive) {
            showNote(btn, audioFailedMessage);
            advanceQueue();
            return;
          }
          closePlayer();
          showNote(btn, audioFailedMessage);
        });
      }
    };

    const openPlayer = ({ btn, src, title, stem }) => {
      ensurePlayer();
      if (window.speechSynthesis) window.speechSynthesis.cancel();
      utterance = null;

      const absolute = resolveAudioUrl(src);
      const sameTrack = activeSourceKey === absolute && !!audioPlayer && !!audioPlayer.src;

      activeBtn = btn;
      activeStem = stem || stemFor(btn);
      startGuardUntil = Date.now() + 450;
      showNote(btn, "");
      if (playerEls.title) {
        const rawTitle = title || "Hekayə";
        playerEls.title.textContent =
          queueActive && queueStems.length
            ? `${queueIndex + 1} / ${queueStems.length}  ·  ${rawTitle}`
            : rawTitle;
      }
      updateQueueChrome(title || "Hekayə");
      showPlayerShell();
      applyAudioSettings();

      if (sameTrack) {
        if (audioPlayer.paused || audioPlayer.ended) {
          if (audioPlayer.ended) audioPlayer.currentTime = 0;
          startPlayback(btn);
        } else {
          markPlaying(btn, true);
        }
        return;
      }

      const token = ++loadToken;
      suppressError = true;
      stopAudioElement({ clearSrc: true });
      window.setTimeout(() => {
        suppressError = false;
      }, 120);

      activeSourceKey = absolute;
      markPlaying(btn, true);
      updatePlayButton(true);

      // Load as a blob so seeking works even when the server lacks Range support.
      fetchController = typeof AbortController === "function" ? new AbortController() : null;
      const fetchOpts = fetchController ? { signal: fetchController.signal } : {};
      fetch(absolute, fetchOpts)
        .then((res) => {
          if (!res.ok) throw new Error("audio fetch failed");
          return res.blob();
        })
        .then((blob) => {
          if (token !== loadToken) return;
          revokeObjectUrl();
          objectUrl = URL.createObjectURL(blob);
          audioPlayer.src = objectUrl;
          startPlayback(btn);
        })
        .catch((err) => {
          if (fetchController && err && err.name === "AbortError") return;
          if (token !== loadToken) return;
          // Fallback: direct URL (may not seek on some local servers).
          try {
            audioPlayer.src = absolute;
            startPlayback(btn);
          } catch (_) {
            closePlayer();
            showNote(btn, audioFailedMessage);
          }
        });
    };

    const playAudioStory = (btn, src, story) => {
      openPlayer({
        btn,
        src,
        title: titleFor(btn, story),
        stem: stemFor(btn),
      });
    };

    const speakStory = async (btn, { fromQueue = false } = {}) => {
      if (!("speechSynthesis" in window) || typeof window.SpeechSynthesisUtterance !== "function") {
        showNote(btn, unsupportedMessage);
        if (fromQueue && queueActive && typeof window.__birinciQueueAdvance === "function") {
          window.__birinciQueueAdvance();
        }
        return;
      }

      const story = resolveStory(btn);
      const text = textForSpeech(story);
      if (!text) {
        showNote(btn, failedMessage);
        if (fromQueue && queueActive && typeof window.__birinciQueueAdvance === "function") {
          window.__birinciQueueAdvance();
        }
        return;
      }

      const voices = await loadVoices();
      const voice = pickVoice(voices);
      if (!voice) {
        if (fromQueue && queueActive) {
          showNote(btn, noVoiceMessage);
          if (typeof window.__birinciQueueAdvance === "function") window.__birinciQueueAdvance();
          return;
        }
        stopSpeech();
        showNote(btn, noVoiceMessage);
        return;
      }

      if (!fromQueue) {
        clearQueue({ keepTrack: false });
        closePlayer();
      } else {
        ensurePlayer();
        stopCurrentMedia();
        showPlayerShell();
        updateQueueChrome(titleFor(btn, story));
      }
      markPlaying(btn, true);
      syncPlayVisibleButton();

      const token = ++speakToken;
      const startSpeak = () => {
        if (token !== speakToken) return;
        utterance = new SpeechSynthesisUtterance(text);
        utterance.lang = (voice.lang || "az-AZ").startsWith("tr") ? "tr-TR" : "az-AZ";
        utterance.voice = voice;
        utterance.rate = 1;
        utterance.pitch = 1;

        utterance.onstart = () => {
          if (token !== speakToken) return;
          markPlaying(btn, true);
          syncPlayVisibleButton();
        };
        utterance.onend = () => {
          if (suppressError || token !== speakToken) return;
          if (fromQueue && queueActive && typeof window.__birinciQueueAdvance === "function") {
            window.__birinciQueueAdvance();
            return;
          }
          clearActive();
          hidePlayerShell();
          syncPlayVisibleButton();
        };
        utterance.onerror = () => {
          if (suppressError || token !== speakToken) return;
          if (fromQueue && queueActive) {
            showNote(btn, failedMessage);
            if (typeof window.__birinciQueueAdvance === "function") window.__birinciQueueAdvance();
            return;
          }
          clearActive();
          showNote(btn, failedMessage);
        };

        try {
          window.speechSynthesis.speak(utterance);
        } catch (err) {
          if (fromQueue && queueActive && typeof window.__birinciQueueAdvance === "function") {
            window.__birinciQueueAdvance();
            return;
          }
          clearActive();
          showNote(btn, unsupportedMessage);
        }
      };

      window.setTimeout(startSpeak, 60);
    };

    const playQueueIndex = (index, { scroll = false, skipCount = 0 } = {}) => {
      if (!queueActive || !queueStems.length) return;
      if (index < 0 || index >= queueStems.length || skipCount >= queueStems.length) {
        closePlayer();
        return;
      }
      queueIndex = index;
      const stem = queueStems[index];
      const story = storyElForStem(stem);
      const btn = listenBtnForStem(stem);
      if (!story || !btn) {
        playQueueIndex(index + 1, { scroll, skipCount: skipCount + 1 });
        return;
      }
      highlightPlaying(story, { scroll });
      updateQueueChrome(titleFor(btn, story));
      const audioSrc = story.dataset.audio;
      if (audioSrc) playAudioStory(btn, audioSrc, story);
      else speakStory(btn, { fromQueue: true });
      syncPlayVisibleButton();
    };

    const advanceQueue = () => {
      if (!queueActive) return;
      if (queueIndex + 1 >= queueStems.length) {
        closePlayer();
        return;
      }
      playQueueIndex(queueIndex + 1, { scroll: true });
    };

    const playVisible = () => {
      const stems = collectVisibleStems();
      if (!stems.length) return;
      if (sameVisibleQueue() && isActivelyPlaying()) {
        return;
      }
      if (sameVisibleQueue() && activeStem && !isActivelyPlaying()) {
        resumeCurrent();
        syncPlayVisibleButton();
        return;
      }
      queueStems = stems.slice();
      queueActive = true;
      queueIndex = 0;
      playQueueIndex(0, { scroll: true });
    };

    window.__birinciPlayVisible = playVisible;
    window.__birinciQueueAdvance = advanceQueue;
    window.__birinciQueuePrev = () => {
      if (!queueActive || queueIndex <= 0) return;
      playQueueIndex(queueIndex - 1, { scroll: true });
    };
    window.__birinciQueueNext = () => {
      if (!queueActive || queueIndex + 1 >= queueStems.length) return;
      playQueueIndex(queueIndex + 1, { scroll: true });
    };
    window.__birinciClearListenQueue = (opts) => clearQueue(opts || { keepTrack: true });
    window.__birinciSyncPlayVisibleUi = syncPlayVisibleButton;

    document.addEventListener("click", (event) => {
      const playVisibleBtn = event.target.closest("[data-tools-play-visible]");
      if (playVisibleBtn) {
        event.preventDefault();
        event.stopPropagation();
        const mode = playVisibleBtn.getAttribute("data-tts-mode") || "listen";
        if (mode === "stop") {
          stopSpeech();
          syncPlayVisibleButton();
          return;
        }
        playVisible();
        return;
      }
      const btn = event.target.closest("[data-story-tts]");
      if (!btn) return;
      event.preventDefault();
      event.stopPropagation();
      if (Date.now() < ignoreClicksUntil) return;

      const story = resolveStory(btn);
      const stem = stemFor(btn);
      const audioSrc = story && story.dataset.audio;
      const same = isSameStoryActive(btn);
      const mode = btn.getAttribute("data-tts-mode") || "listen";

      if (mode === "stop") {
        const root = btn.closest(".tools-bar__views") || btn.parentElement;
        const listenEl = root && root.querySelector('[data-tts-mode="listen"]');
        const pairOn =
          btn.getAttribute("data-tts-state") === "playing" ||
          btn.getAttribute("data-tts-state") === "paused" ||
          (listenEl && listenEl.getAttribute("aria-pressed") === "true");
        if (same || (stem && activeStem === stem) || pairOn) {
          stopSpeech();
        }
        syncTtsPairUi(btn, "idle");
        showNote(btn, "");
        return;
      }

      if (queueActive) {
        const idx = queueStems.indexOf(stem);
        if (idx >= 0) {
          if (same && isPausedPlayback()) {
            resumeCurrent();
            return;
          }
          if (same && isActivelyPlaying()) return;
          if (same && !isActivelyPlaying()) {
            resumeCurrent();
            return;
          }
          playQueueIndex(idx, { scroll: true });
          return;
        }
        clearQueue({ keepTrack: false });
      }

      if (audioSrc) {
        ensurePlayer();
        playAudioStory(btn, audioSrc, story);
        return;
      }

      if (same && isPausedPlayback()) {
        resumeCurrent();
        return;
      }
      if (same && isActivelyPlaying()) return;
      speakStory(btn);
    });

    document.addEventListener(
      "mouseleave",
      (event) => {
        const actions = event.target && event.target.closest && event.target.closest(".story__actions");
        if (!actions) return;
        const btn = actions.querySelector("[data-story-tts]");
        if (btn) showNote(btn, "");
      },
      true
    );

    document.addEventListener(
      "focusout",
      (event) => {
        const actions = event.target && event.target.closest && event.target.closest(".story__actions");
        if (!actions) return;
        if (!actions.contains(event.relatedTarget)) {
          const btn = actions.querySelector("[data-story-tts]");
          if (btn) showNote(btn, "");
        }
      },
      true
    );

    window.addEventListener("beforeunload", stopSpeech);
  };

  const initIllustrationLightbox = () => {
    let overlay = null;
    let dialog = null;
    let imageEl = null;
    let captionEl = null;
    let closeBtn = null;
    let lastFocus = null;

    const ensureOverlay = () => {
      if (overlay) return overlay;
      overlay = document.createElement("div");
      overlay.className = "illustration-lightbox";
      overlay.hidden = true;
      overlay.setAttribute("hidden", "");
      overlay.innerHTML = `
        <div class="illustration-lightbox__dialog" role="dialog" aria-modal="true" aria-label="Böyüdülmüş illüstrasiya">
          <button type="button" class="illustration-lightbox__close" aria-label="Bağla">&times;</button>
          <div class="illustration-lightbox__frame">
            <img class="illustration-lightbox__image" alt="" />
          </div>
          <p class="illustration-lightbox__caption"></p>
        </div>
      `.trim();
      document.body.appendChild(overlay);
      dialog = overlay.querySelector(".illustration-lightbox__dialog");
      imageEl = overlay.querySelector(".illustration-lightbox__image");
      captionEl = overlay.querySelector(".illustration-lightbox__caption");
      closeBtn = overlay.querySelector(".illustration-lightbox__close");

      overlay.addEventListener("click", (event) => {
        if (event.target === overlay) close();
      });
      if (closeBtn) closeBtn.addEventListener("click", close);
      return overlay;
    };

    const close = () => {
      if (!overlay || overlay.hidden) return;
      overlay.hidden = true;
      overlay.setAttribute("hidden", "");
      document.body.classList.remove("illustration-lightbox-open");
      if (imageEl) {
        imageEl.removeAttribute("src");
        imageEl.alt = "";
      }
      if (captionEl) captionEl.textContent = "";
      if (lastFocus && typeof lastFocus.focus === "function") {
        try {
          lastFocus.focus();
        } catch (_) {}
      }
      lastFocus = null;
    };

    const open = (img) => {
      if (!img || !img.getAttribute("src")) return;
      ensureOverlay();
      lastFocus = document.activeElement;
      const src = img.currentSrc || img.getAttribute("src");
      const alt = img.getAttribute("alt") || "";
      imageEl.src = src;
      imageEl.alt = alt;
      captionEl.textContent = alt;
      captionEl.hidden = !alt;
      overlay.hidden = false;
      overlay.removeAttribute("hidden");
      document.body.classList.add("illustration-lightbox-open");
      window.requestAnimationFrame(() => {
        if (closeBtn) closeBtn.focus();
      });
    };

    document.addEventListener("click", (event) => {
      const openBtn = event.target.closest(".story__figure-open");
      if (!openBtn) return;
      const story = openBtn.closest(".story");
      if (story && story.classList.contains("story--figure-hidden")) return;
      const img = openBtn.querySelector("img");
      if (!img) return;
      event.preventDefault();
      open(img);
    });

    document.addEventListener("keydown", (event) => {
      if (!overlay || overlay.hidden) return;
      if (event.key === "Escape") {
        event.preventDefault();
        close();
      }
    });
  };

  const initTextLightbox = () => {
    let overlay = null;
    let titleEl = null;
    let bodyEl = null;
    let closeBtn = null;
    let ttsBtn = null;
    let ttsBtns = [];
    let ttsNote = null;
    let lastFocus = null;

    const ensureOverlay = () => {
      if (overlay) return overlay;
      overlay = document.createElement("div");
      overlay.className = "text-lightbox";
      overlay.hidden = true;
      overlay.setAttribute("hidden", "");
      overlay.innerHTML = `
        <div class="text-lightbox__dialog" role="dialog" aria-modal="true" aria-label="Böyüdülmüş hekayə mətni">
          <button type="button" class="text-lightbox__close" aria-label="Bağla">&times;</button>
          <div class="text-lightbox__header">
            <h2 class="text-lightbox__title"></h2>
          </div>
          <div class="text-lightbox__tts">
            <div class="story__action-group">
              <span class="tools-bar__label">${tUi("story_audio_label", "Səs")}</span>
              <div class="tools-bar__views" role="group" aria-label="${tUi("story_audio_label", "Səs")}">
            <button type="button" class="story-tts tools-bar__view-btn tools-bar__view-btn--icon" data-story-tts data-lightbox-tts data-tts-mode="listen" aria-pressed="false" title="${tUi("listen", "Mətni dinlə")}" aria-label="${tUi("listen", "Mətni dinlə")}">
              ${STORY_ICONS.listen}
            </button>
            <button type="button" class="story-tts tools-bar__view-btn tools-bar__view-btn--icon" data-story-tts data-lightbox-tts data-tts-mode="stop" aria-pressed="true" title="${tUi("stop", "Dayandır")}" aria-label="${tUi("stop", "Dayandır")}">
              ${STORY_ICONS.stop}
            </button>
              </div>
            </div>
            <p class="story-tts__note" data-story-tts-note hidden></p>
          </div>
          <div class="text-lightbox__body"></div>
        </div>
      `.trim();
      document.body.appendChild(overlay);
      titleEl = overlay.querySelector(".text-lightbox__title");
      bodyEl = overlay.querySelector(".text-lightbox__body");
      closeBtn = overlay.querySelector(".text-lightbox__close");
      ttsBtns = Array.from(overlay.querySelectorAll("[data-lightbox-tts]"));
      ttsBtn =
        overlay.querySelector('[data-lightbox-tts][data-tts-mode="listen"]') || ttsBtns[0] || null;
      ttsNote = overlay.querySelector("[data-story-tts-note]");
      overlay.addEventListener("click", (event) => {
        if (event.target === overlay) close();
      });
      if (closeBtn) closeBtn.addEventListener("click", close);
      return overlay;
    };

    const resetTtsUi = () => {
      ttsBtns.forEach((el) => {
        el.removeAttribute("data-story-stem");
        el.setAttribute("data-tts-state", "idle");
        const mode = el.getAttribute("data-tts-mode");
        el.setAttribute("aria-pressed", mode === "stop" ? "true" : "false");
      });
      if (ttsNote) {
        ttsNote.hidden = true;
        ttsNote.textContent = "";
      }
    };

    const close = () => {
      if (!overlay || overlay.hidden) return;
      // Keep sticky audio player running while the text overlay closes.
      overlay.hidden = true;
      overlay.setAttribute("hidden", "");
      document.body.classList.remove("text-lightbox-open");
      if (titleEl) titleEl.textContent = "";
      if (bodyEl) bodyEl.innerHTML = "";
      resetTtsUi();
      if (lastFocus && typeof lastFocus.focus === "function") {
        try {
          lastFocus.focus();
        } catch (_) {}
      }
      lastFocus = null;
    };

    const open = (story, textEl) => {
      if (!story || !textEl) return;
      ensureOverlay();
      lastFocus = document.activeElement;
      const titleNode =
        story.querySelector(".story__title, .card-title") ||
        story.querySelector("h2");
      titleEl.textContent = titleNode ? titleNode.textContent.trim() : "Hekayə";
      bodyEl.innerHTML = textEl.innerHTML;
      const stem = ((story.dataset.stem || story.id) || "").trim();
      ttsBtns.forEach((el) => {
        if (stem) el.setAttribute("data-story-stem", stem);
        else el.removeAttribute("data-story-stem");
        el.setAttribute("data-tts-state", "idle");
        const mode = el.getAttribute("data-tts-mode");
        el.setAttribute("aria-pressed", mode === "stop" ? "true" : "false");
      });
      if (ttsNote) {
        ttsNote.hidden = true;
        ttsNote.textContent = "";
      }
      // Same tap that opens the overlay can land on the listen button (esp. mobile).
      if (typeof window.__birinciIgnoreStoryTtsClicks === "function") {
        window.__birinciIgnoreStoryTtsClicks(500);
      }
      overlay.hidden = false;
      overlay.removeAttribute("hidden");
      document.body.classList.add("text-lightbox-open");
      if (stem && typeof window.__birinciSyncStoryTtsUi === "function") {
        const cardBtn =
          story.querySelector('[data-tts-mode="listen"]') || story.querySelector("[data-story-tts]");
        const pressed = cardBtn && cardBtn.getAttribute("aria-pressed") === "true";
        const pausedLabel =
          cardBtn && cardBtn.getAttribute("data-tts-state") === "paused";
        const active =
          typeof window.__birinciIsStoryAudioActive === "function" &&
          window.__birinciIsStoryAudioActive(stem);
        if (pressed) window.__birinciSyncStoryTtsUi(stem, true);
        else if (active || pausedLabel) window.__birinciSyncStoryTtsUi(stem, false);
      }
      window.requestAnimationFrame(() => {
        // Focus close — not the listen button — to avoid ghost-click start/stop.
        if (closeBtn) closeBtn.focus();
        else if (overlay) overlay.focus();
      });
    };

    document.addEventListener("click", (event) => {
      if (event.target.closest("button, a, input, select, textarea, label")) return;
      const textEl = event.target.closest(".story__text, .story .card-text");
      if (!textEl) return;
      const story = textEl.closest("article.story");
      if (!story || story.classList.contains("story--text-hidden")) return;
      if (textEl.closest(".text-lightbox")) return;
      event.preventDefault();
      open(story, textEl);
    });

    document.addEventListener("keydown", (event) => {
      if (!overlay || overlay.hidden) return;
      if (event.key === "Escape") {
        event.preventDefault();
        close();
      }
    });
  };

  const initStoryFigureToggle = () => {
    const setFigureState = (story, visible) => {
      if (!story) return;
      story.classList.toggle("story--figure-hidden", !visible);
      setStoryModePressed(story, "data-images-mode", visible);
    };

    const setAllFigures = (visible) => {
      document.querySelectorAll("article.story").forEach((story) => {
        setFigureState(story, visible);
      });
    };

    window.__birinciSetStoryFigure = setFigureState;
    window.__birinciSetAllStoryFigures = setAllFigures;

    document.addEventListener("click", (event) => {
      const btn = event.target.closest("[data-images-mode]");
      if (!btn || btn.closest("[data-tools]")) return;
      const story = btn.closest("article.story");
      if (!story) return;
      event.preventDefault();
      setFigureState(story, btn.getAttribute("data-images-mode") === "show");
    });

    setAllFigures(!document.body.classList.contains("images-collapsed"));
  };

  const initStoryTextToggle = () => {
    const setTextState = (story, visible) => {
      if (!story) return;
      story.classList.toggle("story--text-hidden", !visible);
      setStoryModePressed(story, "data-texts-mode", visible);
    };

    const setAllTexts = (visible) => {
      document.querySelectorAll("article.story").forEach((story) => {
        setTextState(story, visible);
      });
    };

    window.__birinciSetStoryText = setTextState;
    window.__birinciSetAllStoryTexts = setAllTexts;

    document.addEventListener("click", (event) => {
      const btn = event.target.closest("[data-texts-mode]");
      if (!btn || btn.closest("[data-tools]")) return;
      const story = btn.closest("article.story");
      if (!story) return;
      event.preventDefault();
      setTextState(story, btn.getAttribute("data-texts-mode") === "show");
    });

    setAllTexts(!document.body.classList.contains("texts-collapsed"));
  };

  try {
    initIllustrationLightbox();
  } catch (err) {
    console.error("initIllustrationLightbox failed", err);
  }
  try {
    initTextLightbox();
  } catch (err) {
    console.error("initTextLightbox failed", err);
  }
  try {
    initStoryFigureToggle();
  } catch (err) {
    console.error("initStoryFigureToggle failed", err);
  }
  try {
    initStoryTextToggle();
  } catch (err) {
    console.error("initStoryTextToggle failed", err);
  }
  try {
    initHomeViews();
  } catch (err) {
    console.error("initHomeViews failed", err);
  }
  try {
    initStoryTts();
  } catch (err) {
    console.error("initStoryTts failed", err);
  }

  document.querySelectorAll(".category-layout").forEach((layout) => {
    try {
      bindStorySidebarLayout(layout);
    } catch (err) {
      console.error("bindStorySidebarLayout failed", err);
    }
  });
})();
"""


def build_one_locale() -> dict:
    """Build the currently configured LANG into LANG_ROOT. Returns catalog."""
    if stories_ready(LANG):
        print(f"[{LANG}] Extracting stories…")
    else:
        print(f"[{LANG}] Building placeholder pages (stories not ready)…")
    catalog = load_catalog()
    total = sum(c["count"] for c in catalog["categories"])
    print(f"[{LANG}] categories={len(catalog['categories'])} stories={total}")

    DATA_JSON.parent.mkdir(parents=True, exist_ok=True)
    slim = {
        "site_title": catalog["site_title"],
        "lang": LANG,
        "categories": [
            {
                "title": c["title"],
                "slug": c["slug"],
                "blurb": c["blurb"],
                "count": c["count"],
                "stories": [
                    {
                        "stem": s["stem"],
                        "title": s["title"],
                        "paragraphs": s["paragraphs"],
                        "hasAudio": (AUDIO_DIR / f"{s['stem']}.mp3").is_file(),
                        "hasImage": bool(s.get("has_image")),
                    }
                    for s in c["stories"]
                ],
            }
            for c in catalog["categories"]
        ],
    }
    DATA_JSON.write_text(json.dumps(slim, ensure_ascii=False, indent=2), encoding="utf-8")

    search_rows = []
    for c in catalog["categories"]:
        for s in c["stories"]:
            text = " ".join(s["paragraphs"])
            search_rows.append(
                {
                    "title": s["title"],
                    "stem": s["stem"],
                    "category": c["title"],
                    "slug": c["slug"],
                    "hay": f'{s["title"]} {c["title"]} {text}'.casefold(),
                }
            )
    search_js = (
        "window.__BIRINCI_SEARCH__ = "
        + json.dumps(search_rows, ensure_ascii=False)
        + ";\n"
    )
    stories_js = (
        "window.__BIRINCI_STORIES__ = "
        + json.dumps(slim, ensure_ascii=False)
        + ";\n"
    )

    js_i18n = dict(LOCALE.get("js") or {})
    for k, v in list(js_i18n.items()):
        if isinstance(v, str):
            js_i18n[k] = v.replace("{lang}", LANG)
    ui_i18n = dict(UI)
    i18n_boot = (
        "window.__BIRINCI_I18N__ = "
        + json.dumps({"lang": LANG, "ui": ui_i18n, "js": js_i18n}, ensure_ascii=False)
        + ";\n"
    )

    for folder in (STORIES, ILLUSTRATIONS, AUDIO_DIR, ASSETS):
        folder.mkdir(parents=True, exist_ok=True)
    sync_shared_assets()

    (ASSETS / "search-index.js").write_text(search_js, encoding="utf-8")
    (ASSETS / "stories-data.js").write_text(stories_js, encoding="utf-8")
    icons_js = (
        "window.__BIRINCI_STORY_ICONS__ = "
        + json.dumps(
            {
                "text": _tools_bar_glyph("text"),
                "text-off": _tools_bar_glyph("text-off"),
                "eye": _tools_bar_glyph("eye"),
                "eye-off": _tools_bar_glyph("eye-off"),
                "listen": _tools_bar_glyph("listen"),
                "stop": _tools_bar_glyph("stop"),
            },
            ensure_ascii=False,
        )
        + ";\n"
    )
    (ASSETS / "site.js").write_text(i18n_boot + icons_js + JS, encoding="utf-8")
    prune_locale_assets()

    PAGES_DIR.mkdir(parents=True, exist_ok=True)
    (LANG_ROOT / "index.html").write_text(build_landing(catalog), encoding="utf-8")

    for cat in catalog["categories"]:
        path = PAGES_DIR / f"{cat['slug']}.html"
        path.write_text(build_category_page(cat), encoding="utf-8")
        print(f"  [{LANG}] wrote {path.name} ({cat['count']})")

    print(f"[{LANG}] landing: {LANG_ROOT / 'index.html'}")
    print(f"[{LANG}] data: {DATA_JSON}")
    return catalog


def write_root_language_gate() -> None:
    (SITE_ROOT / "index.html").write_text(
        """<!DOCTYPE html>
<html lang="az">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1, viewport-fit=cover" />
  <meta name="theme-color" content="#0069b4" />
  <title>Bir inci</title>
  <meta name="description" content="Bir inci" />
  <link rel="preconnect" href="https://fonts.googleapis.com" />
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin />
  <link href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,500;9..144,700&family=Source+Sans+3:wght@400;600;700&display=swap" rel="stylesheet" />
  <style>
    :root{
      --nav-blue:#0069b4;
      --nav-blue-deep:#005a9a;
      --ink:#08263b;
      --font-display:"Fraunces",Georgia,serif;
      --font-ui:"Source Sans 3","Segoe UI","Noto Sans",sans-serif;
    }
    body{font-family:var(--font-ui);margin:0;min-height:100vh;min-height:100dvh;display:grid;place-items:center;
      padding:env(safe-area-inset-top) env(safe-area-inset-right) env(safe-area-inset-bottom) env(safe-area-inset-left);
      background:linear-gradient(165deg,#e8f4fc,#fff8e8);color:var(--ink)}
    .gate{text-align:center;padding:2rem}
    h1{font-family:var(--font-display);font-size:clamp(2rem,5vw,3rem);margin:0 0 .4rem}
    p{opacity:.8;margin:0 0 1.5rem}
    .langs{display:flex;gap:.75rem;justify-content:center;flex-wrap:wrap}
    a,.lang{display:inline-flex;align-items:center;gap:.55rem;padding:.7rem 1.1rem;border-radius:999px;
      background:var(--nav-blue);color:#fff;text-decoration:none;font-weight:700;min-width:7.5rem}
    a:hover{background:var(--nav-blue-deep)}
    .lang[aria-disabled="true"]{opacity:.42;cursor:not-allowed;background:#7aa7c9}
    img{width:1.35rem;height:.9rem;border-radius:2px;object-fit:cover;box-shadow:0 0 0 1px rgba(0,0,0,.12)}
  </style>
  <script>
    (function () {
      function enabledLangs(cfg) {
        return (cfg.languages || []).filter(function (lang) {
          return lang && lang.enabled !== false && lang.code && lang.show_in_switcher !== false;
        });
      }
      function matchesNav(nav, prefixes) {
        return (prefixes || []).some(function (p) { return nav.indexOf(p) === 0; });
      }
      fetch("languages.json")
        .then(function (res) { return res.json(); })
        .then(function (cfg) {
          var langs = enabledLangs(cfg);
          var ready = langs.filter(function (lang) { return lang.implemented !== false; });
          var readyCodes = ready.map(function (lang) { return lang.code; });
          try {
            var saved = localStorage.getItem("birinci-lang");
            if (saved && readyCodes.indexOf(saved) !== -1) {
              location.replace(saved + "/index.html");
              return;
            }
          } catch (e) {}
          var nav = (navigator.language || "").toLowerCase();
          var match = ready.find(function (lang) { return matchesNav(nav, lang.nav_prefixes); });
          if (match) {
            location.replace(match.code + "/index.html");
            return;
          }
          var prompt = document.querySelector("[data-gate-prompt]");
          if (prompt && cfg.gate_prompt) prompt.textContent = cfg.gate_prompt;
          var host = document.querySelector("[data-gate-langs]");
          if (!host) return;
          host.innerHTML = langs.map(function (lang) {
            var flag = lang.flag
              ? '<img src="' + lang.flag + '" alt="" width="22" height="15">'
              : "";
            var name = lang.name || lang.label || lang.code;
            if (lang.implemented === false) {
              return '<span class="lang" aria-disabled="true" title="Coming soon">' +
                flag + "<span>" + name + "</span></span>";
            }
            return '<a href="' + lang.code + '/index.html" hreflang="' + lang.code + '" data-lang="' + lang.code + '">' +
              flag + "<span>" + name + "</span></a>";
          }).join("");
          host.querySelectorAll("[data-lang]").forEach(function (link) {
            link.addEventListener("click", function () {
              try { localStorage.setItem("birinci-lang", link.getAttribute("data-lang") || ""); } catch (e) {}
            });
          });
        })
        .catch(function () {});
    })();
  </script>
</head>
<body>
  <div class="gate">
    <h1>Bir inci</h1>
    <p data-gate-prompt>Choose language</p>
    <div class="langs" data-gate-langs></div>
  </div>
</body>
</html>
""",
        encoding="utf-8",
    )
    print(f"root: {SITE_ROOT / 'index.html'}")


def main(argv: list[str] | None = None) -> int:
    import argparse

    parser = argparse.ArgumentParser(
        description=f"Build Bir inci static site ({'/'.join(SUPPORTED_LANGS)})."
    )
    parser.add_argument(
        "--lang",
        choices=[*SUPPORTED_LANGS, "all"],
        default="all",
        help="Locale to build (default: all)",
    )
    args = parser.parse_args(argv)

    langs = list(SUPPORTED_LANGS) if args.lang == "all" else [args.lang]
    for lang in langs:
        apply_locale(lang)
        build_one_locale()

    write_root_language_gate()
    return 0


if __name__ == "__main__":
    sys.exit(main())
